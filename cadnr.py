import json
import re
import tkinter as tk
import calendar
import base64
from datetime import date, datetime, timedelta
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
from urllib import error, parse, request
import gzip
import zipfile
import shutil
import filecmp
import subprocess
import tempfile
import html
import os
import unicodedata
import hashlib
import socket
import threading
import time
import webbrowser
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer

try:
    from PIL import Image, ImageTk
except ImportError:
    Image = None
    ImageTk = None


class RoundedButton(tk.Canvas):
    def __init__(self, master=None, **kwargs):
        self._text = kwargs.pop("text", "")
        self._command = kwargs.pop("command", None)
        self._state = kwargs.pop("state", "normal")
        estilo = str(kwargs.pop("style", "") or "").lower()
        self._variant = "success" if ("success" in estilo or "green" in estilo or "verde" in estilo) else "default"
        width_chars = kwargs.pop("width", None)
        self._takefocus = kwargs.pop("takefocus", True)
        kwargs.pop("padding", None)

        width_px = 104
        if width_chars is not None:
            try:
                width_px = max(38, int(width_chars) * 7 + 14)
            except (TypeError, ValueError):
                width_px = 104
        height_px = 30
        if self._text == "✓":
            width_px = 24
            height_px = 24

        super().__init__(
            master,
            width=width_px,
            height=height_px,
            highlightthickness=0,
            bd=0,
            takefocus=self._takefocus,
            bg=self._parent_bg(master),
            cursor="hand2" if self._state != "disabled" else "",
        )
        self._items = {}
        self._pressed = False
        self._hover = False
        self._draw()
        self.bind("<ButtonPress-1>", self._on_press)
        self.bind("<ButtonRelease-1>", self._on_release)
        self.bind("<Leave>", self._on_leave)
        self.bind("<Enter>", self._on_enter)

    @staticmethod
    def _parent_bg(master):
        try:
            return master.cget("background")
        except tk.TclError:
            return "#d9d9d9"

    def _palette(self):
        if self._variant == "success":
            if self._state == "disabled":
                return {
                    "fill": "#8fbf98",
                    "outline": "#7eaa87",
                    "text": "#e2f2e5",
                    "shadow": "#789f81",
                }
            if self._pressed:
                return {
                    "fill": "#2f8f4a",
                    "outline": "#297c40",
                    "text": "#effbf1",
                    "shadow": "#256e39",
                }
            if self._hover:
                return {
                    "fill": "#58b86f",
                    "outline": "#49a55f",
                    "text": "#f3fbf4",
                    "shadow": "#3f8e53",
                }
            return {
                "fill": "#49a85f",
                "outline": "#3f9252",
                "text": "#edf9ef",
                "shadow": "#347c45",
            }

        if self._state == "disabled":
            return {
                "fill": "#9db5cf",
                "outline": "#8ca2ba",
                "text": "#dce7f2",
                "shadow": "#8aa0b7",
            }
        if self._pressed:
            return {
                "fill": "#2f6498",
                "outline": "#2a5888",
                "text": "#eaf3ff",
                "shadow": "#234f7c",
            }
        if self._hover:
            return {
                "fill": "#6a9ed1",
                "outline": "#5a8fc2",
                "text": "#f2f8ff",
                "shadow": "#3d6d9d",
            }
        return {
            "fill": "#5c8fc3",
            "outline": "#4d7daf",
            "text": "#eaf3ff",
            "shadow": "#3a6694",
        }

    @staticmethod
    def _rounded_points(x1, y1, x2, y2, radius):
        return [
            x1 + radius, y1,
            x2 - radius, y1,
            x2, y1,
            x2, y1 + radius,
            x2, y2 - radius,
            x2, y2,
            x2 - radius, y2,
            x1 + radius, y2,
            x1, y2,
            x1, y2 - radius,
            x1, y1 + radius,
            x1, y1,
        ]

    def _draw(self):
        self.delete("all")
        w = max(2, int(self.winfo_reqwidth()))
        h = max(2, int(self.winfo_reqheight()))
        radius = max(1, int(h / 2))
        p = self._palette()
        # Sombra leve para dar profundidade.
        self.create_polygon(
            self._rounded_points(2, 3, w - 1, h - 1, radius),
            smooth=True,
            fill=p["shadow"],
            outline="",
        )
        shape = self.create_polygon(
            self._rounded_points(1, 1, w - 1, h - 1, radius),
            smooth=True,
            fill=p["fill"],
            outline=p["outline"],
            width=1,
        )
        label = self.create_text(
            w / 2,
            h / 2,
            text=self._text,
            fill=p["text"],
            font=("Segoe UI", 8, "bold"),
        )
        self._items["shape"] = shape
        self._items["label"] = label

    def _on_press(self, _event):
        if self._state == "disabled":
            return
        self._pressed = True
        self._draw()

    def _on_release(self, event):
        if self._state == "disabled":
            return
        was_pressed = self._pressed
        self._pressed = False
        self._draw()
        inside = 0 <= event.x <= self.winfo_width() and 0 <= event.y <= self.winfo_height()
        if was_pressed and inside and callable(self._command):
            self._command()

    def _on_leave(self, _event):
        self._hover = False
        if self._pressed:
            self._pressed = False
        self._draw()

    def _on_enter(self, _event):
        if self._state == "disabled":
            return
        self._hover = True
        self._draw()

    def configure(self, cnf=None, **kwargs):
        if cnf and isinstance(cnf, dict):
            kwargs.update(cnf)
        if "text" in kwargs:
            self._text = kwargs.pop("text")
        if "command" in kwargs:
            self._command = kwargs.pop("command")
        if "state" in kwargs:
            self._state = kwargs.pop("state")
            super().configure(cursor="hand2" if self._state != "disabled" else "")
        if "style" in kwargs:
            estilo = str(kwargs.pop("style") or "").lower()
            self._variant = "success" if ("success" in estilo or "green" in estilo or "verde" in estilo) else "default"
        if "width" in kwargs:
            try:
                width_px = max(38, int(kwargs.pop("width")) * 7 + 14)
                if self._text == "✓":
                    width_px = 24
                super().configure(width=width_px)
            except (TypeError, ValueError):
                kwargs.pop("width")
        if kwargs:
            super().configure(**kwargs)
        self._draw()

    config = configure

    def invoke(self):
        if self._state != "disabled" and callable(self._command):
            self._command()


ttk.Button = RoundedButton


TIPOS_LOGRADOURO = [
    "Rua",
    "Avenida",
    "Travessa",
    "Alameda",
    "Estrada",
    "Rodovia",
    "Praca",
    "Largo",
]

# Fallback local para uso sem internet.
UF_CIDADES = {
    "AC": ["Rio Branco", "Cruzeiro do Sul", "Sena Madureira"],
    "AL": ["Maceio", "Arapiraca", "Palmeira dos Indios"],
    "AP": ["Macapa", "Santana", "Laranjal do Jari"],
    "AM": ["Manaus", "Parintins", "Itacoatiara"],
    "BA": ["Salvador", "Feira de Santana", "Vitoria da Conquista"],
    "CE": ["Fortaleza", "Caucaia", "Juazeiro do Norte"],
    "DF": ["Brasilia", "Ceilandia", "Taguatinga"],
    "ES": ["Vitoria", "Vila Velha", "Serra"],
    "GO": ["Goiania", "Aparecida de Goiania", "Anapolis"],
    "MA": ["Sao Luis", "Imperatriz", "Caxias"],
    "MT": ["Cuiaba", "Varzea Grande", "Rondonopolis"],
    "MS": ["Campo Grande", "Dourados", "Tres Lagoas"],
    "MG": ["Belo Horizonte", "Uberlandia", "Contagem"],
    "PA": ["Belem", "Ananindeua", "Santarem"],
    "PB": ["Joao Pessoa", "Campina Grande", "Santa Rita"],
    "PR": ["Curitiba", "Londrina", "Maringa"],
    "PE": ["Recife", "Jaboatao dos Guararapes", "Olinda"],
    "PI": ["Teresina", "Parnaiba", "Picos"],
    "RJ": ["Rio de Janeiro", "Sao Goncalo", "Duque de Caxias"],
    "RN": ["Natal", "Mossoro", "Parnamirim"],
    "RS": ["Porto Alegre", "Caxias do Sul", "Pelotas"],
    "RO": ["Porto Velho", "Ji-Parana", "Ariquemes"],
    "RR": ["Boa Vista", "Rorainopolis", "Caracarai"],
    "SC": ["Florianopolis", "Joinville", "Blumenau"],
    "SP": ["Sao Paulo", "Guarulhos", "Campinas"],
    "SE": ["Aracaju", "Nossa Senhora do Socorro", "Lagarto"],
    "TO": ["Palmas", "Araguaina", "Gurupi"],
}

FUNCOES_CBO = {
    "Soldador": "724315",
    "Caldeireiro": "724410",
    "Maçariqueiro": "724310",
    "Motorista": "782510",
    "Montador de Andaimes": "715545",
    "Pintor": "716610",
    "Engenheiro Mecânico": "214405",
    "Operador de Empilhadeira": "782220",
    "Ajudante Geral": "717020",
    "Encarregado": "414210",
    "Montador Industrial": "725205",
    "Mecânico Montador": "725205",
}

OUTROS_DOCUMENTOS_TIPOS = [
    "Carteirinha",
    "Anuencia",
    "Fit Test",
    "ASO",
    "Ordem de Servico",
    "Contrato",
    "Ficha de EPI",
]

class CadastroPopup(tk.Toplevel):
    def __init__(self, master, titulo):
        super().__init__(master)
        self.title(titulo)
        self.transient(master)
        self.grab_set()
        self.columnconfigure(0, weight=1)
        self._scroll_canvas = None
        self._scroll_body = None
        self._scroll_window = None

    def habilitar_rolagem(self):
        if self._scroll_body is not None:
            return self._scroll_body

        self.rowconfigure(0, weight=1)
        self.columnconfigure(0, weight=1)
        canvas = tk.Canvas(self, highlightthickness=0, borderwidth=0)
        canvas.grid(row=0, column=0, sticky="nsew")
        barra = ttk.Scrollbar(self, orient="vertical", command=canvas.yview)
        barra.grid(row=0, column=1, sticky="ns")
        canvas.configure(yscrollcommand=barra.set)

        body = ttk.Frame(canvas)
        body.columnconfigure(0, weight=1)
        window_id = canvas.create_window((0, 0), window=body, anchor="nw")

        def _atualizar_scrollregion(_event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))

        def _ajustar_largura(_event=None):
            try:
                canvas.itemconfigure(window_id, width=canvas.winfo_width())
            except tk.TclError:
                pass

        def _wheel(event):
            delta = int(-1 * (event.delta / 120)) if event.delta else 0
            if delta:
                canvas.yview_scroll(delta, "units")
                return "break"
            return None

        def _wheel_linux_up(_event):
            canvas.yview_scroll(-1, "units")
            return "break"

        def _wheel_linux_down(_event):
            canvas.yview_scroll(1, "units")
            return "break"

        body.bind("<Configure>", _atualizar_scrollregion)
        canvas.bind("<Configure>", _ajustar_largura)
        canvas.bind("<MouseWheel>", _wheel, add="+")
        canvas.bind("<Button-4>", _wheel_linux_up, add="+")
        canvas.bind("<Button-5>", _wheel_linux_down, add="+")
        body.bind("<MouseWheel>", _wheel, add="+")
        body.bind("<Button-4>", _wheel_linux_up, add="+")
        body.bind("<Button-5>", _wheel_linux_down, add="+")
        self.bind("<MouseWheel>", _wheel, add="+")
        self.bind("<Button-4>", _wheel_linux_up, add="+")
        self.bind("<Button-5>", _wheel_linux_down, add="+")

        self._scroll_canvas = canvas
        self._scroll_body = body
        self._scroll_window = window_id
        return body

    def secao(self, texto, row, parent=None):
        host = parent if parent is not None else self
        lbl = ttk.Label(host, text=texto, font=("Segoe UI", 10, "bold"))
        lbl.grid(row=row, column=0, sticky="w", padx=12, pady=(10, 4))

    def ajustar_tamanho(self):
        self.update_idletasks()
        req_w = self.winfo_reqwidth() + 24
        req_h = self.winfo_reqheight() + 24
        if self._scroll_body is not None:
            req_w = max(req_w, self._scroll_body.winfo_reqwidth() + 44)
            req_h = max(req_h, self._scroll_body.winfo_reqheight() + 28)
        max_w = int(self.winfo_screenwidth() * 0.92)
        max_h = int(self.winfo_screenheight() * 0.9)
        width = min(req_w, max_w)
        height = min(req_h, max_h)
        self.geometry(f"{width}x{height}")
        self.minsize(max(560, min(req_w, width)), 420)

class CalendarioPopup(tk.Toplevel):
    MESES_PT = [
        "Janeiro",
        "Fevereiro",
        "Marco",
        "Abril",
        "Maio",
        "Junho",
        "Julho",
        "Agosto",
        "Setembro",
        "Outubro",
        "Novembro",
        "Dezembro",
    ]
    DIAS_PT = ["Seg", "Ter", "Qua", "Qui", "Sex", "Sab", "Dom"]

    def __init__(self, master, entry_data):
        super().__init__(master)
        self.entry_data = entry_data
        self.title("Selecionar data")
        self.resizable(False, False)
        self.transient(master)
        self.grab_set()

        data_inicial = self._parse_data_entry() or date.today()
        self.ano = data_inicial.year
        self.mes = data_inicial.month

        topo = ttk.Frame(self, padding=(10, 8, 10, 4))
        topo.grid(row=0, column=0, sticky="ew")
        topo.columnconfigure(1, weight=1)

        ttk.Button(topo, text="<", width=3, command=self._mes_anterior).grid(row=0, column=0, padx=(0, 8))
        self.lbl_mes = ttk.Label(topo, anchor="center", font=("Segoe UI", 10, "bold"))
        self.lbl_mes.grid(row=0, column=1, sticky="ew")
        ttk.Button(topo, text=">", width=3, command=self._mes_posterior).grid(row=0, column=2, padx=(8, 0))

        self.corpo = ttk.Frame(self, padding=(10, 4, 10, 10))
        self.corpo.grid(row=1, column=0, sticky="nsew")
        self._render_calendario()

    def _parse_data_entry(self):
        texto = self.entry_data.get().strip()
        if not texto:
            return None
        m = re.fullmatch(r"(\d{2})/(\d{2})/(\d{4})", texto)
        if not m:
            return None
        dia, mes, ano = map(int, m.groups())
        try:
            return date(ano, mes, dia)
        except ValueError:
            return None

    def _mes_anterior(self):
        if self.mes == 1:
            self.mes = 12
            self.ano -= 1
        else:
            self.mes -= 1
        self._render_calendario()

    def _mes_posterior(self):
        if self.mes == 12:
            self.mes = 1
            self.ano += 1
        else:
            self.mes += 1
        self._render_calendario()

    def _selecionar_dia(self, dia):
        self.entry_data.delete(0, tk.END)
        self.entry_data.insert(0, f"{dia:02d}/{self.mes:02d}/{self.ano:04d}")
        self.entry_data.event_generate("<<DateSelected>>")
        self.destroy()

    def _render_calendario(self):
        for w in self.corpo.winfo_children():
            w.destroy()

        self.lbl_mes.config(text=f"{self.MESES_PT[self.mes - 1]} {self.ano}")

        for i, nome_dia in enumerate(self.DIAS_PT):
            ttk.Label(self.corpo, text=nome_dia, anchor="center", width=4).grid(row=0, column=i, pady=(0, 4))

        semanas = calendar.monthcalendar(self.ano, self.mes)
        for linha, semana in enumerate(semanas, start=1):
            for col, dia in enumerate(semana):
                if dia == 0:
                    ttk.Label(self.corpo, text=" ", width=4).grid(row=linha, column=col)
                else:
                    ttk.Button(
                        self.corpo,
                        text=str(dia),
                        width=4,
                        command=lambda d=dia: self._selecionar_dia(d),
                    ).grid(row=linha, column=col, padx=1, pady=1)

def mascara_cpf(texto):
    nums = re.sub(r"\D", "", texto)[:11]
    if len(nums) <= 3:
        return nums
    if len(nums) <= 6:
        return f"{nums[:3]}.{nums[3:]}"
    if len(nums) <= 9:
        return f"{nums[:3]}.{nums[3:6]}.{nums[6:]}"
    return f"{nums[:3]}.{nums[3:6]}.{nums[6:9]}-{nums[9:]}"


def mascara_cnpj(texto):
    nums = re.sub(r"\D", "", texto)[:14]
    if len(nums) <= 2:
        return nums
    if len(nums) <= 5:
        return f"{nums[:2]}.{nums[2:]}"
    if len(nums) <= 8:
        return f"{nums[:2]}.{nums[2:5]}.{nums[5:]}"
    if len(nums) <= 12:
        return f"{nums[:2]}.{nums[2:5]}.{nums[5:8]}/{nums[8:]}"
    return f"{nums[:2]}.{nums[2:5]}.{nums[5:8]}/{nums[8:12]}-{nums[12:]}"


def mascara_celular(texto):
    nums = re.sub(r"\D", "", texto)[:11]
    if len(nums) <= 2:
        return nums
    if len(nums) <= 6:
        return f"({nums[:2]}) {nums[2:]}"
    if len(nums) <= 10:
        return f"({nums[:2]}) {nums[2:6]}-{nums[6:]}"
    return f"({nums[:2]}) {nums[2:7]}-{nums[7:]}"


def mascara_data(texto):
    nums = re.sub(r"\D", "", texto)[:8]
    if len(nums) <= 2:
        return nums
    if len(nums) <= 4:
        return f"{nums[:2]}/{nums[2:]}"
    return f"{nums[:2]}/{nums[2:4]}/{nums[4:]}"


def mascara_moeda_br(texto):
    nums = re.sub(r"\D", "", texto)
    if not nums:
        nums = "0"
    valor = int(nums)
    reais = valor // 100
    centavos = valor % 100
    reais_fmt = f"{reais:,}".replace(",", ".")
    return f"R$ {reais_fmt},{centavos:02d}"


class App(tk.Tk):
    def _configurar_estilo_abas(self):
        style = ttk.Style(self)
        temas_disponiveis = set(style.theme_names())
        for tema in ("clam", "vista", "xpnative", "default"):
            if tema in temas_disponiveis:
                try:
                    style.theme_use(tema)
                    break
                except tk.TclError:
                    continue

        fundo_sistema = style.lookup("TFrame", "background") or self.cget("background")
        # Paleta azul escuro + azul claro (sem branco).
        style.configure(
            "TNotebook",
            background=fundo_sistema,
            tabmargins=(6, 0, 6, 0),
            borderwidth=0,
        )
        style.configure(
            "TNotebook.Tab",
            background="#4e7fb3",
            foreground="#dceaff",
            padding=(14, 6),
            font=("Segoe UI", 10, "bold"),
            borderwidth=0,
            relief="flat",
        )
        style.map(
            "TNotebook.Tab",
            background=[
                ("selected", "#1f4f82"),
                ("active", "#3f76ad"),
                ("!active", "#4e7fb3"),
            ],
            foreground=[
                ("selected", "#e9f3ff"),
                ("active", "#e9f3ff"),
                ("!active", "#dceaff"),
            ],
            padding=[
                ("selected", (14, 6)),
                ("active", (14, 6)),
                ("!active", (14, 6)),
            ],
            borderwidth=[
                ("selected", 0),
                ("active", 0),
                ("!active", 0),
            ],
            relief=[
                ("selected", "flat"),
                ("active", "flat"),
                ("!active", "flat"),
            ],
        )
        style.configure(
            "TButton",
            padding=(12, 7),
            background="#5c8fc3",
            foreground="#e9f3ff",
            font=("Segoe UI", 9, "bold"),
            borderwidth=1,
            relief="flat",
        )
        style.map(
            "TButton",
            background=[
                ("pressed", "#2f6498"),
                ("active", "#467caf"),
                ("!disabled", "#5c8fc3"),
            ],
            foreground=[
                ("disabled", "#9cb7d3"),
                ("!disabled", "#e9f3ff"),
            ],
        )

    @staticmethod
    def _chave_ordenacao_nr(item):
        nome = str(item.get("nome", "") or "").strip()
        m = re.search(r"\bNR\s*(\d+)", nome, flags=re.IGNORECASE)
        if m:
            return (int(m.group(1)), nome.lower())
        return (10**9, nome.lower())

    @staticmethod
    def _nr_certificados_padrao():
        itens = [
            {"nome": "NR 1/5/7", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 6", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 11 Emp.", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 11 PR", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 11 Munck", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 11 Garra", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 11 (40)", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 12", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 12 Garra", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 18", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 18 PEMT", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 18 Guin", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 18 PTA", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 18 And", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 20 (08)", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 20 (16)", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 33", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 34", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
            {"nome": "NR 35", "coluna_1": "", "coluna_2": "", "reciclagem": False, "imprimir": False, "imprimir_adicionado": False},
        ]
        return sorted(itens, key=App._chave_ordenacao_nr)

    def __init__(self):
        super().__init__()
        self.title("CADNR - Sistema de Cadastro")
        self.geometry("860x460")
        self.minsize(780, 400)
        self.protocol("WM_DELETE_WINDOW", self._on_app_close)
        self._configurar_estilo_abas()
        self.cidades_cache = {}
        self.empresas = []
        self.funcionarios = []
        self.documentos = []
        self.documentos_salvos = []
        self.prox_empresa_id = 1
        self.prox_funcionario_id = 1
        self.prox_documento_id = 1
        self.main_empresa_ids = []
        self.main_funcionario_ids = []
        self.imprimir_empresa_ids = []
        self.imprimir_funcionario_ids = []
        self.outros_docs_empresa_ids = []
        self.outros_docs_check_vars = {}
        self.outros_docs_check_paths = {}
        self.outros_docs_imprimir = []
        self.cert_imprimir_empresa_ids = []
        self.cert_imprimir_funcionario_ids = []
        self.dados_path = Path(__file__).with_name("cadnr_dados.json")
        self.nr_certificados = self._nr_certificados_padrao()
        self.nr_certificados_widgets = []
        self.nr_filtradas_indices = []
        self._ultima_falha_conversao = ""
        self.assinatura_digital_habilitada = False
        self.assinatura_digital_pfx = ""
        self.assinatura_digital_senha = ""
        self.assinatura_digital_pfx1 = ""
        self.assinatura_digital_senha1 = ""
        self.assinatura_digital_img1 = ""
        self.assinatura_digital_pfx2 = ""
        self.assinatura_digital_senha2 = ""
        self.assinatura_digital_certificados = []
        self._aviso_assinatura_exibido = False
        self._aviso_git_sync_exibido = False
        self._aviso_desktop_sync_exibido = False
        self._aviso_git_auto_exibido = False
        self._qr_http_server = None
        self._qr_http_thread = None
        self._qr_http_base_url = ""
        self._qr_token_map = {}
        self._qr_github_cache = {}
        self._qr_github_ultimo_erro = ""
        self.github_repo = "Elizangela2805/documentos"
        self.github_branch = "main"
        self.github_dir = "_pdf_gerados"
        self.github_pages_base = "https://elizangela2805.github.io/documentos"
        self.github_token = ""
        self._github_nojekyll_ok = set()
        self._git_auto_commit_habilitado = True
        self._git_auto_commit_lock = threading.Lock()
        self._git_auto_commit_pendentes = set()
        self._git_auto_commit_thread = None
        self._docs_monitor_after_id = None
        self._docs_monitor_interval_ms = 3000
        self._docs_monitor_assinatura = ""

        container = ttk.Frame(self, padding=16)
        container.pack(fill="both", expand=True)

        notebook = ttk.Notebook(container)
        notebook.pack(fill="both", expand=True)
        self.notebook = notebook

        aba_cadnr = ttk.Frame(notebook, padding=18)
        aba_outros_documentos = ttk.Frame(notebook, padding=18)
        aba_imprimir = ttk.Frame(notebook, padding=18)
        aba_cadastros = ttk.Frame(notebook, padding=18)
        aba_engrenagem = ttk.Frame(notebook, padding=18)
        self.aba_cadnr = aba_cadnr
        self.aba_outros_documentos = aba_outros_documentos
        self.aba_imprimir = aba_imprimir
        self.aba_cadastros = aba_cadastros
        self.aba_engrenagem = aba_engrenagem

        notebook.add(aba_cadnr, text="CADNR")
        notebook.add(aba_outros_documentos, text="OUTROS DOCUMENTOS")
        notebook.add(aba_imprimir, text="IMPRIMIR")
        notebook.add(aba_cadastros, text="CADASTROS")
        notebook.add(aba_engrenagem, text="⚙")

        cert_imprimir_frame = ttk.Frame(aba_imprimir)
        cert_imprimir_frame.pack(anchor="w", fill="both", expand=True)
        cert_imprimir_frame.columnconfigure(0, weight=1)
        cert_imprimir_frame.rowconfigure(1, weight=1)
        cert_imprimir_frame.rowconfigure(4, weight=1)

        ttk.Label(cert_imprimir_frame, text="NR selecionadas em CERTIFICADOS:").grid(
            row=0, column=0, sticky="w", pady=(2, 4)
        )

        lista_frame = ttk.Frame(cert_imprimir_frame)
        lista_frame.grid(row=1, column=0, sticky="nsew")
        lista_frame.columnconfigure(0, weight=1)
        lista_frame.rowconfigure(0, weight=1)
        self.lista_nr_imprimir = tk.Listbox(lista_frame, height=12)
        self.lista_nr_imprimir.grid(row=0, column=0, sticky="nsew")
        self.lista_nr_imprimir.bind("<Delete>", self._excluir_nr_imprimir_teclado)
        self.lista_nr_imprimir.bind("<BackSpace>", self._excluir_nr_imprimir_teclado)
        lista_scroll = ttk.Scrollbar(
            lista_frame, orient="vertical", command=self.lista_nr_imprimir.yview
        )
        lista_scroll.grid(row=0, column=1, sticky="ns")
        self.lista_nr_imprimir.configure(yscrollcommand=lista_scroll.set)

        acoes_imprimir = ttk.Frame(cert_imprimir_frame)
        acoes_imprimir.grid(row=2, column=0, sticky="w", pady=(10, 0))
        ttk.Button(
            acoes_imprimir,
            text="Excluir",
            command=self._excluir_nr_imprimir_selecionada,
            width=14,
        ).grid(row=0, column=0)

        ttk.Label(cert_imprimir_frame, text="OUTROS DOCUMENTOS adicionados:").grid(
            row=3, column=0, sticky="w", pady=(14, 4)
        )
        lista_outros_frame = ttk.Frame(cert_imprimir_frame)
        lista_outros_frame.grid(row=4, column=0, sticky="nsew")
        lista_outros_frame.columnconfigure(0, weight=1)
        lista_outros_frame.rowconfigure(0, weight=1)
        self.lista_outros_docs_imprimir = tk.Listbox(lista_outros_frame, height=8)
        self.lista_outros_docs_imprimir.grid(row=0, column=0, sticky="nsew")
        lista_outros_scroll = ttk.Scrollbar(
            lista_outros_frame, orient="vertical", command=self.lista_outros_docs_imprimir.yview
        )
        lista_outros_scroll.grid(row=0, column=1, sticky="ns")
        self.lista_outros_docs_imprimir.configure(yscrollcommand=lista_outros_scroll.set)

        acoes_outros_imprimir = ttk.Frame(cert_imprimir_frame)
        acoes_outros_imprimir.grid(row=5, column=0, sticky="w", pady=(10, 0))
        ttk.Button(
            acoes_outros_imprimir,
            text="Excluir",
            command=self._excluir_outro_documento_imprimir_selecionado,
            width=14,
        ).grid(row=0, column=0)

        footer_imprimir = ttk.Frame(cert_imprimir_frame)
        footer_imprimir.grid(row=6, column=0, sticky="ew", pady=(12, 0))
        footer_imprimir.columnconfigure(1, weight=1)
        ttk.Button(
            footer_imprimir,
            text="Testar link QR",
            command=self._testar_link_qr_local,
            width=20,
        ).grid(row=0, column=0, sticky="w")
        ttk.Button(
            footer_imprimir,
            text="Abrir Site de Uploads",
            command=self._abrir_projetta_html_no_chrome,
            width=20,
        ).grid(row=0, column=1, sticky="w", padx=(8, 0))
        ttk.Button(
            footer_imprimir,
            text="Abrir Links PDF",
            command=self._abrir_links_pdfs_publicados,
            width=20,
        ).grid(row=0, column=2, sticky="w", padx=(8, 0))
        ttk.Button(
            footer_imprimir,
            text="Salvar Tudo",
            command=self._gerar_pdf_nr_imprimir,
            width=20,
        ).grid(row=0, column=3, sticky="e")

        engrenagem_wrap = ttk.Frame(aba_engrenagem)
        engrenagem_wrap.pack(fill="both", expand=True)
        engrenagem_wrap.columnconfigure(0, weight=1)
        engrenagem_wrap.rowconfigure(0, weight=1)

        botoes = ttk.Frame(engrenagem_wrap)
        botoes.grid(row=0, column=0)
        botoes.columnconfigure(0, weight=1)
        botoes.columnconfigure(1, weight=1)

        acoes_engrenagem = [
            ("CADASTRAR EMPRESA", self.abrir_cadastro_empresa),
            ("CADASTRAR FUNCIONARIO", self.abrir_cadastro_funcionario),
            ("CADASTRAR FUNCAO", self.abrir_cadastro_funcao),
            ("CADASTRAR DOCUMENTO", self.abrir_cadastro_documento),
            ("ASSINATURA DGT", self._abrir_configuracao_assinatura_digital),
            ("CONFIG SITE", self._abrir_configuracao_github),
        ]
        for idx, (txt, cmd) in enumerate(acoes_engrenagem):
            ttk.Button(
                botoes,
                text=txt,
                style="Success.TButton",
                command=cmd,
                width=24,
            ).grid(row=idx // 2, column=idx % 2, padx=10, pady=8, sticky="ew")

        filtros = ttk.Frame(aba_cadnr)
        filtros.pack(anchor="w", fill="x", pady=(16, 0))
        filtros.columnconfigure(1, weight=1)
        filtros.columnconfigure(3, weight=1)

        ttk.Label(filtros, text="Empresa:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        self.select_empresa = ttk.Combobox(filtros, state="readonly")
        self.select_empresa.grid(row=0, column=1, sticky="ew", padx=(0, 16))
        self.select_empresa.bind("<<ComboboxSelected>>", self._on_empresa_main_selected)

        ttk.Label(filtros, text="Funcionario:").grid(row=0, column=2, sticky="w", padx=(0, 8))
        self.select_funcionario = ttk.Combobox(filtros, state="readonly")
        self.select_funcionario.grid(row=0, column=3, sticky="ew")
        self.select_funcionario.bind("<<ComboboxSelected>>", self._on_funcionario_main_selected)

        nr_frame = ttk.Frame(aba_cadnr)
        nr_frame.pack(anchor="w", fill="both", expand=True, pady=(12, 0))
        nr_frame.columnconfigure(0, weight=1)
        nr_frame.rowconfigure(1, weight=1)
        ttk.Button(
            nr_frame,
            text="SELECIONAR TUDO",
            command=self._selecionar_tudo_nr,
            width=18,
        ).grid(row=0, column=0, sticky="w", pady=(2, 6))
        ttk.Button(
            nr_frame,
            text="DESMARCAR TUDO",
            command=lambda: self._desmarcar_tudo_nr(limpar_datas=True),
            width=18,
        ).grid(row=0, column=0, sticky="w", padx=(170, 0), pady=(2, 6))
        nr_scroll_frame = ttk.Frame(nr_frame)
        nr_scroll_frame.grid(row=1, column=0, sticky="nsew")
        nr_scroll_frame.columnconfigure(0, weight=1)
        nr_scroll_frame.rowconfigure(0, weight=1)

        fundo_interface = ttk.Style(self).lookup("TFrame", "background") or self.cget("background")
        self.nr_canvas = tk.Canvas(
            nr_scroll_frame,
            highlightthickness=0,
            height=300,
            bg=fundo_interface,
        )
        self.nr_canvas.grid(row=0, column=0, sticky="nsew")
        self.nr_scrollbar = ttk.Scrollbar(
            nr_scroll_frame,
            orient="vertical",
            command=self.nr_canvas.yview,
        )
        self.nr_scrollbar.grid(row=0, column=1, sticky="ns")
        self.nr_canvas.configure(yscrollcommand=self.nr_scrollbar.set)

        self.nr_campos_frame = tk.Frame(self.nr_canvas, bg=fundo_interface)
        self.nr_canvas_window = self.nr_canvas.create_window(
            (0, 0),
            window=self.nr_campos_frame,
            anchor="nw",
        )
        self.nr_campos_frame.bind("<Configure>", self._on_nr_frame_configure)
        self.nr_canvas.bind("<Configure>", self._on_nr_canvas_configure)
        self.nr_canvas.bind("<Enter>", self._ativar_scroll_mousewheel_nr)
        self.nr_canvas.bind("<Leave>", self._desativar_scroll_mousewheel_nr)
        ttk.Button(
            nr_frame,
            text="ADICIONAR",
            command=self._adicionar_nr_imprimir,
            width=14,
        ).grid(row=2, column=0, sticky="w", pady=(10, 2))

        self._construir_aba_outros_documentos()

        imprimir_frame = ttk.Frame(aba_cadastros)
        imprimir_frame.pack(anchor="w", fill="x")
        imprimir_frame.columnconfigure(1, weight=1)
        imprimir_frame.columnconfigure(3, weight=1)

        ttk.Label(imprimir_frame, text="Empresa:").grid(row=0, column=0, sticky="w", padx=(0, 8), pady=4)
        self.imprimir_select_empresa = ttk.Combobox(imprimir_frame, state="readonly")
        self.imprimir_select_empresa.grid(row=0, column=1, sticky="ew", padx=(0, 16), pady=4)
        self.imprimir_select_empresa.bind("<<ComboboxSelected>>", self._on_empresa_imprimir_selected)

        ttk.Label(imprimir_frame, text="Funcionario:").grid(row=0, column=2, sticky="w", padx=(0, 8), pady=4)
        self.imprimir_select_funcionario = ttk.Combobox(imprimir_frame, state="readonly")
        self.imprimir_select_funcionario.grid(row=0, column=3, sticky="ew", pady=4)
        self.imprimir_select_funcionario.bind(
            "<<ComboboxSelected>>", self._on_funcionario_imprimir_selected
        )

        ttk.Button(
            imprimir_frame,
            text="EDITAR EMPRESA",
            command=self._abrir_edicao_empresa_imprimir,
            width=24,
        ).grid(row=1, column=1, sticky="w", pady=(10, 0))

        ttk.Button(
            imprimir_frame,
            text="EXCLUIR EMPRESA",
            command=self._excluir_empresa_imprimir,
            width=24,
        ).grid(row=2, column=1, sticky="w", pady=(8, 0))

        ttk.Button(
            imprimir_frame,
            text="EDITAR FUNCIONARIO",
            command=self._abrir_edicao_funcionario_imprimir,
            width=24,
        ).grid(row=1, column=3, sticky="w", pady=(10, 0))

        ttk.Button(
            imprimir_frame,
            text="EXCLUIR FUNCIONARIO",
            command=self._excluir_funcionario_imprimir,
            width=24,
        ).grid(row=2, column=3, sticky="w", pady=(8, 0))

        self._carregar_dados()
        self._desmarcar_tudo_nr(salvar=False, atualizar_ui=False)
        self._atualizar_select_empresas(limpar_nr=False)
        self.after(120, self._pos_inicializacao_pesada)
        self.after(800, self._sincronizar_documentos_salvos_pendentes)
        self.after(1200, self._monitorar_documentos_projeto)

    def _pos_inicializacao_pesada(self):
        alterou_nr = False
        try:
            self._migrar_pastas_documentos_legadas()
        except Exception:
            pass
        try:
            alterou_nr = bool(self._limpar_nr_nao_usadas_no_projeto())
        except Exception:
            alterou_nr = False
        try:
            self._garantir_pastas_empresas()
        except Exception:
            pass
        if alterou_nr:
            try:
                self._salvar_dados()
            except Exception:
                pass
        empresa_id = self._empresa_id_selecionada_main()
        self._aplicar_filtro_nr_por_empresa(empresa_id, limpar_nr=False)

    def _empresa_id_selecionada_main(self):
        idx = self.select_empresa.current()
        if idx < 0 or idx >= len(self.main_empresa_ids):
            return None
        return self.main_empresa_ids[idx]

    def _sincronizar_empresa_entre_abas(self, empresa_id, origem="main"):
        if empresa_id is None:
            return
        if origem == "main":
            if empresa_id in self.imprimir_empresa_ids:
                idx_imp = self.imprimir_empresa_ids.index(empresa_id)
                if self.imprimir_select_empresa.current() != idx_imp:
                    self.imprimir_select_empresa.current(idx_imp)
                    self._atualizar_imprimir_funcionarios(empresa_id)
        else:
            if empresa_id in self.main_empresa_ids:
                idx_main = self.main_empresa_ids.index(empresa_id)
                if self.select_empresa.current() != idx_main:
                    self.select_empresa.current(idx_main)
                    self._atualizar_select_funcionarios(empresa_id)
                    self._aplicar_filtro_nr_por_empresa(empresa_id, limpar_nr=False)
                    self._carregar_outros_documentos_empresa_selecionada()

    def _sincronizar_funcionario_entre_abas(self, funcionario_id, origem="main"):
        if funcionario_id is None:
            return
        if origem == "main":
            if funcionario_id in self.imprimir_funcionario_ids:
                idx_imp = self.imprimir_funcionario_ids.index(funcionario_id)
                if self.imprimir_select_funcionario.current() != idx_imp:
                    self.imprimir_select_funcionario.current(idx_imp)
        else:
            if funcionario_id in self.main_funcionario_ids:
                idx_main = self.main_funcionario_ids.index(funcionario_id)
                if self.select_funcionario.current() != idx_main:
                    self.select_funcionario.current(idx_main)
                    self._carregar_outros_documentos_empresa_selecionada()

    def _on_empresa_main_selected(self, _event=None):
        empresa_id = self._empresa_id_selecionada_main()
        self._atualizar_select_funcionarios(empresa_id)
        self._sincronizar_empresa_entre_abas(empresa_id, origem="main")
        self._aplicar_filtro_nr_por_empresa(empresa_id)
        self._carregar_outros_documentos_empresa_selecionada()

    def _on_funcionario_main_selected(self, _event=None):
        self._sincronizar_funcionario_entre_abas(self._funcionario_id_selecionado_main(), origem="main")
        self._atualizar_preview_nr()
        self._carregar_outros_documentos_empresa_selecionada()

    @staticmethod
    def _empresa_label(empresa):
        nome_pasta = str(empresa.get("nome_pasta", "") or "").strip()
        if nome_pasta:
            return nome_pasta
        return str(empresa.get("nome", "") or "")

    def _empresa_do_funcionario(self, funcionario):
        if not isinstance(funcionario, dict):
            return None
        empresa_id = funcionario.get("empresa_id")
        if isinstance(empresa_id, int):
            empresa = next((e for e in self.empresas if e.get("id") == empresa_id), None)
            if empresa is not None:
                return empresa
        nome_pasta_func = str(funcionario.get("nome_pasta", "") or "").strip()
        if nome_pasta_func:
            nome_norm = nome_pasta_func.casefold()
            for empresa in self.empresas:
                nome_pasta_emp = str(empresa.get("nome_pasta", "") or "").strip().casefold()
                nome_emp = str(empresa.get("nome", "") or "").strip().casefold()
                if nome_norm and (nome_norm == nome_pasta_emp or nome_norm == nome_emp):
                    return empresa
        return None

    def _funcionario_pertence_empresa(self, funcionario, empresa_id):
        empresa = self._empresa_do_funcionario(funcionario)
        return bool(empresa and empresa.get("id") == empresa_id)

    def _pastas_candidatas_empresa(self, empresa):
        candidatos = []
        if isinstance(empresa, dict):
            nomes_base = []
            nome_pasta = str(empresa.get("nome_pasta", "") or "").strip()
            nome_empresa = str(empresa.get("nome", "") or "").strip()
            if nome_pasta:
                nomes_base.append(nome_pasta)
            if nome_empresa and nome_empresa not in nomes_base:
                nomes_base.append(nome_empresa)
            for nome_ref in nomes_base:
                if nome_ref and nome_ref not in candidatos:
                    candidatos.append(nome_ref)
                nome_seguro = self._obter_nome_arquivo_seguro(nome_ref, "")
                if nome_seguro and nome_seguro not in candidatos:
                    candidatos.append(nome_seguro)
        if not candidatos:
            candidatos.append("projetta")
        # Compatibilidade de leitura com estruturas antigas.
        for legado in ("certificados", "documentos"):
            if legado not in candidatos:
                candidatos.append(legado)
        return candidatos

    def _pasta_referencia_empresa(self, empresa):
        if not isinstance(empresa, dict):
            return ""
        base_dir = Path(__file__).resolve().parent
        candidatos = self._pastas_candidatas_empresa(empresa)
        for candidato in candidatos:
            if not candidato:
                continue
            pasta_candidata = base_dir / candidato
            if pasta_candidata.exists() and pasta_candidata.is_dir():
                return candidato
        return next((c for c in candidatos if c), "")

    def _sincronizar_vinculo_funcionarios_empresas(self):
        for funcionario in self.funcionarios:
            if not isinstance(funcionario, dict):
                continue
            empresa = self._empresa_do_funcionario(funcionario)
            if not empresa:
                continue
            funcionario["empresa_id"] = empresa["id"]
            pasta_ref = self._pasta_referencia_empresa(empresa)
            if pasta_ref:
                funcionario["nome_pasta"] = pasta_ref

    @staticmethod
    def _nome_nr_do_arquivo(nome_arquivo, nome_pasta):
        nome = str(nome_arquivo or "").strip()
        if nome_pasta:
            padrao_pasta = re.escape(str(nome_pasta).strip())
            nome = re.sub(rf"\s*\(+\s*{padrao_pasta}\s*\)+\s*$", "", nome, flags=re.IGNORECASE)
        nome = re.sub(r"\s*[\(\)]+\s*$", "", nome).strip()
        nome = re.sub(r"\s+", " ", nome)
        if re.match(r"(?i)^nr\s*", nome):
            resto = re.sub(r"(?i)^nr\s*", "", nome).strip()
            nome = f"NR {resto}" if resto else "NR"
        return nome

    def _nomes_nr_na_pasta_empresa(self, empresa_id):
        if empresa_id is None:
            return None
        empresa = next((e for e in self.empresas if e.get("id") == empresa_id), None)
        if not empresa:
            return {}
        candidatos = self._pastas_candidatas_empresa(empresa)
        if not candidatos:
            return {}
        pasta_base = Path(__file__).resolve().parent
        pasta = None
        nome_referencia = ""
        for candidato in candidatos:
            pasta_candidata = pasta_base / candidato
            if pasta_candidata.exists() and pasta_candidata.is_dir():
                pasta = pasta_candidata
                nome_referencia = candidato
                break
        if pasta is None:
            return {}
        nomes = {}
        for caminho in pasta.iterdir():
            if not caminho.is_file():
                continue
            if caminho.suffix.lower() not in {".docx", ".doc", ".pdf"}:
                continue
            if not self._arquivo_nr_vinculado_empresa(caminho, empresa, nome_referencia):
                continue
            nome_exibicao = self._nome_nr_do_arquivo(caminho.stem, nome_referencia)
            if not re.match(r"(?i)^nr\b", str(nome_exibicao or "").strip()):
                continue
            chave = self._normalizar_nome_nr(nome_exibicao)
            if not chave:
                continue
            nomes[chave] = nome_exibicao
        return nomes

    def _aplicar_filtro_nr_por_empresa(self, empresa_id, limpar_nr=True):
        if limpar_nr:
            self._limpar_nr_nao_usadas_no_projeto()
        nomes_pasta = self._nomes_nr_na_pasta_empresa(empresa_id)
        if nomes_pasta is None:
            self.nr_filtradas_indices = []
        else:
            chaves_existentes = set()
            for item in self.nr_certificados:
                chave = self._normalizar_nome_nr(item.get("nome", ""))
                if chave in nomes_pasta:
                    item["nome"] = nomes_pasta[chave]
                if chave:
                    chaves_existentes.add(chave)
            for chave, nome_exibicao in nomes_pasta.items():
                if chave in chaves_existentes:
                    continue
                self.nr_certificados.append(
                    {
                        "nome": nome_exibicao,
                        "coluna_1": "",
                        "coluna_2": "",
                        "reciclagem": False,
                        "imprimir": False,
                        "imprimir_adicionado": False,
                    }
                )
            self.nr_certificados.sort(key=self._chave_ordenacao_nr)
            self.nr_filtradas_indices = [
                idx
                for idx, item in enumerate(self.nr_certificados)
                if self._normalizar_nome_nr(item.get("nome", "")) in nomes_pasta
            ]
        self._render_campos_nr()
        self._atualizar_lista_nr_imprimir()

    def _mapa_nr_em_uso_no_projeto(self):
        mapa = {}
        for empresa in self.empresas:
            if not isinstance(empresa, dict):
                continue
            empresa_id = empresa.get("id")
            if not isinstance(empresa_id, int):
                continue
            nomes_empresa = self._nomes_nr_na_pasta_empresa(empresa_id) or {}
            for chave, nome in nomes_empresa.items():
                chave_norm = self._normalizar_nome_nr(chave)
                if not chave_norm:
                    continue
                if chave_norm not in mapa:
                    mapa[chave_norm] = str(nome or "").strip() or str(chave or "").strip()
        return mapa

    def _limpar_nr_nao_usadas_no_projeto(self):
        mapa_em_uso = self._mapa_nr_em_uso_no_projeto()
        if not mapa_em_uso:
            return False

        nr_unicos = {}
        for item in self.nr_certificados:
            if not isinstance(item, dict):
                continue
            chave = self._normalizar_nome_nr(self._nome_nr_canonico(item.get("nome", "")))
            if not chave or chave not in mapa_em_uso:
                continue
            nome_base = mapa_em_uso.get(chave, str(item.get("nome", "") or "").strip())
            atual = nr_unicos.get(chave)
            if atual is None:
                nr_unicos[chave] = {
                    "nome": nome_base,
                    "coluna_1": str(item.get("coluna_1", "") or ""),
                    "coluna_2": str(item.get("coluna_2", "") or ""),
                    "dias": self._duracao_nr_item(item),
                    "reciclagem": bool(item.get("reciclagem", False)),
                    "imprimir": bool(item.get("imprimir", False)),
                    "imprimir_adicionado": bool(item.get("imprimir_adicionado", False)),
                }
                continue
            if not atual.get("coluna_1") and str(item.get("coluna_1", "") or "").strip():
                atual["coluna_1"] = str(item.get("coluna_1", "") or "")
            if not atual.get("coluna_2") and str(item.get("coluna_2", "") or "").strip():
                atual["coluna_2"] = str(item.get("coluna_2", "") or "")
            atual["dias"] = max(int(atual.get("dias", 1) or 1), self._duracao_nr_item(item))
            atual["reciclagem"] = bool(atual.get("reciclagem", False) or item.get("reciclagem", False))
            atual["imprimir"] = bool(atual.get("imprimir", False) or item.get("imprimir", False))
            atual["imprimir_adicionado"] = bool(
                atual.get("imprimir_adicionado", False) or item.get("imprimir_adicionado", False)
            )

        for chave, nome in mapa_em_uso.items():
            if chave in nr_unicos:
                continue
            nr_unicos[chave] = {
                "nome": nome,
                "coluna_1": "",
                "coluna_2": "",
                "dias": self._duracao_nr_dias(nome),
                "reciclagem": False,
                "imprimir": False,
                "imprimir_adicionado": False,
            }

        lista_limpa = sorted(nr_unicos.values(), key=self._chave_ordenacao_nr)
        if len(lista_limpa) != len(self.nr_certificados):
            self.nr_certificados = lista_limpa
            return True
        antigas = [self._normalizar_nome_nr(item.get("nome", "")) for item in self.nr_certificados]
        novas = [self._normalizar_nome_nr(item.get("nome", "")) for item in lista_limpa]
        if antigas != novas:
            self.nr_certificados = lista_limpa
            return True
        return False

    def _empresa_id_selecionada_cert_imprimir(self):
        idx = self.cert_imprimir_select_empresa.current()
        if idx < 0 or idx >= len(self.cert_imprimir_empresa_ids):
            return None
        return self.cert_imprimir_empresa_ids[idx]

    def _funcionario_id_selecionado_cert_imprimir(self):
        idx = self.cert_imprimir_select_funcionario.current()
        if idx < 0 or idx >= len(self.cert_imprimir_funcionario_ids):
            return None
        return self.cert_imprimir_funcionario_ids[idx]

    def _on_empresa_cert_imprimir_selected(self, _event=None):
        self._atualizar_cert_imprimir_funcionarios(self._empresa_id_selecionada_cert_imprimir())

    def _on_funcionario_cert_imprimir_selected(self, _event=None):
        self._atualizar_lista_nr_imprimir()

    def _atualizar_cert_imprimir_empresas(self, preservar_id=None):
        self.cert_imprimir_empresa_ids = [None] + [empresa["id"] for empresa in self.empresas]
        self.cert_imprimir_select_empresa["values"] = [""] + [
            self._empresa_label(empresa) for empresa in self.empresas
        ]

        if preservar_id in self.cert_imprimir_empresa_ids:
            idx = self.cert_imprimir_empresa_ids.index(preservar_id)
        else:
            idx = 0
        self.cert_imprimir_select_empresa.current(idx)
        self._atualizar_cert_imprimir_funcionarios(self.cert_imprimir_empresa_ids[idx])

    def _atualizar_cert_imprimir_funcionarios(self, empresa_id):
        if empresa_id is None:
            self.cert_imprimir_funcionario_ids = [None]
            self.cert_imprimir_select_funcionario["values"] = [""]
            self.cert_imprimir_select_funcionario.current(0)
            self._atualizar_lista_nr_imprimir()
            return

        funcs = [f for f in self.funcionarios if self._funcionario_pertence_empresa(f, empresa_id)]
        self.cert_imprimir_funcionario_ids = [None] + [f["id"] for f in funcs]
        self.cert_imprimir_select_funcionario["values"] = [""] + [f["nome"] for f in funcs]
        if funcs:
            self.cert_imprimir_select_funcionario.current(1)
        else:
            self.cert_imprimir_select_funcionario.current(0)
        self._atualizar_lista_nr_imprimir()

    def _linhas_nr_selecionadas(self):
        funcionario = self._funcionario_ativo_para_documento()
        empresa_ref = self._empresa_do_funcionario(funcionario) if funcionario else None
        empresa_id_func = empresa_ref.get("id") if empresa_ref else None
        chaves_empresa = None
        if empresa_id_func is not None:
            nomes_pasta = self._nomes_nr_na_pasta_empresa(empresa_id_func) or {}
            chaves_empresa = set(nomes_pasta.keys())

        linhas = []
        for item in self.nr_certificados:
            if not bool(item.get("imprimir_adicionado", False)):
                continue
            nome = str(item.get("nome", "") or "").strip()
            if chaves_empresa is not None:
                if self._normalizar_nome_nr(nome) not in chaves_empresa:
                    continue
            de = str(item.get("coluna_1", "") or "").strip()
            ate = str(item.get("coluna_2", "") or "").strip()
            if de and ate:
                linhas.append(f"{nome} - De: {de}  Ate: {ate}")
            elif ate:
                linhas.append(f"{nome} - Ate: {ate}")
            elif de:
                linhas.append(f"{nome} - De: {de}")
            else:
                linhas.append(nome)
        return linhas

    def _limpar_areas_selecao_salvar_tudo(self):
        for item in self.nr_certificados:
            item["imprimir_adicionado"] = False
            item["imprimir"] = False
        self._atualizar_lista_nr_imprimir()

        self.outros_docs_imprimir = []
        self._atualizar_lista_outros_docs_imprimir()

        if hasattr(self, "select_empresa"):
            self.select_empresa.current(0)
            self._atualizar_select_funcionarios(None)
            self._aplicar_filtro_nr_por_empresa(None)
            self._carregar_outros_documentos_empresa_selecionada()

        if hasattr(self, "imprimir_select_empresa"):
            self.imprimir_select_empresa.current(0)
            self._atualizar_imprimir_funcionarios(None)

        if hasattr(self, "cert_imprimir_select_empresa"):
            self.cert_imprimir_select_empresa.current(0)
            self._atualizar_cert_imprimir_funcionarios(None)

        self._atualizar_preview_nr()

    def _indices_nr_imprimir_adicionados(self):
        return [
            idx
            for idx, item in enumerate(self.nr_certificados)
            if bool(item.get("imprimir_adicionado", False))
        ]

    def _atualizar_lista_nr_imprimir(self):
        if not hasattr(self, "lista_nr_imprimir"):
            return
        self.lista_nr_imprimir.delete(0, tk.END)
        linhas = self._linhas_nr_selecionadas()
        for linha in linhas:
            self.lista_nr_imprimir.insert(tk.END, linha)

    def _excluir_nr_imprimir_selecionada(self):
        if not hasattr(self, "lista_nr_imprimir"):
            return
        selecao = self.lista_nr_imprimir.curselection()
        if not selecao:
            messagebox.showwarning("IMPRIMIR", "Selecione um documento na lista para excluir.")
            return
        idx_lista = int(selecao[0])
        indices_adicionados = self._indices_nr_imprimir_adicionados()
        if idx_lista < 0 or idx_lista >= len(indices_adicionados):
            return
        idx_nr = indices_adicionados[idx_lista]
        self.nr_certificados[idx_nr]["imprimir_adicionado"] = False
        self._atualizar_lista_nr_imprimir()
        self._salvar_dados()

    def _excluir_nr_imprimir_teclado(self, _event=None):
        self._excluir_nr_imprimir_selecionada()
        return "break"

    def _render_campos_nr(self):
        for widget in self.nr_campos_frame.winfo_children():
            widget.destroy()
        self.nr_certificados_widgets = [None] * len(self.nr_certificados)

        if self._empresa_id_selecionada_main() is None:
            ttk.Label(
                self.nr_campos_frame,
                text="Selecione uma empresa para visualizar as NRs.",
            ).grid(row=0, column=0, sticky="w", pady=(6, 0))
            self.nr_canvas.update_idletasks()
            self.nr_canvas.configure(scrollregion=self.nr_canvas.bbox("all"))
            self.focus_set()
            return

        ttk.Label(self.nr_campos_frame, text="Selecionar").grid(row=0, column=0, padx=(0, 4), sticky="w")
        ttk.Label(self.nr_campos_frame, text="").grid(row=0, column=1, padx=(0, 8))
        ttk.Label(self.nr_campos_frame, text="De").grid(row=0, column=2, sticky="w", padx=(0, 8))
        ttk.Label(self.nr_campos_frame, text="Até").grid(row=0, column=4, sticky="w")
        ttk.Label(self.nr_campos_frame, text="Dias").grid(row=0, column=6, sticky="w", padx=(8, 8))
        ttk.Label(self.nr_campos_frame, text="Reciclagem").grid(row=0, column=7, sticky="w", padx=(10, 0))

        indices_visiveis = (
            list(range(len(self.nr_certificados)))
            if self.nr_filtradas_indices is None
            else list(self.nr_filtradas_indices)
        )
        # Evita indices obsoletos apos excluir/reordenar NRs.
        indices_visiveis = [
            idx for idx in indices_visiveis
            if isinstance(idx, int) and 0 <= idx < len(self.nr_certificados)
        ]
        if self.nr_filtradas_indices is not None:
            self.nr_filtradas_indices = list(indices_visiveis)
        row = 1
        for idx in indices_visiveis:
            certificado = self.nr_certificados[idx]
            nome_certificado = str(certificado.get("nome", "")).strip()
            if not nome_certificado:
                nome_certificado = f"NR {idx + 1}"
            imprimir_var = tk.BooleanVar(value=bool(certificado.get("imprimir", False)))
            ttk.Checkbutton(
                self.nr_campos_frame,
                variable=imprimir_var,
                command=lambda i=idx, c="imprimir": self._atualizar_campo_nr(i, c),
                takefocus=False,
            ).grid(row=row, column=0, sticky="w", padx=(2, 4), pady=3)
            ttk.Label(self.nr_campos_frame, text=nome_certificado).grid(
                row=row, column=1, sticky="w", padx=(0, 8)
            )

            sem_coluna_de = self._nr_sem_coluna_de(nome_certificado) or self._duracao_nr_item(certificado) <= 1
            entrada_coluna_1 = None
            if sem_coluna_de:
                ttk.Label(self.nr_campos_frame, text="-").grid(
                    row=row, column=2, sticky="w", padx=(0, 8)
                )
                self.nr_certificados[idx]["coluna_1"] = ""
            else:
                entrada_coluna_1 = ttk.Entry(self.nr_campos_frame, width=14)
                entrada_coluna_1.grid(row=row, column=2, sticky="w", padx=(0, 8))
                entrada_coluna_1.insert(0, str(certificado.get("coluna_1", "") or ""))
                self._bind_mask(entrada_coluna_1, mascara_data)
                entrada_coluna_1.bind(
                    "<FocusOut>",
                    lambda _e, i=idx, c="coluna_1": self._atualizar_campo_nr(i, c),
                    add="+",
                )
                entrada_coluna_1.bind(
                    "<KeyRelease>",
                    lambda _e, i=idx, c="coluna_1": self._atualizar_campo_nr(i, c),
                    add="+",
                )
                entrada_coluna_1.bind(
                    "<<DateSelected>>",
                    lambda _e, i=idx, c="coluna_1": self._atualizar_campo_nr(i, c),
                    add="+",
                )
                self._attach_calendar(
                    self.nr_campos_frame,
                    entrada_coluna_1,
                    row=row,
                    column=3,
                    padx=(0, 10),
                    button_text="✓",
                    width=4,
                )

            entrada_coluna_2 = ttk.Entry(self.nr_campos_frame, width=14)
            entrada_coluna_2.grid(row=row, column=4, sticky="w", padx=(0, 8))
            entrada_coluna_2.insert(0, str(certificado.get("coluna_2", "") or ""))
            self._bind_mask(entrada_coluna_2, mascara_data)
            entrada_coluna_2.bind(
                "<FocusOut>",
                lambda _e, i=idx, c="coluna_2": self._atualizar_campo_nr(i, c),
                add="+",
            )
            entrada_coluna_2.bind(
                "<KeyRelease>",
                lambda _e, i=idx, c="coluna_2": self._atualizar_campo_nr(i, c),
                add="+",
            )
            entrada_coluna_2.bind(
                "<<DateSelected>>",
                lambda _e, i=idx, c="coluna_2": self._atualizar_campo_nr(i, c),
                add="+",
            )
            self._attach_calendar(
                self.nr_campos_frame,
                entrada_coluna_2,
                row=row,
                column=5,
                button_text="✓",
                width=4,
            )
            try:
                dias_valor = int(certificado.get("dias", self._duracao_nr_dias(nome_certificado)) or 1)
            except (TypeError, ValueError):
                dias_valor = self._duracao_nr_dias(nome_certificado)
            entrada_dias = ttk.Entry(self.nr_campos_frame, width=5)
            entrada_dias.grid(row=row, column=6, sticky="w", padx=(0, 8))
            entrada_dias.insert(0, str(max(1, dias_valor)))
            entrada_dias.bind(
                "<FocusOut>",
                lambda _e, i=idx, c="dias": self._atualizar_campo_nr(i, c),
                add="+",
            )
            entrada_dias.bind(
                "<Return>",
                lambda _e, i=idx, c="dias": self._atualizar_campo_nr(i, c),
                add="+",
            )
            reciclagem_var = tk.BooleanVar(value=bool(certificado.get("reciclagem", False)))
            ttk.Checkbutton(
                self.nr_campos_frame,
                variable=reciclagem_var,
                command=lambda i=idx, c="reciclagem": self._atualizar_campo_nr(i, c),
                takefocus=False,
            ).grid(row=row, column=7, sticky="w", padx=(14, 0))

            self.nr_certificados_widgets[idx] = {
                "nome": nome_certificado,
                "coluna_1": entrada_coluna_1,
                "coluna_2": entrada_coluna_2,
                "dias": entrada_dias,
                "reciclagem_var": reciclagem_var,
                "imprimir_var": imprimir_var,
            }
            row += 1
        self.nr_canvas.update_idletasks()
        self.nr_canvas.configure(scrollregion=self.nr_canvas.bbox("all"))
        self.focus_set()

    def _on_nr_frame_configure(self, _event=None):
        self.nr_canvas.configure(scrollregion=self.nr_canvas.bbox("all"))

    def _on_nr_canvas_configure(self, event):
        self.nr_canvas.itemconfigure(self.nr_canvas_window, width=event.width)

    def _ativar_scroll_mousewheel_nr(self, _event=None):
        self.bind_all("<MouseWheel>", self._on_nr_mousewheel)

    def _desativar_scroll_mousewheel_nr(self, _event=None):
        self.unbind_all("<MouseWheel>")

    def _on_nr_mousewheel(self, event):
        if not hasattr(self, "nr_canvas"):
            return
        self.nr_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _adicionar_linha_nr(self):
        self.nr_certificados.append(
            {
                "nome": f"NR {len(self.nr_certificados) + 1}",
                "coluna_1": "",
                "coluna_2": "",
                "dias": 1,
                "reciclagem": False,
                "imprimir": False,
                "imprimir_adicionado": False,
            }
        )
        self.nr_certificados.sort(key=self._chave_ordenacao_nr)
        self._render_campos_nr()
        self._salvar_dados()

    def _adicionar_nr_imprimir(self):
        total_marcadas = 0
        for idx, item in enumerate(self.nr_certificados_widgets):
            if not item:
                continue
            marcado = bool(item["imprimir_var"].get())
            self.nr_certificados[idx]["imprimir"] = marcado
            self.nr_certificados[idx]["imprimir_adicionado"] = marcado
            if marcado:
                total_marcadas += 1
        if total_marcadas == 0:
            messagebox.showwarning("CERTIFICADOS", "Selecione ao menos uma NR antes de adicionar.")
            return
        self._atualizar_lista_nr_imprimir()
        self._salvar_dados()

    def _selecionar_tudo_nr(self):
        for idx, item in enumerate(self.nr_certificados_widgets):
            if not item:
                continue
            if "imprimir_var" in item:
                item["imprimir_var"].set(True)
            if idx < len(self.nr_certificados):
                self.nr_certificados[idx]["imprimir"] = True
        self._carregar_outros_documentos_empresa_selecionada()
        self._salvar_dados()

    def _desmarcar_tudo_nr(self, salvar=True, atualizar_ui=True, limpar_datas=False):
        for idx, item in enumerate(self.nr_certificados):
            if idx < len(self.nr_certificados_widgets):
                widget_item = self.nr_certificados_widgets[idx]
                if widget_item and "imprimir_var" in widget_item:
                    widget_item["imprimir_var"].set(False)
                if limpar_datas:
                    if widget_item and widget_item.get("coluna_1") is not None:
                        widget_item["coluna_1"].delete(0, tk.END)
                    if widget_item and widget_item.get("coluna_2") is not None:
                        widget_item["coluna_2"].delete(0, tk.END)
            if limpar_datas:
                item["coluna_1"] = ""
                item["coluna_2"] = ""
            item["imprimir"] = False
            item["imprimir_adicionado"] = False
        if atualizar_ui:
            self._atualizar_lista_nr_imprimir()
        self._carregar_outros_documentos_empresa_selecionada()
        if salvar:
            self._salvar_dados()

    def _remover_linha_nr(self, indice):
        if indice < 0 or indice >= len(self.nr_certificados):
            return
        del self.nr_certificados[indice]
        if isinstance(self.nr_filtradas_indices, list):
            nova_lista = []
            for idx in self.nr_filtradas_indices:
                if not isinstance(idx, int):
                    continue
                if idx == indice:
                    continue
                if idx > indice:
                    nova_lista.append(idx - 1)
                else:
                    nova_lista.append(idx)
            self.nr_filtradas_indices = nova_lista
        self.nr_certificados.sort(key=self._chave_ordenacao_nr)
        if isinstance(self.nr_filtradas_indices, list):
            # Recalcula o filtro apos ordenacao para manter alinhamento com a lista final.
            indices_validos = []
            for idx in self.nr_filtradas_indices:
                if isinstance(idx, int) and 0 <= idx < len(self.nr_certificados):
                    indices_validos.append(idx)
            self.nr_filtradas_indices = indices_validos
        self._render_campos_nr()
        self._salvar_dados()

    def _atualizar_campo_nr(self, indice, coluna):
        if coluna not in {"coluna_1", "coluna_2", "dias", "reciclagem", "imprimir"}:
            return
        if indice < 0 or indice >= len(self.nr_certificados_widgets):
            return
        widget_item = self.nr_certificados_widgets[indice]
        if not widget_item:
            return
        if indice >= len(self.nr_certificados):
            return
        if coluna == "reciclagem":
            valor = bool(widget_item["reciclagem_var"].get())
        elif coluna == "imprimir":
            valor = bool(widget_item["imprimir_var"].get())
        elif coluna == "dias":
            raw = widget_item["dias"].get().strip() if widget_item.get("dias") is not None else ""
            raw = re.sub(r"\D", "", raw)
            valor = max(1, int(raw)) if raw else 1
            if widget_item.get("dias") is not None:
                widget_item["dias"].delete(0, tk.END)
                widget_item["dias"].insert(0, str(valor))
        elif coluna == "coluna_1" and widget_item["coluna_1"] is None:
            valor = ""
        elif coluna == "coluna_2" and widget_item["coluna_2"] is None:
            valor = ""
        else:
            valor = widget_item[coluna].get().strip()
        self.nr_certificados[indice][coluna] = valor
        if coluna == "coluna_1":
            data_inicio = self._parse_data_br(valor)
            if data_inicio is not None:
                data_inicio = self._proximo_ou_mesmo_dia_util(data_inicio)
                valor_inicio = self._formatar_data_br(data_inicio)
                if widget_item.get("coluna_1") is not None and widget_item["coluna_1"].get().strip() != valor_inicio:
                    widget_item["coluna_1"].delete(0, tk.END)
                    widget_item["coluna_1"].insert(0, valor_inicio)
                self.nr_certificados[indice]["coluna_1"] = valor_inicio
                duracao = self._duracao_nr_item(self.nr_certificados[indice])
                data_fim = self._adicionar_dias_uteis(data_inicio, duracao - 1)
                valor_fim = self._formatar_data_br(data_fim)
                if widget_item.get("coluna_2") is not None:
                    widget_item["coluna_2"].delete(0, tk.END)
                    widget_item["coluna_2"].insert(0, valor_fim)
                self.nr_certificados[indice]["coluna_2"] = valor_fim
            elif not valor:
                if widget_item.get("coluna_2") is not None:
                    widget_item["coluna_2"].delete(0, tk.END)
                self.nr_certificados[indice]["coluna_2"] = ""
        if coluna == "imprimir" and not bool(valor):
            self._limpar_linha_datas_nr(indice)
        if coluna in {"coluna_1", "coluna_2", "dias"}:
            # Mesmo sem data valida, atualiza a lista de OUTROS DOCUMENTOS
            # para refletir imediatamente alteracoes de vinculo por NR.
            self._carregar_outros_documentos_empresa_selecionada()
        if coluna == "imprimir":
            self._carregar_outros_documentos_empresa_selecionada()
        self._salvar_dados()

    def _atualizar_select_empresas(self, preservar_id=None, limpar_nr=True):
        self.main_empresa_ids = [None] + [empresa["id"] for empresa in self.empresas]
        self.select_empresa["values"] = [""] + [
            self._empresa_label(empresa) for empresa in self.empresas
        ]

        if preservar_id in self.main_empresa_ids:
            idx = self.main_empresa_ids.index(preservar_id)
        else:
            idx = 0

        self.select_empresa.current(idx)
        self._atualizar_select_funcionarios(self.main_empresa_ids[idx])
        self._aplicar_filtro_nr_por_empresa(self.main_empresa_ids[idx], limpar_nr=limpar_nr)
        self._atualizar_imprimir_empresas(preservar_id=preservar_id)
        self._atualizar_outros_documentos_empresas(preservar_id=preservar_id)

    def _atualizar_select_funcionarios(self, empresa_id):
        if empresa_id is None:
            self.main_funcionario_ids = [None]
            self.select_funcionario["values"] = [""]
            self.select_funcionario.current(0)
            self._atualizar_preview_nr()
            return

        funcs = [f for f in self.funcionarios if self._funcionario_pertence_empresa(f, empresa_id)]
        self.main_funcionario_ids = [None] + [f["id"] for f in funcs]
        self.select_funcionario["values"] = [""] + [f["nome"] for f in funcs]
        if funcs:
            self.select_funcionario.current(1)
        else:
            self.select_funcionario.current(0)
        self._atualizar_preview_nr()

    def _empresa_id_selecionada_imprimir(self):
        idx = self.imprimir_select_empresa.current()
        if idx < 0 or idx >= len(self.imprimir_empresa_ids):
            return None
        return self.imprimir_empresa_ids[idx]

    def _funcionario_id_selecionado_imprimir(self):
        idx = self.imprimir_select_funcionario.current()
        if idx < 0 or idx >= len(self.imprimir_funcionario_ids):
            return None
        return self.imprimir_funcionario_ids[idx]

    def _funcionario_id_selecionado_main(self):
        idx = self.select_funcionario.current()
        if idx < 0 or idx >= len(self.main_funcionario_ids):
            return None
        return self.main_funcionario_ids[idx]

    def _funcionario_ativo_para_documento(self):
        fid = None
        aba_atual = None
        if hasattr(self, "notebook"):
            try:
                aba_atual = self.notebook.nametowidget(self.notebook.select())
            except tk.TclError:
                aba_atual = None

        if aba_atual is self.aba_cadnr:
            fid = self._funcionario_id_selecionado_main()
            if fid is None:
                fid = self._funcionario_id_selecionado_imprimir()
        else:
            fid = self._funcionario_id_selecionado_imprimir()
            if fid is None:
                fid = self._funcionario_id_selecionado_main()
        if fid is None:
            return None
        return next((f for f in self.funcionarios if f["id"] == fid), None)

    @staticmethod
    def _title_case_com_excecoes(texto):
        if not texto:
            return ""
        excecoes = {"de", "da", "das"}
        palavras = re.split(r"(\s+)", texto.strip())
        saida = []
        for i, p in enumerate(palavras):
            if not p or p.isspace():
                saida.append(p)
                continue
            base = p.lower()
            if i != 0 and base in excecoes:
                saida.append(base)
            else:
                saida.append(base[:1].upper() + base[1:])
        return "".join(saida)

    @staticmethod
    def _endereco_empresa_documento(empresa):
        if not isinstance(empresa, dict):
            return ""
        tipo = str(empresa.get("logradouro_tipo", "") or "").strip()
        nome = str(empresa.get("logradouro_nome", "") or "").strip()
        numero = str(empresa.get("numero", "") or "").strip()
        complemento = str(empresa.get("complemento", "") or "").strip()
        bairro = str(empresa.get("bairro", "") or "").strip()
        cidade = str(empresa.get("cidade", "") or "").strip()
        uf = str(empresa.get("uf", "") or "").strip()

        logradouro = " ".join(part for part in (tipo, nome) if part).strip()
        if numero:
            logradouro = f"{logradouro}, {numero}" if logradouro else numero
        if complemento:
            logradouro = f"{logradouro} - {complemento}" if logradouro else complemento
        cidade_uf = " ".join(part for part in (cidade, f"({uf})" if uf else "") if part)
        return ", ".join(part for part in (logradouro, bairro, cidade_uf) if part)

    def _formatar_por_placeholder(self, valor, marcador):
        _ = marcador
        return str(valor or "")

    @staticmethod
    def _escape_xml_text(texto):
        return html.escape(str(texto or ""), quote=False)

    @staticmethod
    def _normalizar_chave_placeholder(texto):
        base = unicodedata.normalize("NFKD", str(texto or ""))
        base = "".join(ch for ch in base if not unicodedata.combining(ch))
        base = re.sub(r"[^A-Za-z0-9]+", "_", base.upper()).strip("_")
        return base

    @staticmethod
    def _chave_nr_base(nome_certificado):
        nome = str(nome_certificado or "").upper().strip()
        nome = re.sub(r"[^A-Z0-9]+", "_", nome)
        nome = re.sub(r"_+", "_", nome).strip("_")
        return nome

    @staticmethod
    def _parse_data_br(texto):
        m = re.fullmatch(r"(\d{2})/(\d{2})/(\d{4})", str(texto or "").strip())
        if not m:
            return None
        dia, mes, ano = map(int, m.groups())
        try:
            return date(ano, mes, dia)
        except ValueError:
            return None

    @staticmethod
    def _formatar_data_br(dt):
        return f"{dt.day:02d}/{dt.month:02d}/{dt.year:04d}"

    @staticmethod
    def _formatar_data_extenso_br(dt):
        meses = [
            "janeiro",
            "fevereiro",
            "março",
            "abril",
            "maio",
            "junho",
            "julho",
            "agosto",
            "setembro",
            "outubro",
            "novembro",
            "dezembro",
        ]
        if dt is None:
            return ""
        return f"{dt.day} de {meses[dt.month - 1]} de {dt.year}"

    @staticmethod
    def _formatar_data_dd_mm_aaaa(dt):
        if dt is None:
            return ""
        return f"{dt.day:02d}/{dt.month:02d}/{dt.year:04d}"

    @staticmethod
    def _formatar_data_mm_aaaa(dt):
        if dt is None:
            return ""
        return f"{dt.month:02d}/{dt.year:04d}"

    @staticmethod
    def _formatar_data_mes_ano_extenso_br(dt):
        meses = [
            "janeiro",
            "fevereiro",
            "março",
            "abril",
            "maio",
            "junho",
            "julho",
            "agosto",
            "setembro",
            "outubro",
            "novembro",
            "dezembro",
        ]
        if dt is None:
            return ""
        return f"{meses[dt.month - 1]} de {dt.year}"

    @staticmethod
    def _formatar_data_dd_mes_extenso_br(dt):
        meses = [
            "janeiro",
            "fevereiro",
            "março",
            "abril",
            "maio",
            "junho",
            "julho",
            "agosto",
            "setembro",
            "outubro",
            "novembro",
            "dezembro",
        ]
        if dt is None:
            return ""
        return f"{dt.day:02d} de {meses[dt.month - 1]}"

    def _data_referencia_documento(self):
        for item in self.nr_certificados:
            if not bool(item.get("imprimir", False)):
                continue
            texto = str(item.get("coluna_2", "") or "").strip()
            dt = self._parse_data_br(texto)
            if dt is not None:
                return dt
        for item in self.nr_certificados:
            if not bool(item.get("imprimir", False)):
                continue
            texto = str(item.get("coluna_1", "") or "").strip()
            dt = self._parse_data_br(texto)
            if dt is not None:
                return dt
        return None

    def _periodo_referencia_documento(self):
        for chave_marcacao in ("imprimir_adicionado", "imprimir"):
            for item in self.nr_certificados:
                if not bool(item.get(chave_marcacao, False)):
                    continue
                data_de = self._parse_data_br(str(item.get("coluna_1", "") or "").strip())
                data_ate = self._parse_data_br(str(item.get("coluna_2", "") or "").strip())
                if data_de is not None or data_ate is not None:
                    return data_de, data_ate
        return None, None

    def _data_nr_selecionada_texto(self):
        for chave_marcacao in ("imprimir_adicionado", "imprimir"):
            for item in self.nr_certificados:
                if not bool(item.get(chave_marcacao, False)):
                    continue
                texto_ate = str(item.get("coluna_2", "") or "").strip()
                if texto_ate:
                    return texto_ate
                texto_de = str(item.get("coluna_1", "") or "").strip()
                if texto_de:
                    return texto_de
        return ""

    def _data_nr_35_texto(self):
        for somente_marcadas in (True, False):
            for item in self.nr_certificados:
                if not isinstance(item, dict):
                    continue
                nome_norm = self._normalizar_nome_nr(item.get("nome", ""))
                if not str(nome_norm).startswith("nr35"):
                    continue
                if somente_marcadas and not bool(item.get("imprimir_adicionado", False) or item.get("imprimir", False)):
                    continue
                for chave_data in ("coluna_2", "coluna_1"):
                    texto = str(item.get(chave_data, "") or "").strip()
                    if not texto:
                        continue
                    dt = self._parse_data_br(texto)
                    if dt is not None:
                        return self._formatar_data_br(dt)
                    if re.fullmatch(r"\d{2}/\d{2}/\d{4}", texto):
                        return texto
        return ""

    @staticmethod
    def _formatar_data_extenso_br_sem_ano(dt):
        meses = [
            "janeiro",
            "fevereiro",
            "março",
            "abril",
            "maio",
            "junho",
            "julho",
            "agosto",
            "setembro",
            "outubro",
            "novembro",
            "dezembro",
        ]
        if dt is None:
            return ""
        return f"{dt.day} de {meses[dt.month - 1]}"

    @staticmethod
    def _pascoa(ano):
        # Algoritmo de Meeus/Jones/Butcher
        a = ano % 19
        b = ano // 100
        c = ano % 100
        d = b // 4
        e = b % 4
        f = (b + 8) // 25
        g = (b - f + 1) // 3
        h = (19 * a + b - d - g + 15) % 30
        i = c // 4
        k = c % 4
        l = (32 + 2 * e + 2 * i - h - k) % 7
        m = (a + 11 * h + 22 * l) // 451
        mes = (h + l - 7 * m + 114) // 31
        dia = ((h + l - 7 * m + 114) % 31) + 1
        return date(ano, mes, dia)

    @classmethod
    def _feriados_nacionais_br(cls, ano):
        pascoa = cls._pascoa(ano)
        fixos = {
            date(ano, 1, 1),   # Confraternizacao Universal
            date(ano, 4, 21),  # Tiradentes
            date(ano, 5, 1),   # Dia do Trabalho
            date(ano, 9, 7),   # Independencia
            date(ano, 10, 12), # Nossa Senhora Aparecida
            date(ano, 11, 2),  # Finados
            date(ano, 11, 15), # Proclamacao da Republica
            date(ano, 11, 20), # Dia Nacional de Zumbi e da Consciencia Negra
            date(ano, 12, 25), # Natal
        }
        moveis = {
            pascoa - timedelta(days=48),  # Carnaval (segunda)
            pascoa - timedelta(days=47),  # Carnaval (terca)
            pascoa - timedelta(days=2),   # Sexta-feira Santa
            pascoa,                       # Pascoa
            pascoa + timedelta(days=60),  # Corpus Christi
        }
        return fixos | moveis

    @classmethod
    def _eh_dia_util(cls, dt):
        # Permite sabado nas datas das NR; bloqueia apenas domingo.
        if dt.weekday() >= 6:
            return False
        return dt not in cls._feriados_nacionais_br(dt.year)

    @classmethod
    def _proximo_dia_util(cls, dt):
        atual = dt + timedelta(days=1)
        while not cls._eh_dia_util(atual):
            atual += timedelta(days=1)
        return atual

    @staticmethod
    def _normalizar_nome_nr(nome):
        return re.sub(r"[^a-z0-9]+", "", str(nome or "").lower())

    @staticmethod
    def _chaves_vinculo_empresa_nr(empresa):
        if not isinstance(empresa, dict):
            return set()
        chaves = set()
        for valor in (empresa.get("nome_pasta", ""), empresa.get("nome", "")):
            texto = str(valor or "").strip().lower()
            if not texto:
                continue
            completo = re.sub(r"[^a-z0-9]+", "", texto)
            if completo:
                chaves.add(completo)
            for parte in re.split(r"[^a-z0-9]+", texto):
                parte = parte.strip()
                if len(parte) >= 2:
                    chaves.add(parte)
        return chaves

    def _arquivo_nr_vinculado_empresa(self, caminho_arquivo, empresa, pasta_referencia):
        if not isinstance(caminho_arquivo, Path) or not isinstance(empresa, dict):
            return True
        # Em pasta dedicada da empresa, todos os modelos pertencem a ela.
        # Filtragem por nome da empresa so e necessaria nas pastas compartilhadas legadas.
        if str(pasta_referencia or "").strip().casefold() not in {"certificados", "documentos"}:
            return True
        chaves = self._chaves_vinculo_empresa_nr(empresa)
        stem_norm = re.sub(r"[^a-z0-9]+", "", str(caminho_arquivo.stem or "").lower())
        if not stem_norm:
            return False
        # Modelos genericos de NR no diretorio compartilhado legado
        # (ex.: "NR 6.docx", "NR 20 (08).docx") devem valer para qualquer empresa.
        chaves_nr = {
            self._normalizar_nome_nr(item.get("nome", ""))
            for item in self._nr_certificados_padrao()
            if isinstance(item, dict) and str(item.get("nome", "")).strip()
        }
        if stem_norm in chaves_nr:
            return True
        if not chaves:
            return True
        return any(chave and chave in stem_norm for chave in chaves)

    @staticmethod
    def _nome_nr_canonico(nome):
        # Remove apenas sufixo de duplicacao automatica "(1)", "(2)" etc.
        # Mantem variantes validas como "NR 20 (08)" e "NR 20 (16)".
        return re.sub(r"\s*\((\d+)\)\s*$", lambda m: "" if int(m.group(1)) < 8 else m.group(0), str(nome or "").strip(), flags=re.IGNORECASE)

    def _nr_sem_coluna_de(self, nome_certificado):
        sem_de = {
            "nr157",
            "nr6",
            "nr12",
            "nr12garra",
            "nr18",
            "nr18pta",
            "nr2008",
            "nr34",
            "nr35",
        }
        nome_nr = self._normalizar_nome_nr(nome_certificado)
        if nome_nr in sem_de:
            return True
        # Aceita sufixos de empresa no nome da NR (ex.: "nr2008projetta").
        bases_com_sufixo = {"nr6", "nr12", "nr18pta", "nr34", "nr35", "nr2008"}
        for base in bases_com_sufixo:
            if re.fullmatch(rf"{base}[a-z]+", nome_nr):
                return True
        empresa_id = self._empresa_id_selecionada_main()
        if empresa_id is None:
            return False
        empresa = next((e for e in self.empresas if e.get("id") == empresa_id), None)
        if not empresa:
            return False
        nome_empresa = str(empresa.get("nome_pasta", "") or empresa.get("nome", "") or "").strip().upper()
        return nome_empresa == "RJM" and "brigadista" in nome_nr

    def _primeiro_indice_coluna_de_editavel(self):
        for idx, item in enumerate(self.nr_certificados_widgets):
            if not item:
                continue
            if item.get("coluna_1") is not None:
                return idx
        return None

    def _nr_selecionada_para_data(self, indice):
        if indice < 0 or indice >= len(self.nr_certificados):
            return False
        return bool(self.nr_certificados[indice].get("imprimir", False))

    def _limpar_linha_datas_nr(self, indice):
        if indice < 0 or indice >= len(self.nr_certificados_widgets):
            return
        item_widget = self.nr_certificados_widgets[indice]
        if not item_widget:
            self.nr_certificados[indice]["coluna_1"] = ""
            self.nr_certificados[indice]["coluna_2"] = ""
            return
        entry_de = item_widget["coluna_1"]
        entry_ate = item_widget["coluna_2"]
        if entry_de is not None:
            entry_de.delete(0, tk.END)
        if entry_ate is not None:
            entry_ate.delete(0, tk.END)
        self.nr_certificados[indice]["coluna_1"] = ""
        self.nr_certificados[indice]["coluna_2"] = ""

    def _duracao_nr_dias(self, nome_certificado):
        mapa = {
            "nr5": 3,
            "nr11garra": 2,
            "nr11munck": 2,
            "nr11pr": 2,
            "nr11emp": 2,
            "nr1140": 5,
            "nr18guin": 15,
            "nr18and": 2,
            "nr2016": 2,
            "nr33": 2,
            "certmunck": 5,
        }
        chave = self._normalizar_nome_nr(nome_certificado)
        if re.fullmatch(r"nr11emp[a-z]*", chave):
            return 2
        return max(1, mapa.get(chave, 1))

    def _duracao_nr_item(self, item):
        if isinstance(item, dict):
            raw = item.get("dias", None)
            try:
                dias = int(raw)
                if dias > 0:
                    return dias
            except (TypeError, ValueError):
                pass
            return self._duracao_nr_dias(item.get("nome", ""))
        return 1

    def _adicionar_dias_uteis(self, dt, qtd_dias):
        atual = dt
        for _ in range(max(0, int(qtd_dias))):
            atual = self._proximo_dia_util(atual)
        return atual

    @classmethod
    def _proximo_ou_mesmo_dia_util(cls, dt):
        atual = dt
        while not cls._eh_dia_util(atual):
            atual += timedelta(days=1)
        return atual

    def _preencher_sequencia_dias_uteis_nr(self, coluna):
        if coluna not in {"coluna_1", "coluna_2"}:
            return
        if not self.nr_certificados_widgets:
            return
        if coluna == "coluna_1":
            indice_inicial = None
            texto_inicial = ""
            for idx, widget_item in enumerate(self.nr_certificados_widgets):
                if not widget_item:
                    continue
                if not self._nr_selecionada_para_data(idx):
                    continue
                entry = widget_item.get("coluna_1")
                if entry is None:
                    continue
                texto = entry.get().strip()
                if self._parse_data_br(texto) is not None:
                    indice_inicial = idx
                    texto_inicial = texto
                    break
            if indice_inicial is None:
                return
        else:
            indice_inicial = None
            texto_inicial = ""
            for idx, widget_item in enumerate(self.nr_certificados_widgets):
                if not widget_item:
                    continue
                if not self._nr_selecionada_para_data(idx):
                    continue
                if widget_item.get(coluna) is None:
                    continue
                texto = widget_item[coluna].get().strip()
                if self._parse_data_br(texto) is not None:
                    indice_inicial = idx
                    texto_inicial = texto
                    break
            if indice_inicial is None:
                return
        data_atual = self._parse_data_br(texto_inicial)
        if data_atual is None:
            return

        if coluna == "coluna_1":
            inicio = data_atual
            for idx in range(0, indice_inicial):
                self._limpar_linha_datas_nr(idx)
            for idx in range(indice_inicial, len(self.nr_certificados)):
                if not self._nr_selecionada_para_data(idx):
                    self._limpar_linha_datas_nr(idx)
                    continue
                if not self.nr_certificados_widgets[idx]:
                    continue
                item = self.nr_certificados[idx]
                duracao = self._duracao_nr_item(item)
                fim = self._adicionar_dias_uteis(inicio, duracao - 1)

                valor_inicio = self._formatar_data_br(inicio)
                valor_fim = self._formatar_data_br(fim)
                entry_inicio = self.nr_certificados_widgets[idx]["coluna_1"]
                entry_fim = self.nr_certificados_widgets[idx]["coluna_2"]
                if duracao <= 1:
                    if entry_inicio is not None:
                        entry_inicio.delete(0, tk.END)
                    self.nr_certificados[idx]["coluna_1"] = ""
                else:
                    if entry_inicio is not None:
                        entry_inicio.delete(0, tk.END)
                        entry_inicio.insert(0, valor_inicio)
                        self.nr_certificados[idx]["coluna_1"] = valor_inicio
                    else:
                        self.nr_certificados[idx]["coluna_1"] = ""
                if entry_fim is not None:
                    entry_fim.delete(0, tk.END)
                    entry_fim.insert(0, valor_fim)
                self.nr_certificados[idx]["coluna_2"] = valor_fim

                inicio = self._proximo_dia_util(fim)
            return

        # Sequencia iniciando por "Até" da primeira linha:
        # usa duracao por NR e dias uteis para calcular De/Ate.
        inicio = data_atual
        for idx, item in enumerate(self.nr_certificados):
            if idx < indice_inicial or not self._nr_selecionada_para_data(idx):
                self._limpar_linha_datas_nr(idx)
                continue
            if not self.nr_certificados_widgets[idx]:
                continue
            duracao = self._duracao_nr_item(item)
            fim = self._adicionar_dias_uteis(inicio, duracao - 1)

            valor_inicio = self._formatar_data_br(inicio)
            valor_fim = self._formatar_data_br(fim)

            entry_ate = self.nr_certificados_widgets[idx]["coluna_2"]
            if entry_ate is not None:
                entry_ate.delete(0, tk.END)
                entry_ate.insert(0, valor_fim)
            self.nr_certificados[idx]["coluna_2"] = valor_fim

            entry_de = self.nr_certificados_widgets[idx]["coluna_1"]
            if duracao <= 1:
                if entry_de is not None:
                    entry_de.delete(0, tk.END)
                self.nr_certificados[idx]["coluna_1"] = ""
            else:
                if entry_de is not None:
                    entry_de.delete(0, tk.END)
                    entry_de.insert(0, valor_inicio)
                    self.nr_certificados[idx]["coluna_1"] = valor_inicio
                else:
                    self.nr_certificados[idx]["coluna_1"] = ""

            inicio = self._proximo_dia_util(fim)

    def _montar_campos_documento(self, funcionario):
        empresa = self._empresa_do_funcionario(funcionario) or {}
        data_de_ref, data_ate_ref = self._periodo_referencia_documento()
        data_ref = self._data_referencia_documento()
        if data_ref is None:
            data_ref = data_ate_ref or data_de_ref
        admissao_raw = str(funcionario.get("admissao", "") or "").strip()
        admissao_dt = self._parse_data_br(admissao_raw)
        admissao_extenso = self._formatar_data_extenso_br(admissao_dt) if admissao_dt else admissao_raw
        admissao_extenso_dd = (
            f"{admissao_dt.day:02d} de "
            f"{self._formatar_data_extenso_br(admissao_dt).split(' de ', 1)[1]}"
            if admissao_dt
            else admissao_raw
        )
        data_ref_extenso = self._formatar_data_extenso_br(data_ref)
        data_de_extenso = self._formatar_data_extenso_br(data_de_ref or data_ref)
        data_ate_extenso = self._formatar_data_extenso_br(data_ate_ref or data_ref)
        data_ate_extenso_dd = (
            f"{(data_ate_ref or data_ref).day:02d} de "
            f"{self._formatar_data_extenso_br(data_ate_ref or data_ref).split(' de ', 1)[1]}"
            if (data_ate_ref or data_ref)
            else ""
        )
        data_de_mes_ano = self._formatar_data_dd_mes_extenso_br(data_de_ref or data_ref)
        if data_de_ref and data_ate_ref:
            de_e_ate_extenso = (
                f"{self._formatar_data_extenso_br_sem_ano(data_de_ref)} "
                f"e {self._formatar_data_extenso_br(data_ate_ref)}"
            )
            de_a_ate_extenso = (
                f"{self._formatar_data_extenso_br_sem_ano(data_de_ref)} "
                f"a {self._formatar_data_extenso_br(data_ate_ref)}"
            )
        else:
            data_base = data_ate_ref or data_de_ref or data_ref
            de_e_ate_extenso = self._formatar_data_extenso_br(data_base)
            de_a_ate_extenso = self._formatar_data_extenso_br(data_base)
        tem_reciclagem_selecionada = any(
            bool(item.get("reciclagem", False))
            and bool(item.get("imprimir_adicionado", False) or item.get("imprimir", False))
            for item in self.nr_certificados
        )
        campos = {
            # Novo padrao de placeholders globais.
            "FUNCIONARIO1": str(funcionario.get("nome", "") or "").upper(),
            "CPF1": str(funcionario.get("cpf", "") or "").strip(),
            "FOTO1": self._caminho_foto_funcionario(funcionario),
            "ADMISSAO1": self._formatar_data_dd_mm_aaaa(admissao_dt) if admissao_dt else admissao_raw,
            "FUNCIONARIO2": self._title_case_com_excecoes(funcionario.get("nome", "")),
            "CPF2": str(funcionario.get("cpf", "") or "").strip(),
            "PROFISSAO1": str(funcionario.get("funcao", "") or "").upper(),
            "FUNCIONARIO3": str(funcionario.get("nome", "") or "").upper(),
            "CPF3": str(funcionario.get("cpf", "") or "").strip(),
            "ADMISSAO2": admissao_extenso_dd,
            "FUNCIONARIO4": self._title_case_com_excecoes(funcionario.get("nome", "")),
            "CPF4": str(funcionario.get("cpf", "") or "").strip(),
            "ADMISSAO3": self._formatar_data_mm_aaaa(admissao_dt) if admissao_dt else "",
            "ADMISSAO4": self._formatar_data_dd_mm_aaaa(admissao_dt) if admissao_dt else admissao_raw,
            "COMECO1": data_de_mes_ano,
            "FIM1": data_ate_extenso_dd,
            "FIM2": data_ate_extenso_dd,
            "EMPRESA1": str(empresa.get("nome", "") or "").upper(),
            "CNPJ1": str(empresa.get("cnpj", "") or "").strip(),
            "ENDERECO1": self._endereco_empresa_documento(empresa),
            "CIDADE1": " ".join(
                part
                for part in (
                    str(empresa.get("cidade", "") or "").strip(),
                    f"({str(empresa.get('uf', '') or '').strip()})"
                    if str(empresa.get("uf", "") or "").strip()
                    else "",
                )
                if part
            ),
            "LOGO1": self._caminho_logo_empresa(empresa),
            "LOGO3": self._caminho_logo_empresa(empresa),
            "EMPRESA2": "Projetta Engenharia Consultoria e Treinamentos LTDA",
            "CNPJ2": "34.459.757/0001-60",
            "LOGO2": "_logos_empresas/1_Projetta.png",
            "EMPRESA3": str(empresa.get("nome", "") or "").strip(),
            "CNPJ3": str(empresa.get("cnpj", "") or "").strip(),
            "EMPRESA4": str(empresa.get("nome", "") or "").upper(),
            "CNPJ4": str(empresa.get("cnpj", "") or "").strip(),
            "EMPRESA5": str(empresa.get("nome", "") or "").strip(),
            "CNPJ5": str(empresa.get("cnpj", "") or "").strip(),
            "NOME": funcionario.get("nome", ""),
            "NOME_DO_FUNCIONARIO_1": str(funcionario.get("nome", "") or "").upper(),
            "NOME_DO_FUNCIONARIO_2": self._title_case_com_excecoes(funcionario.get("nome", "")),
            "CPF": funcionario.get("cpf", ""),
            "RG1": funcionario.get("rg", ""),
            "RG": funcionario.get("rg", ""),
            "FUNCAO": funcionario.get("funcao", ""),
            "CBO": funcionario.get("cbo", ""),
            "SALARIO": funcionario.get("salario", ""),
            "ADMISSAO": admissao_extenso,
            "ADMISSAO_NUMERICA": admissao_raw,
            "EMPRESA": empresa.get("nome", ""),
            "CIDADE3": " ".join(
                part
                for part in (
                    str(empresa.get("cidade", "") or "").strip(),
                    f"({str(empresa.get('uf', '') or '').strip()})"
                    if str(empresa.get("uf", "") or "").strip()
                    else "",
                )
                if part
            ),
            "ENDERECO3": self._endereco_empresa_documento(empresa),
            "CNPJ": empresa.get("cnpj", ""),
            "LOGO": self._caminho_logo_empresa(empresa),
            "IMG1": self._caminho_logo_empresa(empresa),
            "IMG3": self._caminho_logo_empresa(empresa),
            "IMG2": "_logos_empresas/1_Projetta.png",
            "DATA": data_ref_extenso,
            "DATA_EXTENSO": data_ref_extenso,
            "DATA_DE_EXTENSO": data_de_extenso,
            "DATA_ATE_EXTENSO": data_ate_extenso,
            "DATANR": self._data_nr_selecionada_texto(),
            "DATA_NR": self._data_nr_selecionada_texto(),
            "DATA35": self._data_nr_35_texto(),
            "DE_E_ATE_EXTENSO": de_e_ate_extenso,
            "DE_A_ATE_EXTENSO": de_a_ate_extenso,
            "TEXTO_CERTIFICA": "A Projetta Engenharia Consultoria e Treinamentos LTDA, CERTIFICA que:",
            "TEXTO_CAPACITACAO": (
                "Obteve aproveitamento satisfatório no curso de capacitação de RECICLAGEM"
                if tem_reciclagem_selecionada
                else "Obteve aproveitamento satisfatório no curso de capacitação"
            ),
        }
        for idx, item in enumerate(self.nr_certificados, start=1):
            base = self._chave_nr_base(item.get("nome", "")) or f"NR_{idx}"
            coluna_1 = str(item.get("coluna_1", "") or "")
            coluna_2 = str(item.get("coluna_2", "") or "")
            reciclagem = bool(item.get("reciclagem", False))
            reciclagem_txt = "SIM" if reciclagem else "NAO"

            # Chaves por nome do certificado (ex.: NR_6_COLUNA_1)
            campos[f"{base}_COLUNA_1"] = coluna_1
            campos[f"{base}_COLUNA_2"] = coluna_2
            campos[f"{base}_COL1"] = coluna_1
            campos[f"{base}_COL2"] = coluna_2
            campos[f"{base}_RECICLAGEM"] = reciclagem_txt
            campos[f"{base}_REC"] = reciclagem_txt

            # Chaves por indice para facilitar modelos genericos.
            campos[f"NR_{idx}_COLUNA_1"] = coluna_1
            campos[f"NR_{idx}_COLUNA_2"] = coluna_2
            campos[f"NR_{idx}_COL1"] = coluna_1
            campos[f"NR_{idx}_COL2"] = coluna_2
            campos[f"NR_{idx}_RECICLAGEM"] = reciclagem_txt
            campos[f"NR_{idx}_REC"] = reciclagem_txt
        return campos

    def _atualizar_preview_nr(self):
        if not hasattr(self, "nr_preview"):
            return
        funcionario = self._funcionario_ativo_para_documento()
        if funcionario is None:
            preview = (
                "Nenhum funcionario selecionado.\n"
                "Selecione em CADNR ou CADASTROS para visualizar os dados do Word."
            )
        else:
            campos = self._montar_campos_documento(funcionario)
            preview = (
                f"NOME: {campos['NOME']}\n"
                f"CPF: {campos['CPF']}\n"
                f"RG: {campos['RG']}\n"
                f"FUNCAO: {campos['FUNCAO']}\n"
                f"CBO: {campos['CBO']}\n"
                f"SALARIO: {campos['SALARIO']}\n"
                f"EMPRESA: {campos['EMPRESA']}"
            )

        self.nr_preview.configure(state="normal")
        self.nr_preview.delete("1.0", tk.END)
        self.nr_preview.insert("1.0", preview)
        self.nr_preview.configure(state="disabled")

    @staticmethod
    def _obter_nome_arquivo_seguro(nome, default="funcionario"):
        nome_seguro = re.sub(r"[<>:\"/\\\\|?*]+", "", str(nome or default)).strip().strip(".")
        if not nome_seguro:
            nome_seguro = default
        return nome_seguro

    @staticmethod
    def _obter_pasta_saida(nome_limpo):
        # Salva fora do diretorio do repositorio Git local.
        desktop_base = App._obter_desktop_base()
        pasta_destino = desktop_base / nome_limpo
        pasta_destino.mkdir(parents=True, exist_ok=True)
        return pasta_destino

    @staticmethod
    def _obter_desktop_base():
        userprofile = Path.home()
        onedrive = os.environ.get("OneDrive", "")
        candidatos = [
            userprofile / "OneDrive" / "Área de Trabalho",
            userprofile / "OneDrive" / "Area de Trabalho",
            Path(onedrive) / "Área de Trabalho" if onedrive else None,
            Path(onedrive) / "Area de Trabalho" if onedrive else None,
            Path(onedrive) / "Desktop" if onedrive else None,
            userprofile / "Desktop",
        ]
        candidatos = [p for p in candidatos if p is not None]
        for pasta in candidatos:
            try:
                if pasta.exists():
                    return pasta
            except OSError:
                continue
        base = candidatos[0] if candidatos else userprofile / "Desktop"
        base.mkdir(parents=True, exist_ok=True)
        return base

    @staticmethod
    def _ip_local_preferencial():
        host_env = str(os.environ.get("CADNR_QR_LOCAL_HOST", "") or "").strip()
        if host_env:
            return host_env
        try:
            with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as s:
                s.connect(("8.8.8.8", 80))
                ip = str(s.getsockname()[0] or "").strip()
                if ip:
                    return ip
        except OSError:
            pass
        return "127.0.0.1"

    @staticmethod
    def _deve_sincronizar_qr_pendentes():
        # Por padrao, nao repete publicacao de documentos antigos no startup,
        # evitando que arquivos removidos manualmente do site "voltem" a aparecer.
        flag = str(os.environ.get("CADNR_QR_SYNC_PENDENTES", "0") or "").strip().lower()
        return flag in {"1", "true", "on", "yes", "sim"}

    @staticmethod
    def _abrir_url_no_chrome(url):
        url_txt = str(url or "").strip()
        if not url_txt:
            return False

        candidatos = []
        program_files = os.environ.get("ProgramFiles", r"C:\Program Files")
        program_files_x86 = os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)")
        local_appdata = os.environ.get("LOCALAPPDATA", "")
        if program_files:
            candidatos.append(Path(program_files) / "Google" / "Chrome" / "Application" / "chrome.exe")
        if program_files_x86:
            candidatos.append(Path(program_files_x86) / "Google" / "Chrome" / "Application" / "chrome.exe")
        if local_appdata:
            candidatos.append(Path(local_appdata) / "Google" / "Chrome" / "Application" / "chrome.exe")

        for chrome_exe in candidatos:
            try:
                if chrome_exe.exists():
                    subprocess.Popen(
                        [str(chrome_exe), url_txt],
                        stdout=subprocess.DEVNULL,
                        stderr=subprocess.DEVNULL,
                    )
                    return True
            except OSError:
                continue

        try:
            browser = webbrowser.get("chrome")
            browser.open(url_txt, new=2)
            return True
        except webbrowser.Error:
            pass

        return webbrowser.open(url_txt, new=2)

    def _iniciar_servidor_qr_local(self):
        if self._qr_http_server is not None and self._qr_http_base_url:
            return True

        app = self

        class QrLocalHandler(BaseHTTPRequestHandler):
            def do_GET(self):
                try:
                    path_req = parse.urlsplit(str(self.path or "")).path
                    if path_req == "/__cadnr_ping__":
                        self.send_response(200)
                        self.end_headers()
                        return
                    arquivo = None

                    # Formato novo e estavel: caminho relativo ao projeto.
                    m_arquivo = re.match(r"^/qr-file/(.+)$", path_req)
                    if m_arquivo:
                        caminho_rel = parse.unquote(m_arquivo.group(1))
                        arquivo = app._resolver_caminho_qr_relativo(caminho_rel)
                    else:
                        # Compatibilidade com formato legado por token.
                        m = re.match(r"^/qr/([A-Za-z0-9_-]{12,})$", path_req)
                        if not m:
                            self.send_error(404)
                            return
                        token = m.group(1)
                        caminho_pdf = app._qr_token_map.get(token)
                        if caminho_pdf:
                            arquivo = Path(str(caminho_pdf or "")).expanduser()
                            if not arquivo.is_absolute():
                                arquivo = (Path(__file__).resolve().parent / arquivo).resolve()

                    if arquivo is None:
                        self.send_error(404)
                        return
                    if not arquivo.exists() or not arquivo.is_file():
                        self.send_error(404)
                        return
                    nome_arq = arquivo.name.encode("utf-8", errors="ignore").decode("utf-8", errors="ignore")
                    dados = arquivo.read_bytes()
                    self.send_response(200)
                    self.send_header("Content-Type", "application/pdf")
                    self.send_header("Content-Length", str(len(dados)))
                    self.send_header("Content-Disposition", f'inline; filename="{nome_arq}"')
                    self.send_header("Cache-Control", "no-cache")
                    self.end_headers()
                    self.wfile.write(dados)
                except Exception:
                    self.send_error(500)

            def log_message(self, fmt, *args):
                return

        try:
            ip = self._ip_local_preferencial()
            porta_env = str(os.environ.get("CADNR_QR_LOCAL_PORT", "") or "").strip()
            try:
                porta = int(porta_env) if porta_env else 8765
            except ValueError:
                porta = 8765
            try:
                servidor = ThreadingHTTPServer(("0.0.0.0", porta), QrLocalHandler)
            except OSError:
                servidor = ThreadingHTTPServer(("0.0.0.0", 0), QrLocalHandler)
            thread = threading.Thread(target=servidor.serve_forever, daemon=True)
            thread.start()
            self._qr_http_server = servidor
            self._qr_http_thread = thread
            porta_real = int(servidor.server_address[1])
            self._qr_http_base_url = f"http://{ip}:{porta_real}"
            return True
        except OSError:
            return False

    @staticmethod
    def _resolver_caminho_qr_relativo(caminho_relativo):
        caminho_rel = str(caminho_relativo or "").strip().replace("\\", "/")
        if not caminho_rel:
            return None
        base = Path(__file__).resolve().parent.resolve()
        try:
            candidato = (base / caminho_rel).resolve()
            candidato.relative_to(base)
        except Exception:
            return None
        return candidato

    def _url_local_qr_para_arquivo(self, caminho_arquivo, permitir_inexistente=False):
        caminho = Path(str(caminho_arquivo or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        if (not permitir_inexistente) and (not caminho.exists() or not caminho.is_file()):
            return ""
        if not self._iniciar_servidor_qr_local():
            return ""
        try:
            base = Path(__file__).resolve().parent.resolve()
            rel = caminho.resolve().relative_to(base)
            rel_txt = str(rel).replace("\\", "/")
            return f"{self._qr_http_base_url}/qr-file/{parse.quote(rel_txt, safe='/')}"
        except Exception:
            pass
        if caminho.exists() and caminho.is_file():
            token_base = f"{caminho.resolve()}|{int(caminho.stat().st_mtime)}|{int(caminho.stat().st_size)}"
        else:
            token_base = f"{caminho.resolve()}|pending"
        token = hashlib.sha1(token_base.encode("utf-8", errors="ignore")).hexdigest()[:20]
        self._qr_token_map[token] = str(caminho.resolve())
        return f"{self._qr_http_base_url}/qr/{token}"

    def _encerrar_servidor_qr_local(self):
        servidor = getattr(self, "_qr_http_server", None)
        if servidor is None:
            return
        try:
            servidor.shutdown()
        except Exception:
            pass
        try:
            servidor.server_close()
        except Exception:
            pass
        self._qr_http_server = None
        self._qr_http_thread = None
        self._qr_http_base_url = ""
        self._qr_token_map = {}

    @staticmethod
    def _normalizar_repo_github(repo_txt):
        repo = str(repo_txt or "").strip()
        if not repo:
            return ""
        repo = re.sub(r"^https?://github\.com/", "", repo, flags=re.IGNORECASE)
        repo = re.sub(r"^git@github\.com:", "", repo, flags=re.IGNORECASE)
        repo = repo.replace("\\", "/").strip("/")
        if repo.lower().endswith(".git"):
            repo = repo[:-4]
        partes = [p for p in repo.split("/") if p]
        if len(partes) < 2:
            return ""
        return f"{partes[0]}/{partes[1]}"

    @staticmethod
    def _normalizar_pages_base(base_txt):
        base = str(base_txt or "").strip()
        if not base:
            return ""
        base = re.sub(r"/+$", "", base)
        base = re.sub(r"/index\.html$", "", base, flags=re.IGNORECASE)
        return base

    @staticmethod
    def _obter_repo_branch_github_preferencial():
        repo_padrao = "Elizangela2805/documentos"
        branch_padrao = "main"
        repo_git = ""
        branch_git = ""
        repo_dir = Path(__file__).resolve().parent

        try:
            proc_remote = subprocess.run(
                ["git", "-C", str(repo_dir), "remote", "get-url", "origin"],
                check=False,
                capture_output=True,
                text=True,
                timeout=10,
            )
            if proc_remote.returncode == 0:
                repo_git = App._normalizar_repo_github(proc_remote.stdout)
        except Exception:
            repo_git = ""

        try:
            proc_branch = subprocess.run(
                ["git", "-C", str(repo_dir), "rev-parse", "--abbrev-ref", "HEAD"],
                check=False,
                capture_output=True,
                text=True,
                timeout=10,
            )
            if proc_branch.returncode == 0:
                b = str(proc_branch.stdout or "").strip()
                if b and b.upper() != "HEAD":
                    branch_git = b
        except Exception:
            branch_git = ""

        repo_env = App._normalizar_repo_github(os.environ.get("CADNR_QR_GITHUB_REPO", ""))
        branch_env = str(os.environ.get("CADNR_QR_GITHUB_BRANCH", "") or "").strip()

        # Prioriza configuracao explicita do app/ambiente; Git local fica como fallback.
        repo = repo_env or repo_git or repo_padrao
        branch = branch_env or branch_git or branch_padrao
        return repo, branch

    @staticmethod
    def _obter_config_qr_github():
        token = str(os.environ.get("CADNR_QR_GITHUB_TOKEN", "") or "").strip()
        repo, branch = App._obter_repo_branch_github_preferencial()
        if not token or not repo:
            return None
        pasta = str(os.environ.get("CADNR_QR_GITHUB_DIR", "") or "").strip().strip("/")
        raw_base = str(os.environ.get("CADNR_QR_GITHUB_RAW_BASE", "") or "").strip().rstrip("/")
        if not raw_base:
            raw_base = f"https://raw.githubusercontent.com/{repo}/{branch}"
        pages_base = str(os.environ.get("CADNR_QR_GITHUB_PAGES_BASE", "") or "").strip().rstrip("/")
        if not pages_base:
            pages_base = "https://elizangela2805.github.io/documentos"
        return {
            "token": token,
            "repo": repo,
            "branch": branch,
            "pasta": pasta,
            "raw_base": raw_base,
            "pages_base": pages_base,
        }

    def _aplicar_configuracao_github_ambiente(self):
        repo = str(self.github_repo or "").strip() or "Elizangela2805/documentos"
        branch = str(self.github_branch or "").strip() or "main"
        pasta = str(self.github_dir or "").strip().strip("/")
        pages_base = self._normalizar_pages_base(self.github_pages_base) or "https://elizangela2805.github.io/documentos"
        token = str(self.github_token or "").strip()
        os.environ["CADNR_QR_GITHUB_REPO"] = repo
        os.environ["CADNR_QR_GITHUB_BRANCH"] = branch
        os.environ["CADNR_QR_GITHUB_DIR"] = pasta
        os.environ["CADNR_QR_GITHUB_PAGES_BASE"] = pages_base.rstrip("/")
        if token:
            os.environ["CADNR_QR_GITHUB_TOKEN"] = token
        elif "CADNR_QR_GITHUB_TOKEN" in os.environ:
            del os.environ["CADNR_QR_GITHUB_TOKEN"]

    @staticmethod
    def _montar_caminho_repo_qr_github(caminho_pdf, pasta_repo=""):
        try:
            base = Path(__file__).resolve().parent.resolve()
            rel = Path(caminho_pdf).resolve().relative_to(base)
            rel_txt = str(rel).replace("\\", "/")
        except Exception:
            try:
                p = Path(caminho_pdf).resolve()
                parent = str(p.parent.name or "").strip()
                rel_txt = f"{parent}/{p.name}" if parent else p.name
            except Exception:
                rel_txt = Path(caminho_pdf).name
        rel_txt = rel_txt.lstrip("/").replace("../", "").replace("..\\", "")
        pasta = str(pasta_repo or "").strip().strip("/")
        if pasta:
            # Evita duplicar prefixo quando o caminho local ja esta dentro da pasta configurada.
            rel_norm = rel_txt.lower()
            pasta_norm = pasta.lower().strip("/")
            if rel_norm == pasta_norm or rel_norm.startswith(f"{pasta_norm}/"):
                return rel_txt
            return f"{pasta}/{rel_txt}"
        return rel_txt

    @staticmethod
    def _url_github_blob_para_arquivo(caminho_arquivo):
        repo, branch = App._obter_repo_branch_github_preferencial()
        pasta = str(os.environ.get("CADNR_QR_GITHUB_DIR", "") or "").strip().strip("/")
        if not repo:
            return ""
        try:
            caminho = Path(str(caminho_arquivo or "")).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            repo_path = App._montar_caminho_repo_qr_github(caminho, pasta)
            if not repo_path:
                return ""
            branch_url = parse.quote(branch, safe="")
            arquivo_url = parse.quote(repo_path, safe="/")
            return f"https://github.com/{repo}/blob/{branch_url}/{arquivo_url}"
        except Exception:
            return ""

    @staticmethod
    def _url_github_pages_para_arquivo(caminho_arquivo):
        repo, _branch = App._obter_repo_branch_github_preferencial()
        pasta = str(os.environ.get("CADNR_QR_GITHUB_DIR", "") or "").strip().strip("/")
        pages_base = App._normalizar_pages_base(os.environ.get("CADNR_QR_GITHUB_PAGES_BASE", ""))
        if not pages_base:
            pages_base = "https://elizangela2805.github.io/documentos"
        if not pages_base:
            return ""
        try:
            caminho = Path(str(caminho_arquivo or "")).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            repo_path = App._montar_caminho_repo_qr_github(caminho, pasta)
            if not repo_path:
                return ""
            arquivo_url = parse.quote(repo_path, safe="/")
            return f"{pages_base}/{arquivo_url}"
        except Exception:
            return ""

    @staticmethod
    def _url_github_raw_para_arquivo(caminho_arquivo):
        repo, branch = App._obter_repo_branch_github_preferencial()
        pasta = str(os.environ.get("CADNR_QR_GITHUB_DIR", "") or "").strip().strip("/")
        if not repo:
            return ""
        try:
            caminho = Path(str(caminho_arquivo or "")).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            repo_path = App._montar_caminho_repo_qr_github(caminho, pasta)
            if not repo_path:
                return ""
            arquivo_url = parse.quote(repo_path, safe="/")
            branch_url = parse.quote(str(branch or "main"), safe="")
            return f"https://raw.githubusercontent.com/{repo}/{branch_url}/{arquivo_url}"
        except Exception:
            return ""

    @staticmethod
    def _slug_url_texto(texto):
        base = unicodedata.normalize("NFKD", str(texto or ""))
        base = "".join(ch for ch in base if not unicodedata.combining(ch))
        return re.sub(r"[^a-z0-9]+", "", base.lower())

    @staticmethod
    def _extrair_data_slug_documento(repo_path, caminho_arquivo):
        def _normalizar_data_match(grupo_a, grupo_b, grupo_c):
            a = str(grupo_a or "")
            b = str(grupo_b or "")
            c = str(grupo_c or "")
            if len(a) == 4:
                return f"{c.zfill(2)}{b.zfill(2)}{a}"
            return f"{a.zfill(2)}{b.zfill(2)}{c}"

        texto_busca = " ".join(
            [
                str(repo_path or ""),
                str(Path(str(caminho_arquivo or "")).stem or ""),
            ]
        )
        m_data = re.search(r"(?<!\d)(\d{2})[-_/]?(\d{2})[-_/]?(\d{4})(?!\d)", texto_busca)
        if not m_data:
            m_data = re.search(r"(?<!\d)(\d{4})[-_/]?(\d{2})[-_/]?(\d{2})(?!\d)", texto_busca)
        if m_data:
            return _normalizar_data_match(m_data.group(1), m_data.group(2), m_data.group(3))
        try:
            caminho = Path(str(caminho_arquivo or "")).expanduser()
            if caminho.exists():
                dt = datetime.fromtimestamp(caminho.stat().st_mtime)
                return dt.strftime("%d%m%Y")
        except Exception:
            pass
        return datetime.now().strftime("%d%m%Y")

    @staticmethod
    def _extrair_documento_slug_documento(repo_path, caminho_arquivo):
        stem = str(Path(str(repo_path or "")).stem or "").strip()
        if not stem:
            stem = str(Path(str(caminho_arquivo or "")).stem or "").strip()
        stem_ascii = unicodedata.normalize("NFKD", stem)
        stem_ascii = "".join(ch for ch in stem_ascii if not unicodedata.combining(ch))
        stem_ascii = stem_ascii.lower()
        if "carteirinha" in stem_ascii:
            slug_cart = App._slug_url_texto(stem_ascii)
            return slug_cart or "carteirinha"
        m_nr = re.search(r"\bnr\s*0*([0-9]{1,3})\s*([a-z]{2,6})?\b", stem_ascii)
        if m_nr:
            numero = str(m_nr.group(1) or "").strip()
            sufixo = str(m_nr.group(2) or "").strip()
            if sufixo not in {"pta"}:
                sufixo = ""
            return f"nr{numero}{sufixo}"
        slug = App._slug_url_texto(stem)
        return slug or "documento"

    @staticmethod
    def _extrair_funcionario_slug_documento(repo_path, caminho_arquivo):
        repo_norm = str(repo_path or "").replace("\\", "/").strip("/")
        partes = [p for p in repo_norm.split("/") if p]
        if partes and str(partes[0]).casefold() == "_pdf_gerados":
            partes = partes[1:]
        if partes and str(partes[-1]).lower().endswith(".pdf"):
            partes = partes[:-1]

        candidato = ""
        for parte in partes:
            parte_txt = str(parte or "").strip()
            if not parte_txt:
                continue
            if re.fullmatch(r"\d{8}", parte_txt):
                continue
            if re.fullmatch(r"\d{4}[-_/]\d{2}[-_/]\d{2}", parte_txt):
                continue
            if parte_txt.lower() == "_qrcodes":
                continue
            candidato = parte_txt
            break

        if not candidato:
            try:
                caminho = Path(str(caminho_arquivo or "")).expanduser()
                candidato = str(caminho.parent.name or "").strip()
            except Exception:
                candidato = ""
        slug = App._slug_url_texto(candidato)
        return slug or "funcionario"

    @staticmethod
    def _url_site_consulta_para_arquivo(caminho_arquivo):
        pasta = str(os.environ.get("CADNR_QR_GITHUB_DIR", "") or "").strip().strip("/")
        pages_base = App._normalizar_pages_base(os.environ.get("CADNR_QR_GITHUB_PAGES_BASE", ""))
        if not pages_base:
            pages_base = "https://elizangela2805.github.io/documentos"
        if not pages_base:
            return ""
        try:
            caminho = Path(str(caminho_arquivo or "")).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            repo_path = App._montar_caminho_repo_qr_github(caminho, pasta)
            if not repo_path:
                return ""
            funcionario = parse.quote(App._extrair_funcionario_slug_documento(repo_path, caminho), safe="")
            documento = parse.quote(App._extrair_documento_slug_documento(repo_path, caminho), safe="")
            data_ref = parse.quote(App._extrair_data_slug_documento(repo_path, caminho), safe="")
            return f"{pages_base}/{funcionario}/{documento}/{data_ref}"
        except Exception:
            return ""

    @staticmethod
    def _assinatura_rota_documento(repo_path, caminho_arquivo):
        funcionario = App._extrair_funcionario_slug_documento(repo_path, caminho_arquivo)
        documento = App._extrair_documento_slug_documento(repo_path, caminho_arquivo)
        data_ref = App._extrair_data_slug_documento(repo_path, caminho_arquivo)
        return {
            "funcionario": str(funcionario or "").strip().lower(),
            "documento": str(documento or "").strip().lower(),
            "data_ref": str(data_ref or "").strip().lower(),
        }

    def _atualizar_indice_documento_site(self, caminho_pdf):
        config = self._obter_config_qr_github()
        if not config:
            return False
        try:
            caminho = Path(str(caminho_pdf or "")).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            if caminho.suffix.lower() != ".pdf" or not caminho.exists() or not caminho.is_file():
                return False

            repo_path = self._montar_caminho_repo_qr_github(caminho, config.get("pasta", ""))
            if not repo_path:
                return False
            assinatura = self._assinatura_rota_documento(repo_path, caminho)
            consulta = self._url_site_consulta_para_arquivo(caminho)

            indice_rel = "_pdf_gerados/_indice_documentos.json"
            api_indice = (
                f"https://api.github.com/repos/{config['repo']}/contents/"
                f"{parse.quote(indice_rel, safe='/')}"
            )
            headers = {
                "Authorization": f"Bearer {config['token']}",
                "Accept": "application/vnd.github+json",
                "User-Agent": "CADNR/1.0",
                "X-GitHub-Api-Version": "2022-11-28",
            }

            sha_indice = None
            conteudo_atual = {}
            try:
                req_get = request.Request(
                    f"{api_indice}?ref={parse.quote(config['branch'])}",
                    headers=headers,
                    method="GET",
                )
                with request.urlopen(req_get, timeout=20) as resp:
                    dados_get = json.loads(resp.read().decode("utf-8"))
                sha_indice = str(dados_get.get("sha", "") or "").strip() or None
                conteudo_b64 = str(dados_get.get("content", "") or "").replace("\n", "")
                if conteudo_b64:
                    txt = base64.b64decode(conteudo_b64).decode("utf-8", errors="ignore")
                    parsed = json.loads(txt)
                    if isinstance(parsed, dict):
                        conteudo_atual = parsed
            except error.HTTPError as exc:
                if exc.code != 404:
                    self._qr_github_ultimo_erro = f"Indice GET HTTP {exc.code}"
            except Exception as exc:
                self._qr_github_ultimo_erro = f"Indice GET falhou: {exc}"

            itens = conteudo_atual.get("items", []) if isinstance(conteudo_atual, dict) else []
            if not isinstance(itens, list):
                itens = []
            repo_path_norm = str(repo_path or "").replace("\\", "/").strip("/")
            nome_arquivo = str(Path(repo_path_norm).name or "").strip()
            funcionario_exib = str(Path(repo_path_norm).parent.name or "").strip() or assinatura["funcionario"]
            rota = f"{assinatura['funcionario']}/{assinatura['documento']}/{assinatura['data_ref']}"

            novo_item = {
                "path": repo_path_norm,
                "nome": nome_arquivo,
                "funcionario": funcionario_exib,
                "consulta": consulta,
                "route_funcionario": assinatura["funcionario"],
                "route_documento": assinatura["documento"],
                "route_data": assinatura["data_ref"],
                "route": rota,
            }

            atualizou = False
            for idx, item in enumerate(itens):
                if not isinstance(item, dict):
                    continue
                path_item = str(item.get("path", "") or "").replace("\\", "/").strip("/")
                if path_item == repo_path_norm:
                    itens[idx] = {**item, **novo_item}
                    atualizou = True
                    break
            if not atualizou:
                itens.append(novo_item)

            payload_indice = {
                "updated_at": datetime.utcnow().isoformat(timespec="seconds") + "Z",
                "total": len(itens),
                "items": itens,
            }
            conteudo_novo_b64 = base64.b64encode(
                json.dumps(payload_indice, ensure_ascii=False, indent=2).encode("utf-8")
            ).decode("ascii")
            payload_put = {
                "message": f"CADNR indice update: {nome_arquivo or 'documento'}",
                "content": conteudo_novo_b64,
                "branch": config["branch"],
            }
            if sha_indice:
                payload_put["sha"] = sha_indice

            req_put = request.Request(
                api_indice,
                headers=headers,
                method="PUT",
                data=json.dumps(payload_put).encode("utf-8"),
            )
            with request.urlopen(req_put, timeout=30):
                pass
            return True
        except error.HTTPError as exc:
            detalhe = ""
            try:
                detalhe = exc.read().decode("utf-8", errors="ignore")
            except Exception:
                detalhe = ""
            self._qr_github_ultimo_erro = f"Indice PUT HTTP {exc.code} {detalhe[:120].strip()}"
            return False
        except Exception as exc:
            self._qr_github_ultimo_erro = f"Indice PUT falhou: {exc}"
            return False

    def _publicar_arquivo_no_site(self, caminho_arquivo, permitir_inexistente=False):
        config = self._obter_config_qr_github()
        if not config:
            return ""
        try:
            self._garantir_nojekyll_no_site(config)
        except Exception:
            pass

        caminho = Path(str(caminho_arquivo or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        if (not permitir_inexistente) and (not caminho.exists() or not caminho.is_file()):
            return ""
        repo_path = self._montar_caminho_repo_qr_github(caminho, config.get("pasta", ""))
        if not repo_path:
            return ""

        cache_chave = str(caminho.resolve())
        try:
            stat = caminho.stat()
            cache_tag = f"{int(stat.st_mtime)}:{int(stat.st_size)}"
        except OSError:
            cache_tag = "pending"
        cache_item = self._qr_github_cache.get(cache_chave)
        if isinstance(cache_item, dict) and cache_item.get("tag") == cache_tag and cache_item.get("url"):
            return str(cache_item.get("url", ""))

        if not caminho.exists() or not caminho.is_file():
            return ""
        try:
            conteudo_b64 = base64.b64encode(caminho.read_bytes()).decode("ascii")
        except OSError:
            return ""

        api_base = f"https://api.github.com/repos/{config['repo']}/contents/{parse.quote(repo_path, safe='/')}"
        headers = {
            "Authorization": f"Bearer {config['token']}",
            "Accept": "application/vnd.github+json",
            "User-Agent": "CADNR/1.0",
            "X-GitHub-Api-Version": "2022-11-28",
        }

        sha_existente = None
        try:
            req_get = request.Request(
                f"{api_base}?ref={parse.quote(config['branch'])}",
                headers=headers,
                method="GET",
            )
            with request.urlopen(req_get, timeout=20) as resp:
                atual = json.loads(resp.read().decode("utf-8"))
                sha_existente = str(atual.get("sha", "") or "").strip() or None
        except error.HTTPError as exc:
            if exc.code != 404:
                self._qr_github_ultimo_erro = f"Publicacao GET HTTP {exc.code}"
        except Exception:
            pass

        payload = {
            "message": f"CADNR QR update: {caminho.name}",
            "content": conteudo_b64,
            "branch": config["branch"],
        }
        if sha_existente:
            payload["sha"] = sha_existente

        try:
            req_put = request.Request(
                api_base,
                headers=headers,
                method="PUT",
                data=json.dumps(payload).encode("utf-8"),
            )
            with request.urlopen(req_put, timeout=30) as resp:
                dados = json.loads(resp.read().decode("utf-8"))
            content = dados.get("content", {}) if isinstance(dados, dict) else {}
            # Prioriza rota amigavel no site (documento/data) para abrir em nova guia.
            url = self._url_site_consulta_para_arquivo(caminho)
            if not url:
                url = self._url_github_raw_para_arquivo(caminho)
            if not url:
                url = self._url_github_pages_para_arquivo(caminho)
            if not url:
                url = str(content.get("download_url", "") or "").strip()
            if not url:
                url = f"{config['raw_base']}/{parse.quote(repo_path, safe='/')}"
            self._qr_github_cache[cache_chave] = {"tag": cache_tag, "url": url}
            self._qr_github_ultimo_erro = ""
            return url
        except error.HTTPError as exc:
            detalhe = ""
            try:
                detalhe = exc.read().decode("utf-8", errors="ignore")
            except Exception:
                detalhe = ""
            detalhe_curto = detalhe[:180].strip()
            msg = f"Publicacao PUT HTTP {exc.code} {detalhe_curto}"
            if exc.code == 404:
                msg += (
                    " | Verifique token com permissao de escrita em Contents "
                    "(Read and write) para o repositorio configurado."
                )
            self._qr_github_ultimo_erro = msg
        except Exception as exc:
            self._qr_github_ultimo_erro = f"Publicacao PUT falhou: {exc}"
        return ""

    def _garantir_nojekyll_no_site(self, config=None):
        cfg = config if isinstance(config, dict) else self._obter_config_qr_github()
        if not cfg:
            return False
        repo = str(cfg.get("repo", "") or "").strip()
        branch = str(cfg.get("branch", "") or "").strip() or "main"
        token = str(cfg.get("token", "") or "").strip()
        if not repo or not token:
            return False

        chave = f"{repo}@{branch}"
        if chave in self._github_nojekyll_ok:
            return True

        api_nojekyll = f"https://api.github.com/repos/{repo}/contents/.nojekyll"
        headers = {
            "Authorization": f"Bearer {token}",
            "Accept": "application/vnd.github+json",
            "User-Agent": "CADNR/1.0",
            "X-GitHub-Api-Version": "2022-11-28",
        }
        sha_existente = None
        try:
            req_get = request.Request(
                f"{api_nojekyll}?ref={parse.quote(branch)}",
                headers=headers,
                method="GET",
            )
            with request.urlopen(req_get, timeout=20) as resp:
                dados = json.loads(resp.read().decode("utf-8"))
            sha_existente = str(dados.get("sha", "") or "").strip() or None
        except error.HTTPError as exc:
            if exc.code != 404:
                self._qr_github_ultimo_erro = f".nojekyll GET HTTP {exc.code}"
                return False
        except Exception as exc:
            self._qr_github_ultimo_erro = f".nojekyll GET falhou: {exc}"
            return False

        if sha_existente:
            self._github_nojekyll_ok.add(chave)
            return True

        payload = {
            "message": "CADNR setup: add .nojekyll for Pages static files",
            "content": "",
            "branch": branch,
        }
        try:
            req_put = request.Request(
                api_nojekyll,
                headers=headers,
                method="PUT",
                data=json.dumps(payload).encode("utf-8"),
            )
            with request.urlopen(req_put, timeout=30):
                pass
            self._github_nojekyll_ok.add(chave)
            return True
        except error.HTTPError as exc:
            detalhe = ""
            try:
                detalhe = exc.read().decode("utf-8", errors="ignore")
            except Exception:
                detalhe = ""
            self._qr_github_ultimo_erro = f".nojekyll PUT HTTP {exc.code} {detalhe[:120].strip()}"
            return False
        except Exception as exc:
            self._qr_github_ultimo_erro = f".nojekyll PUT falhou: {exc}"
            return False

    def _espelhar_arquivo_no_repo_local(self, caminho_arquivo):
        try:
            origem = Path(str(caminho_arquivo or "")).expanduser()
            if not origem.is_absolute():
                origem = (Path(__file__).resolve().parent / origem).resolve()
            if not origem.exists() or not origem.is_file():
                return ""
            base = Path(__file__).resolve().parent.resolve()
            repo_path = self._montar_caminho_repo_qr_github(origem, os.environ.get("CADNR_QR_GITHUB_DIR", ""))
            if not repo_path:
                return ""
            destino = (base / Path(repo_path)).resolve()
            destino.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(origem, destino)
            return self._normalizar_caminho_documento_db(destino)
        except Exception:
            return ""

    def _url_github_qr_para_arquivo(self, caminho_arquivo, permitir_inexistente=False):
        caminho = Path(str(caminho_arquivo or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        if caminho.suffix.lower() != ".pdf":
            return ""
        return self._publicar_arquivo_no_site(caminho, permitir_inexistente=permitir_inexistente)

    def _on_app_close(self):
        try:
            if self._docs_monitor_after_id is not None:
                self.after_cancel(self._docs_monitor_after_id)
                self._docs_monitor_after_id = None
        except Exception:
            pass
        self._encerrar_servidor_qr_local()
        self.destroy()

    def _resolver_caminho_assinatura_pfx(self, slot=1):
        if int(slot) == 1:
            caminho_txt = str(self.assinatura_digital_pfx1 or self.assinatura_digital_pfx or "").strip()
        else:
            caminho_txt = str(self.assinatura_digital_pfx or "").strip()
        if not caminho_txt:
            return None
        caminho = Path(caminho_txt).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        return caminho

    def _resolver_caminho_assinatura_imagem1(self):
        caminho_txt = str(self.assinatura_digital_img1 or "").strip()
        if not caminho_txt:
            return None
        caminho = Path(caminho_txt).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        return caminho

    def _certificados_assinatura_configurados(self):
        certificados = []
        pfx1 = self._resolver_caminho_assinatura_pfx(1)
        if pfx1 is not None and pfx1.exists() and pfx1.is_file():
            certificados.append(
                {
                    "nome": "Projetta",
                    "pfx": str(pfx1),
                    "senha": str(self.assinatura_digital_senha1 or ""),
                }
            )
        pfx2 = self._resolver_caminho_assinatura_pfx(2)
        if pfx2 is not None and pfx2.exists() and pfx2.is_file():
            certificados.append(
                {
                    "nome": "Emerson",
                    "pfx": str(pfx2),
                    "senha": str(self.assinatura_digital_senha2 or self.assinatura_digital_senha or ""),
                }
            )
        for cert in self.assinatura_digital_certificados:
            if not isinstance(cert, dict):
                continue
            pfx_txt = str(cert.get("pfx", "") or "").strip()
            if not pfx_txt:
                continue
            pfx = Path(pfx_txt).expanduser()
            if not pfx.is_absolute():
                pfx = (Path(__file__).resolve().parent / pfx).resolve()
            if not pfx.exists() or not pfx.is_file():
                continue
            certificados.append(
                {
                    "nome": str(cert.get("nome", "") or pfx.stem).strip() or pfx.stem,
                    "pfx": str(pfx),
                    "senha": str(cert.get("senha", "") or ""),
                }
            )
        unicos = []
        vistos = set()
        for c in certificados:
            chave = str(c.get("pfx", "") or "").strip().lower()
            if not chave or chave in vistos:
                continue
            vistos.add(chave)
            unicos.append(c)
        return unicos

    def _obter_certificado_para_marcador(self, marcador):
        marcador = str(marcador or "").strip().lower()
        if marcador == "assinatura1":
            pfx = self._resolver_caminho_assinatura_pfx(1)
            if pfx is not None and pfx.exists() and pfx.is_file():
                return str(pfx), str(self.assinatura_digital_senha1 or ""), "Projetta"
        if marcador == "assinatura2":
            pfx = self._resolver_caminho_assinatura_pfx(2)
            if pfx is not None and pfx.exists() and pfx.is_file():
                return str(pfx), str(self.assinatura_digital_senha2 or self.assinatura_digital_senha or ""), "Emerson"
        certificados = self._certificados_assinatura_configurados()
        if certificados:
            c = certificados[0]
            return str(c.get("pfx", "")), str(c.get("senha", "")), str(c.get("nome", "Certificado"))
        return "", "", ""

    def _assinatura_digital_ativa(self):
        if not bool(self.assinatura_digital_habilitada):
            return False
        caminho = self._resolver_caminho_assinatura_pfx(1)
        return bool(caminho is not None and caminho.exists() and caminho.is_file())

    def _assinar_pdf_digital(self, caminho_pdf):
        caminho = Path(str(caminho_pdf or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        if not caminho.exists() or not caminho.is_file() or caminho.suffix.lower() != ".pdf":
            return False, "arquivo PDF invalido para assinatura"

        pfx = self._resolver_caminho_assinatura_pfx(1)
        if pfx is None or not pfx.exists() or not pfx.is_file():
            return False, "certificado PFX da Projetta nao encontrado"
        senha = str(self.assinatura_digital_senha1 or self.assinatura_digital_senha or "")
        return self._assinar_pdf_digital_com_cert(
            caminho_pdf,
            pfx_path=pfx,
            senha=senha,
            field_name="SignatureProjetta",
            reason="Assinatura digital Projetta",
        )

    def _assinar_pdf_digital_com_cert(self, caminho_pdf, pfx_path, senha="", field_name="Signature1", reason=""):
        caminho = Path(str(caminho_pdf or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        pfx = Path(str(pfx_path or "")).expanduser()
        if not pfx.is_absolute():
            pfx = (Path(__file__).resolve().parent / pfx).resolve()
        try:
            from pyhanko.sign import signers
            from pyhanko.pdf_utils.incremental_writer import IncrementalPdfFileWriter
        except Exception:
            return False, "biblioteca pyHanko nao instalada (pip install pyHanko)"

        senha_bytes = senha.encode("utf-8") if senha else b""
        signer = signers.SimpleSigner.load_pkcs12(
            pfx_file=str(pfx),
            passphrase=senha_bytes,
        )
        if signer is None:
            return False, "nao foi possivel carregar certificado PFX"

        for tentativa in range(2):
            temp_saida = caminho.with_name(f"{caminho.stem}.__signed__.pdf")
            try:
                with open(caminho, "rb") as arq_in, open(temp_saida, "wb") as arq_out:
                    writer = IncrementalPdfFileWriter(arq_in)
                    meta = signers.PdfSignatureMetadata(
                        field_name=str(field_name or "Signature1"),
                        reason=str(reason or "Assinatura digital CADNR"),
                    )
                    pdf_signer = signers.PdfSigner(meta, signer=signer)
                    pdf_signer.sign_pdf(writer, output=arq_out)

                shutil.move(str(temp_saida), str(caminho))
                return True, ""
            except Exception as ex:
                try:
                    if temp_saida.exists():
                        temp_saida.unlink()
                except OSError:
                    pass
                msg = str(ex or "")
                msg_norm = msg.lower()
                erro_xmp = ("xmp" in msg_norm) and ("parse" in msg_norm)
                if tentativa == 0 and erro_xmp:
                    self._recriar_pdf_sem_xmp(caminho)
                    continue
                return False, msg
        return False, "falha desconhecida na assinatura digital"

    def _localizar_marcadores_assinatura_pdf(self, caminho_pdf):
        try:
            import fitz
        except Exception:
            return []
        caminho = Path(str(caminho_pdf or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        if not caminho.exists() or not caminho.is_file():
            return []
        encontrados = []
        try:
            doc = fitz.open(str(caminho))
            for pidx, pagina in enumerate(doc):
                for marcador in ("assinatura1", "assinatura2"):
                    rects = []
                    for termo in (marcador, marcador.upper(), marcador.title()):
                        try:
                            rects.extend(pagina.search_for(termo))
                        except Exception:
                            pass
                    vistos = set()
                    for r in rects:
                        chave = (round(r.x0, 2), round(r.y0, 2), round(r.x1, 2), round(r.y1, 2))
                        if chave in vistos:
                            continue
                        vistos.add(chave)
                        encontrados.append(
                            {
                                "marcador": marcador,
                                "pagina": pidx,
                                "rect": fitz.Rect(r.x0, r.y0, r.x1, r.y1),
                            }
                        )
            doc.close()
        except Exception:
            return []
        return encontrados

    def _aplicar_visual_assinatura_em_marcador_pdf(self, caminho_pdf, marcador_info):
        try:
            import fitz
        except Exception:
            return False, "PyMuPDF (fitz) indisponivel"
        caminho = Path(str(caminho_pdf or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        if not caminho.exists() or not caminho.is_file():
            return False, "arquivo PDF nao encontrado"

        marcador = str((marcador_info or {}).get("marcador", "") or "").strip().lower()
        pagina_idx = int((marcador_info or {}).get("pagina", 0) or 0)
        rect_ref = marcador_info.get("rect", None)
        if rect_ref is None:
            return False, "marcador de assinatura invalido"

        temp_saida = caminho.with_name(f"{caminho.stem}.__vis__.pdf")
        try:
            doc = fitz.open(str(caminho))
            if pagina_idx < 0 or pagina_idx >= doc.page_count:
                doc.close()
                return False, "pagina do marcador fora do intervalo"
            pagina = doc[pagina_idx]

            rect = fitz.Rect(rect_ref)
            pagina.draw_rect(rect, fill=(1, 1, 1), color=(1, 1, 1), overlay=True)

            if marcador == "assinatura1":
                # Card visual maior para assinatura1 (modelo Projetta).
                larg = 260.0
                alt = 150.0
                x0 = max(0.0, rect.x0)
                y0 = max(0.0, rect.y0 - 6.0)
                box = fitz.Rect(x0, y0, x0 + larg, y0 + alt)
                img = self._resolver_caminho_assinatura_imagem1()
                if img is not None and img.exists() and img.is_file():
                    pagina.insert_image(box, filename=str(img), keep_proportion=True, overlay=True)
                else:
                    pagina.draw_rect(box, color=(0, 0, 0.4), width=0.5, fill=(1, 1, 1), overlay=True)
                    box_txt = fitz.Rect(box.x0 + 6, box.y0 + 4, box.x1 - 6, box.y1 - 4)
                    pagina.insert_textbox(
                        box_txt,
                        "ASSINADO DIGITALMENTE\nPROJETTA ENGENHARIA",
                        fontsize=10,
                        color=(0, 0.13, 0.42),
                        align=0,
                    )
            else:
                larg = 290.0
                alt = 120.0
                x0 = max(0.0, rect.x0)
                y0 = max(0.0, rect.y0 - 6.0)
                box = fitz.Rect(x0, y0, x0 + larg, y0 + alt)
                pagina.draw_rect(box, color=(1, 1, 1), width=0, fill=(1, 1, 1), overlay=True)
                dt_ass = datetime.now().strftime("%d/%m/%Y %H:%M:%S -03:00")
                topo_esq = fitz.Point(box.x0 + 6, box.y0 + 14)
                topo_dir = fitz.Point(box.x0 + 154, box.y0 + 14)
                pagina.insert_text(
                    topo_esq,
                    "EMERSON PEREIRA DE\nOLIVEIRA:37232085842",
                    fontsize=8.7,
                    color=(0, 0, 0),
                )
                pagina.insert_text(
                    topo_dir,
                    f"Assinado de forma digital por\n"
                    f"EMERSON PEREIRA DE\n"
                    f"OLIVEIRA:37232085842\n"
                    f"Dados: {dt_ass}",
                    fontsize=6.3,
                    color=(0, 0, 0),
                )
                y_linha = box.y0 + 42
                pagina.draw_line(
                    fitz.Point(box.x0 + 6, y_linha),
                    fitz.Point(box.x1 - 6, y_linha),
                    color=(0, 0, 0),
                    width=0.8,
                )
                box_txt = fitz.Rect(box.x0 + 8, box.y0 + 45, box.x1 - 8, box.y1 - 6)
                pagina.insert_textbox(
                    box_txt,
                    "INSTRUTOR E RESPONSÁVEL TÉCNICO\n"
                    "Emerson Pereira de Oliveira\n"
                    "Engenheiro Mecânico\n"
                    "Eng. Segurança do Trabalho\n"
                    "CREASP 5069946034\n"
                    "RNP 2616205491",
                    fontsize=7.7,
                    color=(0, 0, 0),
                    align=1,
                )

            doc.save(str(temp_saida), garbage=4, deflate=True)
            doc.close()
            shutil.move(str(temp_saida), str(caminho))
            return True, ""
        except Exception as ex:
            try:
                if temp_saida.exists():
                    temp_saida.unlink()
            except OSError:
                pass
            return False, str(ex)

    def _assinar_pdf_por_marcadores(self, caminho_pdf):
        # Modo simplificado: assinatura invisivel unica com certificado da Projetta.
        return self._assinar_pdf_digital(caminho_pdf)

    def _abrir_configuracao_assinatura_digital(self):
        popup = CadastroPopup(self, "Assinatura Digital")
        popup.columnconfigure(0, weight=1)

        r = 0
        popup.secao("Certificado Projetta (invisivel)", r)
        r += 1
        frame_pfx1 = ttk.Frame(popup)
        frame_pfx1.grid(row=r, column=0, sticky="ew", padx=12)
        frame_pfx1.columnconfigure(1, weight=1)
        ttk.Label(frame_pfx1, text="Arquivo PFX:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        pfx1_var = tk.StringVar(value=str(self.assinatura_digital_pfx1 or ""))
        ttk.Entry(frame_pfx1, textvariable=pfx1_var, state="readonly").grid(row=0, column=1, sticky="ew")

        def escolher_pfx1():
            caminho = filedialog.askopenfilename(
                title="Selecionar certificado Projetta (.pfx/.p12)",
                filetypes=[("Certificados", "*.pfx *.p12"), ("Todos os arquivos", "*.*")],
            )
            if caminho:
                pfx1_var.set(caminho)

        ttk.Button(frame_pfx1, text="Pesquisar...", command=escolher_pfx1).grid(row=0, column=2, padx=(8, 0))

        r += 1
        frame_senha1 = ttk.Frame(popup)
        frame_senha1.grid(row=r, column=0, sticky="ew", padx=12, pady=(6, 0))
        frame_senha1.columnconfigure(1, weight=1)
        ttk.Label(frame_senha1, text="Senha PFX:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        senha1_var = tk.StringVar(value=str(self.assinatura_digital_senha1 or ""))
        ttk.Entry(frame_senha1, textvariable=senha1_var, show="*").grid(row=0, column=1, sticky="ew")

        r += 1
        frame_flag = ttk.Frame(popup)
        frame_flag.grid(row=r, column=0, sticky="ew", padx=12, pady=(8, 0))
        habilitada_var = tk.BooleanVar(value=bool(self.assinatura_digital_habilitada))
        ttk.Checkbutton(
            frame_flag,
            text="Habilitar assinatura digital invisivel",
            variable=habilitada_var,
        ).grid(row=0, column=0, sticky="w")

        r += 1
        aviso = ttk.Label(
            popup,
            text=(
                "A assinatura digital requer pyHanko instalado.\n"
                "Todos os PDFs serao assinados digitalmente em modo invisivel."
            ),
            justify="left",
        )
        aviso.grid(row=r, column=0, sticky="w", padx=12, pady=(8, 0))

        r += 1
        botoes = ttk.Frame(popup)
        botoes.grid(row=r, column=0, sticky="e", padx=12, pady=(16, 10))

        def salvar_config():
            habilitar = bool(habilitada_var.get())
            pfx1_txt = str(pfx1_var.get() or "").strip()
            senha1_txt = str(senha1_var.get() or "")
            if habilitar:
                if not pfx1_txt:
                    messagebox.showerror("Assinatura Digital", "Selecione o arquivo PFX da Projetta.")
                    return
                ppath = Path(pfx1_txt).expanduser()
                if not ppath.is_absolute():
                    ppath = (Path(__file__).resolve().parent / ppath).resolve()
                if not ppath.exists() or not ppath.is_file():
                    messagebox.showerror("Assinatura Digital", "Arquivo PFX nao encontrado.")
                    return
                if ppath.suffix.lower() not in {".pfx", ".p12"}:
                    messagebox.showerror("Assinatura Digital", "Selecione um arquivo .pfx/.p12.")
                    return
                self.assinatura_digital_pfx1 = str(Path(pfx1_txt).expanduser().resolve()) if pfx1_txt else ""
            else:
                self.assinatura_digital_pfx1 = pfx1_txt
            self.assinatura_digital_pfx2 = ""
            self.assinatura_digital_img1 = ""
            self.assinatura_digital_habilitada = habilitar
            self.assinatura_digital_senha1 = senha1_txt
            self.assinatura_digital_senha2 = ""
            self.assinatura_digital_certificados = []
            # Compatibilidade com campo antigo.
            self.assinatura_digital_pfx = self.assinatura_digital_pfx1
            self.assinatura_digital_senha = self.assinatura_digital_senha1
            self._aviso_assinatura_exibido = False
            self._salvar_dados()
            messagebox.showinfo("Assinatura Digital", "Configuracao salva com sucesso.")
            popup.destroy()

        ttk.Button(botoes, text="Cancelar", command=popup.destroy).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(botoes, text="Salvar", command=salvar_config).grid(row=0, column=1)

        popup.ajustar_tamanho()

    def _abrir_configuracao_github(self):
        popup = CadastroPopup(self, "Config Site")
        popup.columnconfigure(0, weight=1)

        r = 0
        popup.secao("Publicacao automatica e QR (Site de Publicacao)", r)

        r += 1
        f_token = ttk.Frame(popup)
        f_token.grid(row=r, column=0, sticky="ew", padx=12)
        f_token.columnconfigure(1, weight=1)
        ttk.Label(f_token, text="Token:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        token_var = tk.StringVar(value=str(self.github_token or os.environ.get("CADNR_QR_GITHUB_TOKEN", "") or ""))
        ttk.Entry(f_token, textvariable=token_var, show="*").grid(row=0, column=1, sticky="ew")

        r += 1
        f_repo = ttk.Frame(popup)
        f_repo.grid(row=r, column=0, sticky="ew", padx=12, pady=(6, 0))
        f_repo.columnconfigure(1, weight=1)
        ttk.Label(f_repo, text="Repo:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        repo_var = tk.StringVar(value=str(self.github_repo or "Elizangela2805/documentos"))
        ttk.Entry(f_repo, textvariable=repo_var).grid(row=0, column=1, sticky="ew")

        r += 1
        f_branch = ttk.Frame(popup)
        f_branch.grid(row=r, column=0, sticky="ew", padx=12, pady=(6, 0))
        f_branch.columnconfigure(1, weight=1)
        ttk.Label(f_branch, text="Branch:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        branch_var = tk.StringVar(value=str(self.github_branch or "main"))
        ttk.Entry(f_branch, textvariable=branch_var).grid(row=0, column=1, sticky="ew")

        r += 1
        f_dir = ttk.Frame(popup)
        f_dir.grid(row=r, column=0, sticky="ew", padx=12, pady=(6, 0))
        f_dir.columnconfigure(1, weight=1)
        ttk.Label(f_dir, text="Pasta no site:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        dir_var = tk.StringVar(value=str(self.github_dir or "_pdf_gerados"))
        ttk.Entry(f_dir, textvariable=dir_var).grid(row=0, column=1, sticky="ew")

        r += 1
        f_pages = ttk.Frame(popup)
        f_pages.grid(row=r, column=0, sticky="ew", padx=12, pady=(6, 0))
        f_pages.columnconfigure(1, weight=1)
        ttk.Label(f_pages, text="Base Pages:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        pages_var = tk.StringVar(value=str(self.github_pages_base or "https://elizangela2805.github.io/documentos"))
        ttk.Entry(f_pages, textvariable=pages_var).grid(row=0, column=1, sticky="ew")

        r += 1
        aviso = ttk.Label(
            popup,
            text="Essas configuracoes sao usadas para upload automatico dos PDFs e links de QR.",
        )
        aviso.grid(row=r, column=0, sticky="w", padx=12, pady=(8, 0))

        r += 1
        botoes = ttk.Frame(popup)
        botoes.grid(row=r, column=0, sticky="e", padx=12, pady=(12, 12))

        def salvar_config():
            repo_txt = str(repo_var.get() or "").strip() or "Elizangela2805/documentos"
            branch_txt = str(branch_var.get() or "").strip() or "main"
            dir_txt = str(dir_var.get() or "").strip().strip("/")
            pages_txt = str(pages_var.get() or "").strip() or "https://elizangela2805.github.io/documentos"
            token_txt = str(token_var.get() or "").strip()
            repo_norm = self._normalizar_repo_github(repo_txt)
            if not repo_norm:
                messagebox.showerror("Config Site", "Repositorio invalido. Use formato dono/repositorio.")
                return
            if not pages_txt.lower().startswith("http://") and not pages_txt.lower().startswith("https://"):
                messagebox.showerror("Config Site", "Base publica invalida. Informe URL iniciando com http(s).")
                return
            self.github_repo = repo_norm
            self.github_branch = branch_txt
            self.github_dir = dir_txt or "_pdf_gerados"
            self.github_pages_base = self._normalizar_pages_base(pages_txt)
            self.github_token = token_txt
            self._aplicar_configuracao_github_ambiente()
            self._salvar_dados()
            messagebox.showinfo("Config Site", "Configuracao salva com sucesso.")
            popup.destroy()

        def testar_publicacao():
            repo_txt = str(repo_var.get() or "").strip() or "Elizangela2805/documentos"
            branch_txt = str(branch_var.get() or "").strip() or "main"
            token_txt = str(token_var.get() or "").strip()
            repo_norm = self._normalizar_repo_github(repo_txt)
            if not repo_norm:
                messagebox.showerror("Config Site", "Repositorio invalido. Use formato dono/repositorio.")
                return
            if not token_txt:
                messagebox.showerror("Config Site", "Informe o token para testar a publicacao.")
                return

            headers = {
                "Authorization": f"Bearer {token_txt}",
                "Accept": "application/vnd.github+json",
                "User-Agent": "CADNR/1.0",
                "X-GitHub-Api-Version": "2022-11-28",
            }

            try:
                req_repo = request.Request(
                    f"https://api.github.com/repos/{repo_norm}",
                    headers=headers,
                    method="GET",
                )
                with request.urlopen(req_repo, timeout=20):
                    pass
            except error.HTTPError as exc:
                detalhe = ""
                try:
                    detalhe = exc.read().decode("utf-8", errors="ignore")
                except Exception:
                    detalhe = ""
                detalhe = detalhe[:220].strip()
                sufixo = f"\n\nDetalhe: {detalhe}" if detalhe else ""
                messagebox.showerror(
                    "Teste Publicacao",
                    f"Falha ao acessar o repositorio ({exc.code}).{sufixo}",
                )
                return
            except Exception as exc:
                messagebox.showerror("Teste Publicacao", f"Falha de conexao: {exc}")
                return

            probe_nome = datetime.now().strftime("cadnr_probe_%Y%m%d_%H%M%S.txt")
            probe_path = f"_cadnr_probe/{probe_nome}"
            probe_api = f"https://api.github.com/repos/{repo_norm}/contents/{parse.quote(probe_path, safe='/')}"
            probe_b64 = base64.b64encode(
                f"cadnr probe {datetime.now().isoformat(timespec='seconds')}".encode("utf-8")
            ).decode("ascii")
            payload_put = {
                "message": f"CADNR probe upload: {probe_nome}",
                "content": probe_b64,
                "branch": branch_txt,
            }

            sha_probe = ""
            try:
                req_put = request.Request(
                    probe_api,
                    headers=headers,
                    method="PUT",
                    data=json.dumps(payload_put).encode("utf-8"),
                )
                with request.urlopen(req_put, timeout=25) as resp:
                    dados = json.loads(resp.read().decode("utf-8"))
                if isinstance(dados, dict):
                    content = dados.get("content", {}) if isinstance(dados.get("content"), dict) else {}
                    sha_probe = str(content.get("sha", "") or "").strip()
            except error.HTTPError as exc:
                detalhe = ""
                try:
                    detalhe = exc.read().decode("utf-8", errors="ignore")
                except Exception:
                    detalhe = ""
                detalhe = detalhe[:220].strip()
                sufixo = f"\n\nDetalhe: {detalhe}" if detalhe else ""
                if exc.code == 404:
                    sufixo += (
                        "\n\nVerifique permissao de escrita em Contents (Read and write)"
                        " para o repositorio."
                    )
                messagebox.showerror(
                    "Teste Publicacao",
                    f"Falha ao publicar arquivo de teste ({exc.code}).{sufixo}",
                )
                return
            except Exception as exc:
                messagebox.showerror("Teste Publicacao", f"Falha ao publicar teste: {exc}")
                return

            aviso_limpeza = ""
            if sha_probe:
                payload_del = {
                    "message": f"CADNR probe cleanup: {probe_nome}",
                    "branch": branch_txt,
                    "sha": sha_probe,
                }
                try:
                    req_del = request.Request(
                        probe_api,
                        headers=headers,
                        method="DELETE",
                        data=json.dumps(payload_del).encode("utf-8"),
                    )
                    with request.urlopen(req_del, timeout=20):
                        pass
                except Exception as exc:
                    aviso_limpeza = f"\n\nAviso: teste publicado, mas nao foi possivel remover o arquivo de prova ({exc})."

            messagebox.showinfo(
                "Teste Publicacao",
                "Publicacao validada com sucesso.\n"
                f"Repo: {repo_norm}\n"
                f"Branch: {branch_txt}{aviso_limpeza}",
            )

        ttk.Button(botoes, text="Cancelar", command=popup.destroy).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(botoes, text="Testar Publicacao", command=testar_publicacao).grid(row=0, column=1, padx=(0, 8))
        ttk.Button(botoes, text="Salvar", command=salvar_config).grid(row=0, column=2)

        popup.ajustar_tamanho()

    def _testar_link_qr_local(self):
        caminho_pdf = None
        for item in reversed(self.documentos_salvos):
            caminho_txt = self._normalizar_caminho_documento_db(item.get("caminho", ""))
            if not caminho_txt:
                continue
            caminho = Path(caminho_txt).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            if caminho.suffix.lower() != ".pdf":
                continue
            if caminho.exists() and caminho.is_file():
                caminho_pdf = caminho
                break

        if caminho_pdf is not None:
            url_github = self._url_github_qr_para_arquivo(caminho_pdf)
            if url_github:
                titulo = "QR - Site"
                msg = (
                    "PDF publicado no site para leitura via QR.\n\n"
                    f"Arquivo: {caminho_pdf.name}\n"
                    f"URL: {url_github}\n\n"
                    "O link foi copiado para a area de transferencia."
                )
                try:
                    self.clipboard_clear()
                    self.clipboard_append(url_github)
                except Exception:
                    pass
                try:
                    self._abrir_url_no_chrome(url_github)
                except Exception:
                    pass
                messagebox.showinfo(titulo, msg)
                return

        if not self._iniciar_servidor_qr_local():
            detalhe_gh = str(self._qr_github_ultimo_erro or "").strip()
            sufixo = f"\n\nDetalhe da publicacao: {detalhe_gh}" if detalhe_gh else ""
            messagebox.showerror(
                "QR",
                "Nao foi possivel iniciar o servidor local para testes de QR."
                f"{sufixo}",
            )
            return

        if caminho_pdf is not None:
            url_teste = self._url_local_qr_para_arquivo(caminho_pdf)
            titulo = "QR - Teste"
            msg = (
                "Link QR local pronto.\n\n"
                f"Arquivo: {caminho_pdf.name}\n"
                f"URL: {url_teste}\n\n"
                "O link foi copiado para a area de transferencia."
            )
        else:
            url_teste = f"{self._qr_http_base_url}/__cadnr_ping__"
            titulo = "QR - Servidor Local"
            msg = (
                "Servidor local iniciado, mas nenhum PDF salvo foi encontrado para teste.\n\n"
                f"URL de teste: {url_teste}\n\n"
                "Gere um PDF e teste novamente para validar o QR completo."
            )

        try:
            self.clipboard_clear()
            self.clipboard_append(url_teste)
        except Exception:
            pass
        try:
            self._abrir_url_no_chrome(url_teste)
        except Exception:
            pass
        messagebox.showinfo(titulo, msg)

    def _abrir_links_pdfs_publicados(self):
        urls = []
        vistos = set()
        for item in self.documentos_salvos:
            if not isinstance(item, dict):
                continue
            caminho_txt = self._normalizar_caminho_documento_db(item.get("caminho", ""))
            if not caminho_txt:
                continue
            caminho = Path(caminho_txt).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            if caminho.suffix.lower() != ".pdf":
                continue
            url = self._url_site_consulta_para_arquivo(caminho)
            if not url:
                url = self._url_github_pages_para_arquivo(caminho)
            if not url:
                url = self._url_github_raw_para_arquivo(caminho)
            if not url:
                continue
            chave = str(url).strip()
            if not chave or chave in vistos:
                continue
            vistos.add(chave)
            urls.append(chave)

        if not urls:
            messagebox.showwarning("Links PDF", "Nenhum link de PDF foi gerado.")
            return

        total_abertas = 0
        limite = 30
        for url in urls[:limite]:
            try:
                if self._abrir_url_no_chrome(url):
                    total_abertas += 1
            except Exception:
                continue

        msg = (
            f"Links gerados: {len(urls)}\n"
            f"Abas abertas: {total_abertas}"
        )
        if len(urls) > limite:
            msg += f"\n\nForam abertas as primeiras {limite} abas."
        messagebox.showinfo("Links PDF", msg)

    def _abrir_projetta_html_no_chrome(self):
        base = Path(__file__).resolve().parent
        candidatos = [
            (base / "index.html").resolve(),
            (base / "projetta.html").resolve(),
        ]
        caminho_html = next((p for p in candidatos if p.exists() and p.is_file()), None)
        if caminho_html is None:
            messagebox.showwarning(
                "Site de Uploads",
                f"Arquivo nao encontrado:\n{candidatos[0]}",
            )
            return
        try:
            url_html = caminho_html.as_uri()
        except ValueError:
            url_html = str(caminho_html)
        self._abrir_url_no_chrome(url_html)

    @staticmethod
    def _obter_arquivo_pdf_livre(pasta_destino, nome_base, reutilizar_existente=False):
        candidato = pasta_destino / f"{nome_base}.pdf"
        if reutilizar_existente or not candidato.exists():
            return candidato
        i = 1
        while True:
            candidato = pasta_destino / f"{nome_base} ({i}).pdf"
            if not candidato.exists():
                return candidato
            i += 1

    def _criar_pasta_empresa_cadnr(self, nome_pasta):
        nome_limpo = self._obter_nome_arquivo_seguro(nome_pasta, "")
        if not nome_limpo:
            return None
        destino = Path(__file__).resolve().parent / nome_limpo
        destino.mkdir(parents=True, exist_ok=True)
        return destino

    def _garantir_pastas_empresas(self):
        for empresa in self.empresas:
            if not isinstance(empresa, dict):
                continue
            nome_ref = str(empresa.get("nome_pasta", "") or empresa.get("nome", "") or "").strip()
            nome_seguro = self._obter_nome_arquivo_seguro(nome_ref, "")
            if not nome_seguro:
                continue
            try:
                self._criar_pasta_empresa_cadnr(nome_seguro)
            except OSError:
                continue

    @staticmethod
    def _pasta_logos_empresas():
        return Path(__file__).resolve().parent / "_logos_empresas"

    @staticmethod
    def _resolver_logo_empresa(logo_ref):
        logo_txt = str(logo_ref or "").strip()
        if not logo_txt:
            return None
        caminho_logo = Path(logo_txt)
        if not caminho_logo.is_absolute():
            caminho_logo = Path(__file__).resolve().parent / caminho_logo
        return caminho_logo

    def _caminho_logo_empresa(self, empresa):
        if not isinstance(empresa, dict):
            return ""
        caminho_logo = self._resolver_logo_empresa(empresa.get("logo", ""))
        if caminho_logo is None:
            return ""
        try:
            return str(caminho_logo.resolve()) if caminho_logo.exists() else ""
        except OSError:
            return ""

    def _caminho_foto_funcionario(self, funcionario):
        if not isinstance(funcionario, dict):
            return ""
        caminho_foto = self._resolver_foto_funcionario(funcionario.get("foto", ""))
        if caminho_foto is None:
            return ""
        try:
            return str(caminho_foto.resolve()) if caminho_foto.exists() else ""
        except OSError:
            return ""

    def _salvar_logo_empresa(self, arquivo_origem, empresa_id, nome_referencia=""):
        origem = Path(str(arquivo_origem or "")).expanduser()
        if not origem.exists() or not origem.is_file():
            raise OSError("Arquivo de logo invalido.")

        extensoes_validas = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"}
        ext = origem.suffix.lower()
        if ext not in extensoes_validas:
            raise OSError("Formato de imagem nao suportado para logo.")

        base_dir = Path(__file__).resolve().parent
        pasta_logos = self._pasta_logos_empresas()
        pasta_logos.mkdir(parents=True, exist_ok=True)

        nome_seguro = self._obter_nome_arquivo_seguro(nome_referencia, "") or f"empresa_{empresa_id}"
        destino = pasta_logos / f"{empresa_id}_{nome_seguro}{ext}"
        shutil.copy2(origem, destino)
        return str(destino.relative_to(base_dir)).replace("\\", "/")

    def _remover_logo_empresa(self, logo_ref):
        caminho_logo = self._resolver_logo_empresa(logo_ref)
        if caminho_logo is None:
            return
        try:
            caminho_resolvido = caminho_logo.resolve()
            pasta_logos = self._pasta_logos_empresas().resolve()
            try:
                caminho_resolvido.relative_to(pasta_logos)
            except ValueError:
                return
            if caminho_resolvido.exists() and caminho_resolvido.is_file():
                caminho_resolvido.unlink()
        except OSError:
            return

    @staticmethod
    def _pasta_fotos_funcionarios():
        return Path(__file__).resolve().parent / "_fotos_funcionarios"

    @staticmethod
    def _resolver_foto_funcionario(foto_ref):
        foto_txt = str(foto_ref or "").strip()
        if not foto_txt:
            return None
        caminho_foto = Path(foto_txt)
        if not caminho_foto.is_absolute():
            caminho_foto = Path(__file__).resolve().parent / caminho_foto
        return caminho_foto

    def _salvar_foto_funcionario(self, arquivo_origem, funcionario_id, nome_referencia=""):
        origem = Path(str(arquivo_origem or "")).expanduser()
        if not origem.exists() or not origem.is_file():
            raise OSError("Arquivo de foto invalido.")

        extensoes_validas = {".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"}
        ext = origem.suffix.lower()
        if ext not in extensoes_validas:
            raise OSError("Formato de imagem nao suportado para foto.")

        base_dir = Path(__file__).resolve().parent
        pasta_fotos = self._pasta_fotos_funcionarios()
        pasta_fotos.mkdir(parents=True, exist_ok=True)

        nome_seguro = self._obter_nome_arquivo_seguro(nome_referencia, "") or f"funcionario_{funcionario_id}"
        destino = pasta_fotos / f"{funcionario_id}_{nome_seguro}{ext}"
        shutil.copy2(origem, destino)
        return str(destino.relative_to(base_dir)).replace("\\", "/")

    def _remover_foto_funcionario(self, foto_ref):
        caminho_foto = self._resolver_foto_funcionario(foto_ref)
        if caminho_foto is None:
            return
        try:
            caminho_resolvido = caminho_foto.resolve()
            pasta_fotos = self._pasta_fotos_funcionarios().resolve()
            try:
                caminho_resolvido.relative_to(pasta_fotos)
            except ValueError:
                return
            if caminho_resolvido.exists() and caminho_resolvido.is_file():
                caminho_resolvido.unlink()
        except OSError:
            return

    def _salvar_documento_no_projeto(
        self,
        arquivo_origem,
        empresa_nome="",
        tipo_documento="documento",
        empresa_id=None,
        funcionario_id=None,
    ):
        origem = Path(str(arquivo_origem or "")).expanduser()
        if not origem.exists() or not origem.is_file():
            raise OSError("Arquivo de documento invalido.")
        base_dir = Path(__file__).resolve().parent
        nome_empresa = self._obter_nome_arquivo_seguro(empresa_nome, "") or "projetta"
        pasta_docs = base_dir / nome_empresa
        pasta_docs.mkdir(parents=True, exist_ok=True)

        # Mantem o nome original do arquivo selecionado no cadastro.
        destino = pasta_docs / origem.name
        try:
            if destino.exists() and destino.resolve() == origem.resolve():
                destino_rel = str(destino.relative_to(base_dir)).replace("\\", "/")
                self._registrar_documento_salvo(
                    destino_rel,
                    origem="cadastro_documento",
                    empresa_id=empresa_id,
                    funcionario_id=funcionario_id,
                    tipo_documento=tipo_documento,
                )
                return destino_rel
        except OSError:
            pass

        if destino.exists():
            try:
                if filecmp.cmp(str(origem), str(destino), shallow=False):
                    destino_rel = str(destino.relative_to(base_dir)).replace("\\", "/")
                    self._registrar_documento_salvo(
                        destino_rel,
                        origem="cadastro_documento",
                        empresa_id=empresa_id,
                        funcionario_id=funcionario_id,
                        tipo_documento=tipo_documento,
                    )
                    return destino_rel
            except OSError:
                pass
            i = 1
            ext = origem.suffix
            stem = origem.stem
            while True:
                candidato = pasta_docs / f"{stem}_{i}{ext}"
                if not candidato.exists():
                    destino = candidato
                    break
                i += 1
        shutil.copy2(origem, destino)
        destino_rel = str(destino.relative_to(base_dir)).replace("\\", "/")
        self._registrar_documento_salvo(
            destino_rel,
            origem="cadastro_documento",
            empresa_id=empresa_id,
            funcionario_id=funcionario_id,
            tipo_documento=tipo_documento,
        )
        return destino_rel

    @staticmethod
    def _iso_datetime_now():
        return datetime.now().isoformat(timespec="seconds")

    def _normalizar_caminho_documento_db(self, caminho):
        caminho_txt = str(caminho or "").strip()
        if not caminho_txt:
            return ""
        # Compatibilidade com formatos antigos:
        # - documentos/certificados/arquivo.ext
        # - documentos/outros/arquivo.ext
        # - certificados/arquivo.ext
        # - outros/arquivo.ext
        caminho_antigo = caminho_txt.replace("\\", "/")
        if caminho_antigo.lower().startswith("documentos/"):
            sufixo = caminho_antigo[len("documentos/") :]
            if "/" in sufixo:
                caminho_txt = f"projetta/{sufixo.split('/', 1)[1]}"
            else:
                caminho_txt = f"projetta/{sufixo}"
        elif caminho_antigo.lower().startswith("certificados/"):
            caminho_txt = f"projetta/{caminho_antigo.split('/', 1)[1]}"
        elif caminho_antigo.lower().startswith("outros/"):
            caminho_txt = f"projetta/{caminho_antigo.split('/', 1)[1]}"
        base_dir = Path(__file__).resolve().parent
        p = Path(caminho_txt).expanduser()
        try:
            if p.is_absolute():
                resolvido = p.resolve()
                try:
                    return str(resolvido.relative_to(base_dir)).replace("\\", "/")
                except ValueError:
                    return str(resolvido)
            resolvido_rel = (base_dir / p).resolve()
            if resolvido_rel.exists():
                try:
                    return str(resolvido_rel.relative_to(base_dir)).replace("\\", "/")
                except ValueError:
                    return str(resolvido_rel)
        except OSError:
            pass
        return caminho_txt.replace("\\", "/")

    def _migrar_pastas_documentos_legadas(self):
        base_dir = Path(__file__).resolve().parent
        destino_padrao = base_dir / "projetta"
        destino_padrao.mkdir(parents=True, exist_ok=True)

        fontes = [
            base_dir / "certificados",
            base_dir / "outros",
            base_dir / "documentos",
            (base_dir / "documentos" / "certificados"),
            (base_dir / "documentos" / "outros"),
        ]
        for origem in fontes:
            if not origem.exists() or not origem.is_dir():
                continue
            if origem.resolve() == destino_padrao.resolve():
                continue
            for arquivo in origem.iterdir():
                if not arquivo.is_file():
                    continue
                alvo = destino_padrao / arquivo.name
                if alvo.exists():
                    i = 1
                    while True:
                        candidato = destino_padrao / f"{arquivo.stem}_{i}{arquivo.suffix}"
                        if not candidato.exists():
                            alvo = candidato
                            break
                        i += 1
                try:
                    shutil.move(str(arquivo), str(alvo))
                except OSError:
                    continue
        try:
            for pasta in (
                base_dir / "documentos" / "certificados",
                base_dir / "documentos" / "outros",
                base_dir / "certificados",
                base_dir / "outros",
                base_dir / "documentos",
            ):
                if pasta.exists() and pasta.is_dir() and not any(pasta.iterdir()):
                    pasta.rmdir()
        except OSError:
            return

    def _avisar_falha_git_auto(self, detalhe):
        if self._aviso_git_auto_exibido:
            return
        self._aviso_git_auto_exibido = True
        detalhe_txt = str(detalhe or "").strip() or "erro nao identificado"
        messagebox.showwarning(
            "Git Auto",
            "Nao foi possivel executar commit + push automatico.\n"
            f"Detalhe: {detalhe_txt}",
        )

    def _normalizar_caminho_git_relativo(self, caminho):
        base = Path(__file__).resolve().parent.resolve()
        p = Path(str(caminho or "")).expanduser()
        if not p.is_absolute():
            p = (base / p).resolve()
        try:
            rel = p.resolve().relative_to(base)
            return str(rel).replace("\\", "/")
        except Exception:
            return ""

    def _enfileirar_git_auto_commit(self, caminhos):
        if not self._git_auto_commit_habilitado:
            return
        itens = caminhos if isinstance(caminhos, (list, tuple, set)) else [caminhos]
        normalizados = set()
        for item in itens:
            rel = self._normalizar_caminho_git_relativo(item)
            if rel:
                normalizados.add(rel)
        if not normalizados:
            return
        with self._git_auto_commit_lock:
            self._git_auto_commit_pendentes.update(normalizados)
            thread_ativa = self._git_auto_commit_thread is not None and self._git_auto_commit_thread.is_alive()
            if thread_ativa:
                return
            self._git_auto_commit_thread = threading.Thread(
                target=self._worker_git_auto_commit,
                daemon=True,
            )
            self._git_auto_commit_thread.start()

    def _worker_git_auto_commit(self):
        repo_dir = Path(__file__).resolve().parent
        while True:
            with self._git_auto_commit_lock:
                if not self._git_auto_commit_pendentes:
                    self._git_auto_commit_thread = None
                    return
                lote = sorted(self._git_auto_commit_pendentes)
                self._git_auto_commit_pendentes.clear()

            try:
                proc_repo = subprocess.run(
                    ["git", "-C", str(repo_dir), "rev-parse", "--is-inside-work-tree"],
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=12,
                )
                if proc_repo.returncode != 0 or str(proc_repo.stdout or "").strip().lower() != "true":
                    continue

                add_cmd = ["git", "-C", str(repo_dir), "add", "--", *lote]
                proc_add = subprocess.run(
                    add_cmd,
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                if proc_add.returncode != 0:
                    raise RuntimeError((proc_add.stderr or proc_add.stdout or "git add falhou").strip())

                diff_cmd = ["git", "-C", str(repo_dir), "diff", "--cached", "--name-only", "--", *lote]
                proc_diff = subprocess.run(
                    diff_cmd,
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=20,
                )
                if proc_diff.returncode != 0:
                    raise RuntimeError((proc_diff.stderr or proc_diff.stdout or "git diff falhou").strip())
                alterados = [ln.strip() for ln in str(proc_diff.stdout or "").splitlines() if ln.strip()]
                if not alterados:
                    continue

                foco = alterados[0]
                msg = f"Auto save docs: {foco}"
                proc_commit = subprocess.run(
                    ["git", "-C", str(repo_dir), "commit", "-m", msg, "--", *alterados],
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=40,
                )
                if proc_commit.returncode != 0:
                    saida = (proc_commit.stderr or proc_commit.stdout or "").strip()
                    if "nothing to commit" in saida.lower():
                        continue
                    raise RuntimeError(saida or "git commit falhou")

                proc_push = subprocess.run(
                    ["git", "-C", str(repo_dir), "push"],
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
                if proc_push.returncode != 0:
                    raise RuntimeError((proc_push.stderr or proc_push.stdout or "git push falhou").strip())

            except Exception as exc:
                try:
                    self.after(0, lambda e=exc: self._avisar_falha_git_auto(e))
                except Exception:
                    pass
                continue

    def _registrar_documento_salvo(
        self,
        caminho,
        origem="sistema",
        empresa_id=None,
        funcionario_id=None,
        tipo_documento="",
        inserir_qr_pdf=True,
    ):
        caminho_norm = self._normalizar_caminho_documento_db(caminho)
        if not caminho_norm:
            return
        try:
            base_repo = Path(__file__).resolve().parent.resolve()
            caminho_abs = Path(str(caminho_norm)).expanduser()
            if not caminho_abs.is_absolute():
                caminho_abs = (base_repo / caminho_abs).resolve()
            dentro_repo = True
            try:
                caminho_abs.relative_to(base_repo)
            except Exception:
                dentro_repo = False
            if (not dentro_repo) and caminho_abs.suffix.lower() == ".pdf":
                espelhado = self._espelhar_arquivo_no_repo_local(caminho_abs)
                if espelhado:
                    caminho_norm = espelhado
        except Exception:
            pass
        caminhos_git = [caminho_norm]
        item = {
            "caminho": caminho_norm,
            "origem": str(origem or "sistema"),
            "empresa_id": empresa_id if isinstance(empresa_id, int) else None,
            "funcionario_id": funcionario_id if isinstance(funcionario_id, int) else None,
            "tipo_documento": str(tipo_documento or "").strip(),
            "data_ultima_gravacao": self._iso_datetime_now(),
        }
        qr_caminho = self._gerar_qrcode_documento_salvo(caminho_norm, item["tipo_documento"])
        if qr_caminho:
            item["qrcode"] = qr_caminho
            caminhos_git.append(qr_caminho)
            caminho_pdf = Path(str(caminho_norm or "")).expanduser()
            if not caminho_pdf.is_absolute():
                caminho_pdf = (Path(__file__).resolve().parent / caminho_pdf).resolve()
            if inserir_qr_pdf and caminho_pdf.suffix.lower() == ".pdf":
                tipo_norm = str(item.get("tipo_documento", "") or "").strip().casefold()
                origem_norm = str(item.get("origem", "") or "").strip().casefold()
                if tipo_norm in {"fit test", "fittest"}:
                    self._inserir_qrcode_no_pdf(
                        caminho_pdf,
                        qr_caminho,
                        tamanho_cm=3.0,
                        margem_esquerda_cm=2.0,
                        margem_inferior_cm=2.0,
                    )
                elif origem_norm == "imprimir_nr" or tipo_norm.startswith("nr"):
                    self._inserir_qrcode_no_pdf(
                        caminho_pdf,
                        qr_caminho,
                        tamanho_cm=2.0,
                        margem_esquerda_cm=2.0,
                        margem_inferior_cm=0.5,
                    )
                elif tipo_norm in {"ordem de servico", "ordem de serviço", "anuencia", "anuência"}:
                    self._inserir_qrcode_no_pdf(
                        caminho_pdf,
                        qr_caminho,
                        tamanho_cm=3.0,
                        margem_esquerda_cm=2.0,
                        margem_inferior_cm=1.5,
                        somente_ultima_pagina=True,
                    )
                else:
                    self._inserir_qrcode_no_pdf(
                        caminho_pdf,
                        qr_caminho,
                        tamanho_cm=2.0,
                    )
                self._limpar_metadados_pdf(caminho_pdf)
        caminho_pdf = Path(str(caminho_norm or "")).expanduser()
        if not caminho_pdf.is_absolute():
            caminho_pdf = (Path(__file__).resolve().parent / caminho_pdf).resolve()
        if caminho_pdf.suffix.lower() == ".pdf":
            if self._assinatura_digital_ativa():
                ok_assin, msg_assin = self._assinar_pdf_por_marcadores(caminho_pdf)
                item["assinatura_digital"] = "assinado" if ok_assin else "falha"
                if not ok_assin:
                    item["assinatura_erro"] = str(msg_assin or "").strip()
                    if not self._aviso_assinatura_exibido:
                        self._aviso_assinatura_exibido = True
                        messagebox.showwarning(
                            "Assinatura Digital",
                            "Nao foi possivel assinar digitalmente um ou mais PDFs.\n"
                            f"Detalhe: {msg_assin}",
                        )
            else:
                item["assinatura_digital"] = "nao_configurada"
            self._sincronizar_pdf_na_area_de_trabalho(caminho_pdf, item.get("qrcode", ""))
            self._sincronizar_pdf_no_github(caminho_pdf, item.get("qrcode", ""))
        for idx, existente in enumerate(self.documentos_salvos):
            if str(existente.get("caminho", "") or "").strip() == caminho_norm:
                self.documentos_salvos[idx] = item
                self._enfileirar_git_auto_commit(caminhos_git)
                return
        self.documentos_salvos.append(item)
        self._enfileirar_git_auto_commit(caminhos_git)

    def _avisar_falha_git_sync(self, detalhe):
        if self._aviso_git_sync_exibido:
            return
        self._aviso_git_sync_exibido = True
        detalhe_txt = str(detalhe or "").strip() or "erro nao identificado"
        messagebox.showwarning(
            "Publicacao",
            "PDF gerado localmente, mas nao foi possivel publicar no site automaticamente.\n"
            f"Detalhe: {detalhe_txt}",
        )

    def _sincronizar_documentos_salvos_pendentes(self):
        if not self._deve_sincronizar_qr_pendentes():
            return
        vistos = set()
        for item in list(self.documentos_salvos):
            if not isinstance(item, dict):
                continue
            caminho_txt = self._normalizar_caminho_documento_db(item.get("caminho", ""))
            if not caminho_txt:
                continue
            caminho_pdf = Path(str(caminho_txt)).expanduser()
            if not caminho_pdf.is_absolute():
                caminho_pdf = (Path(__file__).resolve().parent / caminho_pdf).resolve()
            if caminho_pdf.suffix.lower() != ".pdf" or not caminho_pdf.exists() or not caminho_pdf.is_file():
                continue
            chave = str(caminho_pdf.resolve())
            if chave in vistos:
                continue
            vistos.add(chave)
            caminho_qr = str(item.get("qrcode", "") or "").strip()
            self._sincronizar_pdf_no_github(caminho_pdf, caminho_qr, avisar_falha=False)

    def _avisar_falha_desktop_sync(self, detalhe):
        if self._aviso_desktop_sync_exibido:
            return
        self._aviso_desktop_sync_exibido = True
        detalhe_txt = str(detalhe or "").strip() or "erro nao identificado"
        messagebox.showwarning(
            "Area de Trabalho",
            "PDF gerado localmente, mas nao foi possivel salvar a copia na Area de Trabalho.\n"
            f"Detalhe: {detalhe_txt}",
        )

    def _obter_nome_funcionario_para_desktop(self, caminho_pdf):
        caminho = Path(str(caminho_pdf or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()

        candidato = ""
        repo_dir = Path(__file__).resolve().parent
        try:
            rel = caminho.resolve().relative_to(repo_dir.resolve())
            partes = list(rel.parts)
            if len(partes) >= 2 and str(partes[0]).casefold() == "_pdf_gerados":
                candidato = str(partes[1] or "").strip()
        except Exception:
            pass

        if not candidato:
            candidato = str(caminho.parent.name or "").strip()
        return self._obter_nome_arquivo_seguro(candidato, "funcionario")

    def _copiar_arquivo_para_desktop(self, caminho_arquivo, pasta_funcionario="", subpasta=""):
        arquivo = Path(str(caminho_arquivo or "")).expanduser()
        if not arquivo.is_absolute():
            arquivo = (Path(__file__).resolve().parent / arquivo).resolve()
        if not arquivo.exists() or not arquivo.is_file():
            return

        desktop_base = self._obter_desktop_base()
        nome_pasta = self._obter_nome_arquivo_seguro(pasta_funcionario, "funcionario")
        pasta_desktop = desktop_base / nome_pasta
        subpasta_txt = str(subpasta or "").strip().replace("\\", "/").strip("/")
        if subpasta_txt:
            pasta_desktop = pasta_desktop / Path(subpasta_txt)
        pasta_desktop.mkdir(parents=True, exist_ok=True)

        destino = pasta_desktop / arquivo.name
        destino.parent.mkdir(parents=True, exist_ok=True)
        ultima_exc = None
        # Alguns arquivos PDF ficam bloqueados por poucos instantes apos geracao/conversao.
        for _ in range(8):
            try:
                shutil.copy2(arquivo, destino)
                return
            except PermissionError as exc:
                ultima_exc = exc
            except OSError as exc:
                ultima_exc = exc
                if getattr(exc, "winerror", None) != 32:
                    raise
            time.sleep(0.25)
        if ultima_exc is not None:
            raise ultima_exc

    def _sincronizar_pdf_na_area_de_trabalho(self, caminho_pdf, caminho_qr=""):
        try:
            caminho = Path(str(caminho_pdf or "")).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            if caminho.suffix.lower() != ".pdf" or not caminho.exists() or not caminho.is_file():
                return
            nome_funcionario = self._obter_nome_funcionario_para_desktop(caminho)
            self._copiar_arquivo_para_desktop(caminho, pasta_funcionario=nome_funcionario)

            qr_txt = str(caminho_qr or "").strip()
            if qr_txt:
                self._copiar_arquivo_para_desktop(
                    qr_txt,
                    pasta_funcionario=nome_funcionario,
                    subpasta="_qrcodes",
                )
        except Exception as exc:
            self._avisar_falha_desktop_sync(exc)

    def _sincronizar_pdf_no_github(self, caminho_pdf, caminho_qr="", avisar_falha=True):
        try:
            caminho = Path(str(caminho_pdf or "")).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            if caminho.suffix.lower() != ".pdf" or not caminho.exists() or not caminho.is_file():
                return
            # Publica o PDF no repositorio remoto do site via API.
            url_publicada = self._publicar_arquivo_no_site(caminho)
            if url_publicada:
                self._atualizar_indice_documento_site(caminho)
                qr_txt = str(caminho_qr or "").strip()
                if qr_txt:
                    caminho_qr_arq = Path(qr_txt).expanduser()
                    if not caminho_qr_arq.is_absolute():
                        caminho_qr_arq = (Path(__file__).resolve().parent / caminho_qr_arq).resolve()
                    if caminho_qr_arq.exists() and caminho_qr_arq.is_file():
                        self._publicar_arquivo_no_site(caminho_qr_arq)
            if not url_publicada:
                # Fallback: espelha no repositorio local e deixa o Git Auto publicar.
                caminhos_fallback = []
                rel_pdf = self._espelhar_arquivo_no_repo_local(caminho)
                if rel_pdf:
                    caminhos_fallback.append(rel_pdf)
                qr_txt = str(caminho_qr or "").strip()
                if qr_txt:
                    rel_qr = self._espelhar_arquivo_no_repo_local(qr_txt)
                    if rel_qr:
                        caminhos_fallback.append(rel_qr)
                if caminhos_fallback:
                    self._enfileirar_git_auto_commit(caminhos_fallback)
                    self._qr_github_ultimo_erro = ""
                    return
                detalhe = str(self._qr_github_ultimo_erro or "").strip() or "falha ao publicar no site"
                if avisar_falha:
                    self._avisar_falha_git_sync(detalhe)
        except Exception as exc:
            if avisar_falha:
                self._avisar_falha_git_sync(exc)
            return

    def _montar_payload_qrcode_documento(self, caminho_documento, permitir_arquivo_inexistente=False):
        caminho_txt = str(caminho_documento or "").strip()
        if not caminho_txt:
            return ""
        caminho = Path(caminho_txt).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        # Padrao oficial de destino do QR: rota amigavel no site
        # /nomedofuncionario/documento/data
        url_consulta = self._url_site_consulta_para_arquivo(caminho)
        url_github = self._url_github_qr_para_arquivo(
            caminho,
            permitir_inexistente=bool(permitir_arquivo_inexistente),
        )
        if url_github:
            # Mesmo com publicacao concluida, mantem o payload no formato de rota amigavel.
            if url_consulta:
                return url_consulta
            return url_github
        if url_consulta:
            return url_consulta
        url_pages = self._url_github_pages_para_arquivo(caminho)
        if url_pages:
            return url_pages
        url_raw = self._url_github_raw_para_arquivo(caminho)
        if url_raw:
            return url_raw
        url_blob = self._url_github_blob_para_arquivo(caminho)
        if url_blob:
            return url_blob
        base_url = str(os.environ.get("CADNR_QR_BASE_URL", "") or "").strip()
        if base_url:
            base_url = re.sub(r"/+$", "", base_url)
            try:
                rel = caminho.resolve().relative_to(Path(__file__).resolve().parent.resolve())
                rel_txt = str(rel).replace("\\", "/")
            except Exception:
                rel_txt = caminho.name
            return f"{base_url}/{parse.quote(rel_txt, safe='/')}"
        url_local = self._url_local_qr_para_arquivo(
            caminho,
            permitir_inexistente=bool(permitir_arquivo_inexistente),
        )
        if url_local:
            return url_local
        try:
            return caminho.resolve().as_uri()
        except Exception:
            return str(caminho.resolve())

    def _gerar_qrcode_documento_salvo(self, caminho_documento, tipo_documento=""):
        caminho_txt = str(caminho_documento or "").strip()
        if not caminho_txt:
            return ""
        try:
            caminho = Path(caminho_txt).expanduser()
            if not caminho.is_absolute():
                caminho = (Path(__file__).resolve().parent / caminho).resolve()
            if not caminho.exists() or not caminho.is_file():
                return ""
            if caminho.suffix.lower() in {".png", ".jpg", ".jpeg"} and caminho.stem.lower().endswith("_qrcode"):
                return ""

            pasta_qr = caminho.parent / "_qrcodes"
            pasta_qr.mkdir(parents=True, exist_ok=True)
            destino_qr = pasta_qr / f"{caminho.stem}_qrcode.png"
            payload = self._montar_payload_qrcode_documento(caminho)

            import qrcode

            qr = qrcode.QRCode(
                version=None,
                error_correction=qrcode.constants.ERROR_CORRECT_M,
                box_size=10,
                border=4,
            )
            qr.add_data(payload)
            qr.make(fit=True)
            img = qr.make_image()
            img.save(str(destino_qr))
            return self._normalizar_caminho_documento_db(destino_qr)
        except Exception:
            return ""

    def _inserir_qrcode_em_carteirinha_docx(
        self,
        caminho_docx,
        caminho_qr,
        tamanho_cm=1.8,
        margem_esquerda_cm=0.1,
    ):
        inseriu_marcador, _ = self._inserir_imagem_por_marcador_docx(
            caminho_docx,
            caminho_qr,
            ["qrcode1", "qrcode", "qr code", "qr_code", "qr"],
            largura_cm=float(tamanho_cm),
            altura_cm=float(tamanho_cm),
            incluir_cabecalho_rodape=False,
        )
        if inseriu_marcador:
            return True
        try:
            from docx import Document
            from docx.shared import Cm
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except Exception:
            return False

        try:
            doc = Document(str(caminho_docx))
            tabelas_alvo = []
            for tabela in doc.tables:
                texto_tabela = " ".join(cell.text or "" for row in tabela.rows for cell in row.cells)
                texto_norm = self._normalizar_texto_filtro(texto_tabela)
                score = 0
                if "operador" in texto_norm:
                    score += 5
                if "nr" in texto_norm:
                    score += 2
                if score > 0:
                    tabelas_alvo.append((score, tabela))
            if tabelas_alvo:
                tabelas_alvo = [t for _, t in sorted(tabelas_alvo, key=lambda x: x[0], reverse=True)]
            else:
                tabelas_alvo = list(doc.tables)

            for tabela in tabelas_alvo:
                if not tabela.rows:
                    continue
                # Posiciona no canto inferior esquerdo da tabela alvo.
                cell = tabela.rows[-1].cells[0] if tabela.rows[-1].cells else None
                if cell is None:
                    continue
                try:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
                except Exception:
                    pass
                p = cell.paragraphs[-1] if cell.paragraphs else cell.add_paragraph()
                try:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                except Exception:
                    pass
                try:
                    p.paragraph_format.left_indent = Cm(float(margem_esquerda_cm))
                    p.paragraph_format.space_before = Cm(0)
                    p.paragraph_format.space_after = Cm(0)
                except Exception:
                    pass
                run = p.add_run()
                run.add_picture(
                    str(caminho_qr),
                    width=Cm(float(tamanho_cm)),
                    height=Cm(float(tamanho_cm)),
                )
                doc.save(str(caminho_docx))
                return True
        except Exception:
            return False
        return False

    def _inserir_qrcode_no_pdf(
        self,
        caminho_pdf,
        caminho_qr,
        tamanho_cm=2.0,
        margem_esquerda_cm=2.0,
        margem_inferior_cm=1.2,
        somente_ultima_pagina=False,
    ):
        try:
            import fitz
        except Exception:
            return False, "PyMuPDF (fitz) indisponivel"

        try:
            pdf = Path(str(caminho_pdf or "")).expanduser()
            qr = Path(str(caminho_qr or "")).expanduser()
            if not pdf.is_absolute():
                pdf = (Path(__file__).resolve().parent / pdf).resolve()
            if not qr.is_absolute():
                qr = (Path(__file__).resolve().parent / qr).resolve()
            if not pdf.exists() or not qr.exists():
                return False, "arquivo PDF ou QR nao encontrado"

            doc = fitz.open(str(pdf))
            if doc.page_count <= 0:
                doc.close()
                return False, "pdf sem paginas"

            pt_por_cm = 28.3464567
            tam = float(tamanho_cm) * pt_por_cm
            margem_esquerda = float(margem_esquerda_cm) * pt_por_cm
            margem_inferior = float(margem_inferior_cm) * pt_por_cm
            paginas_alvo = [doc[-1]] if bool(somente_ultima_pagina) else list(doc)
            for pagina in paginas_alvo:
                x0 = max(0.0, margem_esquerda)
                y1 = pagina.rect.height - margem_inferior
                y0 = max(0.0, y1 - tam)
                x1 = x0 + tam
                rect = fitz.Rect(x0, y0, x1, y1)
                pagina.insert_image(rect, filename=str(qr), keep_proportion=True, overlay=True)

            temp_saida = pdf.with_name(f"{pdf.stem}.__qr__.pdf")
            doc.save(str(temp_saida), garbage=4, deflate=True)
            doc.close()
            shutil.move(str(temp_saida), str(pdf))
            return True, ""
        except Exception as ex:
            try:
                temp_saida = Path(str(caminho_pdf or "")).with_name(
                    f"{Path(str(caminho_pdf or '')).stem}.__qr__.pdf"
                )
                if temp_saida.exists():
                    temp_saida.unlink()
            except Exception:
                pass
            return False, str(ex)

    @staticmethod
    def _inserir_imagem_por_marcador_docx(
        docx_path,
        imagem_path,
        marcadores,
        largura_cm=2,
        altura_cm=None,
        incluir_cabecalho_rodape=True,
    ):
        caminho_docx = Path(docx_path)
        caminho_img = Path(imagem_path)
        if not caminho_docx.exists() or not caminho_img.exists():
            return False, "Arquivo DOCX ou imagem nao encontrado."
        try:
            from docx import Document
            from docx.shared import Cm
        except Exception:
            return False, "Dependencia ausente: python-docx (pip install python-docx)."

        lista_marcadores = [str(m).strip() for m in (marcadores or []) if str(m).strip()]
        if not lista_marcadores:
            return False, "Marcador de imagem nao informado."
        termos = []
        for marcador in lista_marcadores:
            letras = r"\s*".join(re.escape(ch) for ch in marcador)
            termos.append(rf"\[\s*{letras}\s*\]")
            termos.append(rf"(?<![A-Za-zÀ-ÿ0-9_]){letras}(?![A-Za-zÀ-ÿ0-9_])")
        padrao = re.compile("|".join(termos), flags=re.IGNORECASE)

        def _limpar_marcador_runs(paragrafo):
            textos = [run.text or "" for run in paragrafo.runs]
            if not textos:
                return False
            completo = "".join(textos)
            novo = padrao.sub("", completo)
            if novo == completo:
                return False
            if paragrafo.runs:
                paragrafo.runs[0].text = novo
                for run in paragrafo.runs[1:]:
                    run.text = ""
            return True

        def _iter_paragrafos_tabelas(tabela):
            for row in tabela.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p
                    for tabela_interna in cell.tables:
                        yield from _iter_paragrafos_tabelas(tabela_interna)

        def _kwargs_tamanho_imagem():
            largura_max = float(largura_cm) if largura_cm is not None else None
            altura_max = float(altura_cm) if altura_cm is not None else None
            if largura_max is not None and largura_max <= 0:
                largura_max = None
            if altura_max is not None and altura_max <= 0:
                altura_max = None

            if largura_max is None and altura_max is None:
                return {}

            # Fallback sem PIL: aplica pelo menos a largura maxima.
            if Image is None:
                if largura_max is not None:
                    return {"width": Cm(largura_max)}
                return {"height": Cm(altura_max)}

            try:
                with Image.open(caminho_img) as img:
                    largura_px, altura_px = img.size
            except Exception:
                if largura_max is not None:
                    return {"width": Cm(largura_max)}
                return {"height": Cm(altura_max)}

            if largura_px <= 0 or altura_px <= 0:
                if largura_max is not None:
                    return {"width": Cm(largura_max)}
                return {"height": Cm(altura_max)}

            proporcao = float(largura_px) / float(altura_px)
            if largura_max is not None and altura_max is not None:
                proporcao_caixa = float(largura_max) / float(altura_max)
                if proporcao >= proporcao_caixa:
                    largura_final = largura_max
                    altura_final = largura_final / proporcao
                else:
                    altura_final = altura_max
                    largura_final = altura_final * proporcao
                return {"width": Cm(largura_final), "height": Cm(altura_final)}

            if largura_max is not None:
                return {"width": Cm(largura_max)}
            return {"height": Cm(altura_max)}

        try:
            doc = Document(str(caminho_docx))
            alterou = False
            paragrafos_alvo = list(doc.paragraphs)
            for tabela in doc.tables:
                paragrafos_alvo.extend(list(_iter_paragrafos_tabelas(tabela)))
            if incluir_cabecalho_rodape:
                for sec in doc.sections:
                    paragrafos_alvo.extend(list(sec.header.paragraphs))
                    paragrafos_alvo.extend(list(sec.footer.paragraphs))
                    for tabela in sec.header.tables:
                        paragrafos_alvo.extend(list(_iter_paragrafos_tabelas(tabela)))
                    for tabela in sec.footer.tables:
                        paragrafos_alvo.extend(list(_iter_paragrafos_tabelas(tabela)))

            tamanho_kwargs = _kwargs_tamanho_imagem()
            for p in paragrafos_alvo:
                if _limpar_marcador_runs(p):
                    run_img = p.add_run()
                    try:
                        run_img.add_picture(str(caminho_img), **tamanho_kwargs)
                        alterou = True
                    except Exception:
                        pass
            if alterou:
                doc.save(str(caminho_docx))
                return True, ""
            return False, "Marcador de imagem nao encontrado no documento."
        except Exception:
            return False, "Falha ao processar o DOCX para inserir a imagem."

    @staticmethod
    def _inserir_img1_logo_em_tabela_docx(docx_path, logo_path, incluir_cabecalho_rodape=True):
        return App._inserir_imagem_por_marcador_docx(
            docx_path,
            logo_path,
            ["logo1", "img1", "img"],
            largura_cm=2.0,
            incluir_cabecalho_rodape=incluir_cabecalho_rodape,
        )

    @staticmethod
    def _inserir_img2_logo_em_tabela_docx(docx_path, logo_path, incluir_cabecalho_rodape=True):
        return App._inserir_imagem_por_marcador_docx(
            docx_path,
            logo_path,
            ["logo2", "img2"],
            largura_cm=2.0,
            incluir_cabecalho_rodape=incluir_cabecalho_rodape,
        )

    @staticmethod
    def _inserir_img3_logo_em_tabela_docx(docx_path, logo_path, incluir_cabecalho_rodape=True):
        return App._inserir_imagem_por_marcador_docx(
            docx_path,
            logo_path,
            ["logo3", "img3"],
            largura_cm=1.3,
            incluir_cabecalho_rodape=incluir_cabecalho_rodape,
        )

    @staticmethod
    def _carregar_preview_logo(caminho_logo, largura_max=180, altura_max=110):
        caminho = Path(str(caminho_logo or "")).expanduser()
        if not caminho.exists() or not caminho.is_file():
            return None

        if Image is not None and ImageTk is not None:
            try:
                with Image.open(caminho) as img:
                    imagem = img.copy()
                imagem.thumbnail((largura_max, altura_max))
                return ImageTk.PhotoImage(imagem)
            except OSError:
                pass

        try:
            imagem_tk = tk.PhotoImage(file=str(caminho))
            fator_x = max(1, (imagem_tk.width() + largura_max - 1) // largura_max)
            fator_y = max(1, (imagem_tk.height() + altura_max - 1) // altura_max)
            fator = max(fator_x, fator_y)
            return imagem_tk.subsample(fator, fator) if fator > 1 else imagem_tk
        except tk.TclError:
            return None

    def _atualizar_preview_logo(self, label_widget, caminho_logo):
        img = self._carregar_preview_logo(caminho_logo)
        if img is None:
            label_widget.configure(image="", text="Pre-visualizacao indisponivel")
            label_widget.image = None
            return
        label_widget.configure(image=img, text="")
        label_widget.image = img

    def _origem_word_por_nr_selecionada(self, funcionario):
        base_dir = Path(__file__).resolve().parent
        empresa = self._empresa_do_funcionario(funcionario) or {}
        empresa_id = empresa.get("id")
        pasta_empresa = None
        nome_referencia_pasta = ""
        for candidato in self._pastas_candidatas_empresa(empresa):
            if not candidato:
                continue
            pasta_candidata = base_dir / candidato
            if pasta_candidata.exists() and pasta_candidata.is_dir():
                pasta_empresa = pasta_candidata
                nome_referencia_pasta = candidato
                break

        selecionadas = []
        for item in self.nr_certificados:
            if bool(item.get("imprimir_adicionado", False)):
                selecionadas.append(str(item.get("nome", "") or "").strip())
        if not selecionadas:
            for item in self.nr_certificados:
                if bool(item.get("imprimir", False)):
                    selecionadas.append(str(item.get("nome", "") or "").strip())
        nomes_empresa = self._nomes_nr_na_pasta_empresa(empresa_id) or {}
        chaves_empresa = set(nomes_empresa.keys())
        if chaves_empresa:
            selecionadas = [
                n for n in selecionadas if self._normalizar_nome_nr(n) in chaves_empresa
            ]

        arquivos_docx = []
        if pasta_empresa and pasta_empresa.exists() and pasta_empresa.is_dir():
            for arq in sorted(pasta_empresa.glob("*.docx")):
                if self._arquivo_nr_vinculado_empresa(arq, empresa, nome_referencia_pasta):
                    arquivos_docx.append(arq)

        def _numero_nr(texto):
            m = re.search(r"\bnr\s*0*(\d+)\b", str(texto or ""), flags=re.IGNORECASE)
            return m.group(1) if m else ""

        if selecionadas and arquivos_docx:
            for nome_nr in selecionadas:
                chave = self._normalizar_nome_nr(nome_nr)
                nr_numero = _numero_nr(nome_nr)
                candidatos_mesma_nr = []
                for arq in arquivos_docx:
                    nome_limpo_arquivo = self._nome_nr_do_arquivo(arq.stem, nome_referencia_pasta)
                    chave_arquivo = self._normalizar_nome_nr(nome_limpo_arquivo)
                    if nr_numero and _numero_nr(nome_limpo_arquivo) == nr_numero:
                        candidatos_mesma_nr.append((chave_arquivo, arq))
                for chave_arquivo, arq in candidatos_mesma_nr:
                    if chave_arquivo == chave:
                        return arq
                if candidatos_mesma_nr:
                    return candidatos_mesma_nr[0][1]
                for arq in arquivos_docx:
                    nome_limpo_arquivo = self._nome_nr_do_arquivo(arq.stem, nome_referencia_pasta)
                    chave_arquivo = self._normalizar_nome_nr(nome_limpo_arquivo)
                    # Aceita igualdade e casos de nome composto no arquivo.
                    if (
                        chave_arquivo == chave
                        or chave in chave_arquivo
                        or chave_arquivo in chave
                    ):
                        return arq

        for fallback in ("nr6.docx", "nr_certificados.docx", "nr_documentos.docx", "nr_projetta.docx"):
            caminho = base_dir / fallback
            if caminho.exists():
                return caminho

        if pasta_empresa and pasta_empresa.exists() and pasta_empresa.is_dir():
            primeiro = next(iter(sorted(pasta_empresa.glob("*.docx"))), None)
            if primeiro is not None:
                return primeiro
        return None

    def _atualizar_word_nr(self, pasta_destino=None, exibir_mensagem=True):
        funcionario = self._funcionario_ativo_para_documento()
        if funcionario is None:
            messagebox.showwarning(
                "NR",
                "Selecione um funcionario em CADNR ou CADASTROS para atualizar o Word.",
            )
            return None

        origem = self._origem_word_por_nr_selecionada(funcionario)
        if origem is None or not origem.exists():
            messagebox.showerror(
                "NR",
                "Nenhum arquivo .docx correspondente foi encontrado para a NR selecionada.",
            )
            return None

        campos = self._montar_campos_documento(funcionario)
        foto1_caminho = self._resolver_foto_funcionario(campos.get("FOTO1", ""))
        logo_img1_caminho = self._resolver_logo_empresa(campos.get("LOGO1", ""))
        logo_img2_caminho = self._resolver_logo_empresa(campos.get("LOGO2", ""))
        logo_img3_caminho = self._resolver_logo_empresa(campos.get("LOGO3", ""))

        nome_limpo = re.sub(
            r"[^A-Za-z0-9]+",
            "_",
            str(funcionario.get("nome", "funcionario")),
        ).strip("_")
        if not nome_limpo:
            nome_limpo = "funcionario"
        if pasta_destino is None:
            destino = Path(__file__).with_name(f"nr6_{nome_limpo}.docx")
        else:
            destino = pasta_destino / f"nr6_{nome_limpo}.docx"

        try:
            with zipfile.ZipFile(origem, "r") as zin, zipfile.ZipFile(destino, "w") as zout:
                for info in zin.infolist():
                    data = zin.read(info.filename)
                    if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                        texto = data.decode("utf-8")

                        def repl(m):
                            marcador = m.group(2).strip()
                            chave = marcador.upper()
                            if chave in {"EMPRESA", "CNPJ", "FOTO1", "LOGO1", "LOGO2", "LOGO3", "IMG1", "IMG2", "IMG3"}:
                                return m.group(0)
                            if chave not in campos:
                                return m.group(0)
                            valor = self._formatar_por_placeholder(campos[chave], marcador)
                            return self._escape_xml_text(valor)

                        pattern = r"([‘'])(\s*[A-Za-zÀ-ÿ0-9_]+\s*)([’'])"
                        texto = re.sub(pattern, repl, texto)

                        de_e_ate_token = "__CADNR_DE_E_ATE__"
                        de_a_ate_token = "__CADNR_DE_A_ATE__"

                        alias_campos = {}
                        for k, v in campos.items():
                            chave_alias = self._normalizar_chave_placeholder(k)
                            if chave_alias in {"EMPRESA", "CNPJ", "FOTO1", "LOGO1", "LOGO2", "LOGO3", "IMG1", "IMG2", "IMG3"}:
                                continue
                            alias_campos[chave_alias] = v
                        alias_campos["NOME_DO_FUNCIONARIO"] = campos.get("NOME", "")
                        alias_campos["DD_DE_MES_DE_AAAA"] = campos.get("DATA_EXTENSO", "")
                        alias_campos["FIM1"] = campos.get("DATA_ATE_EXTENSO", "")
                        alias_campos["COMECO1_E_FIM1"] = campos.get("DE_E_ATE_EXTENSO", "")
                        alias_campos["COMECO1_A_FIM1"] = campos.get("DE_A_ATE_EXTENSO", "")

                        def repl_colchete(m):
                            marcador_raw = m.group(1)
                            marcador = re.sub(r"<[^>]+>", "", marcador_raw).strip()
                            if not marcador:
                                return m.group(0)
                            chave_norm = self._normalizar_chave_placeholder(marcador)
                            if chave_norm in alias_campos:
                                if chave_norm == "COMECO1_E_FIM1":
                                    return de_e_ate_token
                                if chave_norm == "COMECO1_A_FIM1":
                                    return de_a_ate_token
                                valor = self._formatar_por_placeholder(alias_campos[chave_norm], marcador)
                                return self._escape_xml_text(valor)
                            # Fallback para texto de data literal dentro de colchetes.
                            if re.search(r"\d", marcador) and "de" in marcador.lower():
                                valor = str(campos.get("DATA_EXTENSO", "") or "")
                                if valor:
                                    # Mantem prefixo original (ex.: cidade) e troca somente a data.
                                    marcador_atualizado = re.sub(
                                        r"(\d{1,2}\s+de\s+[A-Za-zÀ-ÿ]+\s+de\s+\d{4}\.?)",
                                        valor,
                                        marcador,
                                        count=1,
                                        flags=re.IGNORECASE,
                                    )
                                    return self._escape_xml_text(marcador_atualizado)
                            return m.group(0)

                        texto = re.sub(r"\[\s*(.*?)\s*\]", repl_colchete, texto, flags=re.DOTALL)

                        # Reforco para placeholders conhecidos que podem escapar da regex
                        # (por exemplo, quando o Word quebra runs de forma incomum).
                        substituicoes_diretas = [
                            ("NOME DO FUNCIONARIO", campos.get("NOME", "")),
                            ("NOME DO FUNCIONÁRIO", campos.get("NOME", "")),
                            ("COMEÇO1 E FIM1", de_e_ate_token),
                            ("COMECO1 E FIM1", de_e_ate_token),
                            ("COMEÇO1 A FIM1", de_a_ate_token),
                            ("COMECO1 A FIM1", de_a_ate_token),
                            ("FIM1", campos.get("FIM1", "")),
                            ("COMEÇO1", campos.get("COMECO1", "")),
                            ("COMECO1", campos.get("COMECO1", "")),
                            ("DATANR", campos.get("DATANR", "")),
                            ("DATA_NR", campos.get("DATANR", "")),
                        ]
                        for marcador_txt, valor_txt in substituicoes_diretas:
                            valor_xml = self._escape_xml_text(valor_txt)
                            texto = texto.replace(f"[{marcador_txt}]", valor_xml)
                            texto = texto.replace(marcador_txt, valor_xml)

                        def _substituir_marcador_runs(texto_xml, marcador, token):
                            marcador = str(marcador or "").strip()
                            if not marcador:
                                return texto_xml
                            # Word pode fragmentar texto com varias tags intermediarias (run, proofErr, smartTag etc.).
                            # Aqui permitimos qualquer tag entre os caracteres do marcador.
                            sep_runs = r"(?:\s|<[^>]+>)*"
                            letras = [re.escape(ch) for ch in marcador]
                            padrao_letras = sep_runs.join(letras)
                            aspas = r"[\"“”‘’']"
                            texto_xml = re.sub(
                                rf"(?i)\[\s*{padrao_letras}\s*\]",
                                token,
                                texto_xml,
                                flags=re.DOTALL,
                            )
                            texto_xml = re.sub(
                                rf"(?i)(?<![A-Za-zÀ-ÿ0-9_]){padrao_letras}(?![A-Za-zÀ-ÿ0-9_])",
                                token,
                                texto_xml,
                                flags=re.DOTALL,
                            )
                            # Variação entre aspas (Word pode quebrar em vários runs)
                            texto_xml = re.sub(
                                rf"({aspas}){sep_runs}{padrao_letras}{sep_runs}({aspas})",
                                token,
                                texto_xml,
                                flags=re.IGNORECASE | re.DOTALL,
                            )
                            return texto_xml

                        empresa1_token = "__CADNR_EMPRESA1__"
                        empresa3_token = "__CADNR_EMPRESA3__"
                        empresa2_token = "__CADNR_EMPRESA2__"
                        empresa4_token = "__CADNR_EMPRESA4__"
                        cidade3_token = "__CADNR_CIDADE3__"
                        cidade1_token = "__CADNR_CIDADE1__"
                        endereco3_token = "__CADNR_ENDERECO3__"
                        endereco1_token = "__CADNR_ENDERECO1__"
                        cnpj1_token = "__CADNR_CNPJ1__"
                        cnpj2_token = "__CADNR_CNPJ2__"
                        cnpj3_token = "__CADNR_CNPJ3__"
                        cnpj4_token = "__CADNR_CNPJ4__"
                        cpf1_token = "__CADNR_CPF1__"
                        cpf2_token = "__CADNR_CPF2__"
                        cpf3_token = "__CADNR_CPF3__"
                        cpf4_token = "__CADNR_CPF4__"
                        profissao1_token = "__CADNR_PROFISSAO1__"
                        funcionario1_token = "__CADNR_FUNCIONARIO1__"
                        funcionario2_token = "__CADNR_FUNCIONARIO2__"
                        funcionario3_token = "__CADNR_FUNCIONARIO3__"
                        funcionario4_token = "__CADNR_FUNCIONARIO4__"
                        admissao1_token = "__CADNR_ADMISSAO1__"
                        admissao2_token = "__CADNR_ADMISSAO2__"
                        admissao3_token = "__CADNR_ADMISSAO3__"
                        admissao4_token = "__CADNR_ADMISSAO4__"
                        comeco1_token = "__CADNR_COMECO1__"
                        fim1_token = "__CADNR_FIM1__"
                        fim2_token = "__CADNR_FIM2__"
                        datanr_token = "__CADNR_DATANR__"
                        foto1_token = "__CADNR_FOTO1__"
                        logo1_token = "__CADNR_LOGO1__"
                        logo2_token = "__CADNR_LOGO2__"
                        logo3_token = "__CADNR_LOGO3__"
                        img1_token = "__CADNR_IMG1__"
                        img2_token = "__CADNR_IMG2__"
                        img3_token = "__CADNR_IMG3__"
                        texto = _substituir_marcador_runs(texto, "EMPRESA1", empresa1_token)
                        texto = _substituir_marcador_runs(texto, "EMPRESA3", empresa3_token)
                        texto = _substituir_marcador_runs(texto, "EMPRESA2", empresa2_token)
                        texto = _substituir_marcador_runs(texto, "EMPRESA4", empresa4_token)
                        texto = _substituir_marcador_runs(texto, "CIDADE3", cidade3_token)
                        texto = _substituir_marcador_runs(texto, "CIDADE1", cidade1_token)
                        texto = _substituir_marcador_runs(texto, "ENDERECO3", endereco3_token)
                        texto = _substituir_marcador_runs(texto, "ENDEREÇO3", endereco3_token)
                        texto = _substituir_marcador_runs(texto, "ENDERECO1", endereco1_token)
                        texto = _substituir_marcador_runs(texto, "ENDEREÇO1", endereco1_token)
                        texto = _substituir_marcador_runs(texto, "CNPJ1", cnpj1_token)
                        texto = _substituir_marcador_runs(texto, "CNPJ2", cnpj2_token)
                        texto = _substituir_marcador_runs(texto, "CNPJ3", cnpj3_token)
                        texto = _substituir_marcador_runs(texto, "CNPJ4", cnpj4_token)
                        texto = _substituir_marcador_runs(texto, "CPF1", cpf1_token)
                        texto = _substituir_marcador_runs(texto, "CPF2", cpf2_token)
                        texto = _substituir_marcador_runs(texto, "CPF3", cpf3_token)
                        texto = _substituir_marcador_runs(texto, "CPF4", cpf4_token)
                        texto = _substituir_marcador_runs(texto, "PROFISSAO1", profissao1_token)
                        texto = _substituir_marcador_runs(texto, "FUNCIONARIO1", funcionario1_token)
                        texto = _substituir_marcador_runs(texto, "FUNCIONARIO2", funcionario2_token)
                        texto = _substituir_marcador_runs(texto, "FUNCIONARIO3", funcionario3_token)
                        texto = _substituir_marcador_runs(texto, "FUNCIONARIO4", funcionario4_token)
                        texto = _substituir_marcador_runs(texto, "ADMISSAO1", admissao1_token)
                        texto = _substituir_marcador_runs(texto, "ADMISSAO2", admissao2_token)
                        texto = _substituir_marcador_runs(texto, "ADMISSAO3", admissao3_token)
                        texto = _substituir_marcador_runs(texto, "ADMISSAO4", admissao4_token)
                        texto = _substituir_marcador_runs(texto, "COMECO1", comeco1_token)
                        texto = _substituir_marcador_runs(texto, "COMEÇO1", comeco1_token)
                        texto = _substituir_marcador_runs(texto, "FIM1", fim1_token)
                        texto = _substituir_marcador_runs(texto, "FIM2", fim2_token)
                        texto = _substituir_marcador_runs(texto, "DATANR", datanr_token)
                        texto = _substituir_marcador_runs(texto, "DATA_NR", datanr_token)
                        texto = _substituir_marcador_runs(texto, "FOTO1", foto1_token)
                        texto = _substituir_marcador_runs(texto, "LOGO1", logo1_token)
                        texto = _substituir_marcador_runs(texto, "LOGO2", logo2_token)
                        texto = _substituir_marcador_runs(texto, "LOGO3", logo3_token)
                        texto = _substituir_marcador_runs(texto, "IMG1", img1_token)
                        texto = _substituir_marcador_runs(texto, "IMG2", img2_token)
                        texto = _substituir_marcador_runs(texto, "IMG3", img3_token)

                        de_e_ate_xml = self._escape_xml_text(campos.get("DE_E_ATE_EXTENSO", ""))
                        if de_e_ate_xml:
                            texto = re.sub(
                                r"\[\s*Come(?:c|ç)o1\s+e\s+Fim1\s*\]",
                                de_e_ate_token,
                                texto,
                                flags=re.IGNORECASE,
                            )
                        de_a_ate_xml = self._escape_xml_text(campos.get("DE_A_ATE_EXTENSO", ""))
                        if de_a_ate_xml:
                            texto = re.sub(
                                r"\[\s*Come(?:c|ç)o1\s+a\s+Fim1\s*\]",
                                de_a_ate_token,
                                texto,
                                flags=re.IGNORECASE,
                            )

                        # Fallback tolerante para variacoes de espaco/quebra de linha.
                        nome_func_fallback = self._escape_xml_text(campos.get("NOME", ""))
                        if nome_func_fallback:
                            texto = re.sub(
                                r"(?i)\bNOME\s+DO\s+FUNCION(?:A|Á)RIO\b",
                                nome_func_fallback,
                                texto,
                            )
                            texto = re.sub(
                                r"(?i)\s+do\s+funcion(?:a|á)rio\b",
                                "",
                                texto,
                            )

                        # Remove colchetes remanescentes apos as substituicoes.
                        texto = texto.replace("[", "").replace("]", "")

                        # Fallback para marcador "FIM1" que pode sobrar sem colchetes.
                        data_ate_xml = self._escape_xml_text(campos.get("FIM1", ""))
                        if data_ate_xml:
                            texto = re.sub(
                                r"(?i)(?<![A-Za-zÀ-ÿ0-9])fim1(?![A-Za-zÀ-ÿ0-9])",
                                data_ate_xml,
                                texto,
                            )
                        data_de_xml = self._escape_xml_text(campos.get("DATA_DE_EXTENSO", ""))
                        if data_de_xml and data_ate_xml:
                            # Caso simples: texto continuo "De a <data ate>".
                            texto = re.sub(
                                rf"(?i)(?<![A-Za-zÀ-ÿ0-9])de\s+a\s+{re.escape(data_ate_xml)}",
                                f"{data_de_xml} a {data_ate_xml}",
                                texto,
                            )
                            # Caso Word com quebras em runs (<w:t>...):
                            # permite tags entre "De", "a" e a data final.
                            sep_runs = r"(?:\s|</w:t>\s*</w:r>\s*<w:r[^>]*>\s*<w:t[^>]*>)+"
                            texto = re.sub(
                                rf"(?i)(?<![A-Za-zÀ-ÿ0-9])de{sep_runs}a{sep_runs}{re.escape(data_ate_xml)}",
                                f"{data_de_xml} a {data_ate_xml}",
                                texto,
                                flags=re.DOTALL,
                            )
                        if data_de_xml:
                            # Regra generica: "De a <data>" vira "<data_de> a <data>".
                            texto = re.sub(
                                r"(?i)(?<![A-Za-zÀ-ÿ0-9])de\s+a\s+(\d{1,2}\s+de\s+[A-Za-zÀ-ÿ]+\s+de\s+\d{4})",
                                rf"{data_de_xml} a \1",
                                texto,
                            )
                            # Regra generica: "Comeco/Começo a <data>" vira "<data_de> a <data>".
                            texto = re.sub(
                                r"(?i)(?<![A-Za-zÀ-ÿ0-9])come(?:c|ç)o\s+a\s+(\d{1,2}\s+de\s+[A-Za-zÀ-ÿ]+\s+de\s+\d{4})",
                                rf"{data_de_xml} a \1",
                                texto,
                            )
                            # Versao tolerante a quebras de runs no Word.
                            sep_runs = r"(?:\s|</w:t>\s*</w:r>\s*<w:r[^>]*>\s*<w:t[^>]*>)+"
                            data_pattern_runs = (
                                rf"\d{{1,2}}{sep_runs}de{sep_runs}[A-Za-zÀ-ÿ]+{sep_runs}de{sep_runs}\d{{4}}"
                            )
                            texto = re.sub(
                                rf"(?i)(?<![A-Za-zÀ-ÿ0-9])de{sep_runs}a{sep_runs}({data_pattern_runs})",
                                lambda m: f"{data_de_xml} a {m.group(1)}",
                                texto,
                                flags=re.DOTALL,
                            )
                            texto = re.sub(
                                rf"(?i)(?<![A-Za-zÀ-ÿ0-9])come(?:c|ç)o{sep_runs}a{sep_runs}({data_pattern_runs})",
                                lambda m: f"{data_de_xml} a {m.group(1)}",
                                texto,
                                flags=re.DOTALL,
                            )
                        nome_upper_xml = self._escape_xml_text(
                            str(campos.get("NOME_DO_FUNCIONARIO_1", "") or "")
                        )
                        if nome_upper_xml:
                            texto = re.sub(r"\bNOME\b", nome_upper_xml, texto)
                            texto = re.sub(
                                r"(?i)\bFUNCION(?:A|Á)RIO\b",
                                nome_upper_xml,
                                texto,
                            )
                            texto = re.sub(
                                r"(<w:t[^>]*>)\s*N\s*</w:t>\s*</w:r>\s*<w:r[^>]*>.*?<w:t[^>]*>\s*O\s*</w:t>\s*</w:r>\s*<w:r[^>]*>.*?<w:t[^>]*>\s*M\s*</w:t>\s*</w:r>\s*<w:r[^>]*>.*?<w:t[^>]*>\s*E\s*(</w:t>)",
                                rf"\g<1>{nome_upper_xml}\g<2>",
                                texto,
                                flags=re.IGNORECASE | re.DOTALL,
                            )

                        # Fallback para textos fixos do modelo que nao estao com placeholder.
                        nome_func = self._escape_xml_text(campos.get("NOME", ""))
                        nome_func_upper = self._escape_xml_text(
                            str(campos.get("NOME_DO_FUNCIONARIO_1", "") or "")
                        )
                        if nome_func:
                            texto = texto.replace("<w:t>Nome</w:t>", f"<w:t>{nome_func}</w:t>")
                            texto = re.sub(
                                r"(<w:t[^>]*>\s*)NOME(\s*</w:t>)",
                                rf"\g<1>{nome_func_upper or nome_func}\g<2>",
                                texto,
                                flags=re.IGNORECASE,
                            )
                        funcao_func = self._escape_xml_text(
                            str(campos.get("FUNCAO", "") or "").upper()
                        )
                        if funcao_func:
                            texto = re.sub(
                                r"(?i)\bPROFISSAO1\b",
                                funcao_func,
                                texto,
                            )
                        texto_certifica = self._escape_xml_text(campos.get("TEXTO_CERTIFICA", ""))
                        if texto_certifica:
                            texto = re.sub(
                                r"(?i)A\s+[^<]{0,220}?CERTIFICA\s+que:",
                                texto_certifica,
                                texto,
                            )
                            sep_runs = r"(?:\s|</w:t>\s*</w:r>\s*<w:r[^>]*>\s*<w:t[^>]*>)+"
                            texto = re.sub(
                                rf"(?i)A{sep_runs}[^<]{{0,220}}?{sep_runs}CERTIFICA{sep_runs}que:",
                                texto_certifica,
                                texto,
                                flags=re.DOTALL,
                            )
                        texto = re.sub(
                            r"(?i)Obteve aproveitamento satisfat[oó]rio no curso de capacita[cç][aã]o(?:\s+de\s+RECICLAGEM)?",
                            self._escape_xml_text(campos.get("TEXTO_CAPACITACAO", "")),
                            texto,
                        )
                        def _adicionar_bold_run(run_xml):
                            bloco_rpr = re.search(r"<w:rPr>.*?</w:rPr>", run_xml, flags=re.DOTALL)
                            if bloco_rpr:
                                rpr_txt = bloco_rpr.group(0)
                                if "<w:b" in rpr_txt:
                                    return run_xml
                                novo_rpr = rpr_txt.replace("</w:rPr>", "<w:b/></w:rPr>")
                                return run_xml.replace(rpr_txt, novo_rpr, 1)
                            idx = run_xml.find(">")
                            if idx < 0:
                                return run_xml
                            return run_xml[: idx + 1] + "<w:rPr><w:b/></w:rPr>" + run_xml[idx + 1 :]

                        def _aplicar_negrito_trechos(texto_xml, trechos):
                            trechos_validos = [t for t in trechos if t]
                            if not trechos_validos:
                                return texto_xml

                            def _processar_paragrafo(mp):
                                paragrafo = mp.group(1)
                                runs = []
                                texto_total = ""
                                pos = 0
                                for mr in re.finditer(r"<w:r\b[^>]*>.*?</w:r>", paragrafo, flags=re.DOTALL):
                                    run_xml = mr.group(0)
                                    partes_t = re.findall(r"<w:t[^>]*>(.*?)</w:t>", run_xml, flags=re.DOTALL)
                                    run_texto = "".join(partes_t)
                                    ini = pos
                                    fim = pos + len(run_texto)
                                    runs.append(
                                        {
                                            "xml_ini": mr.start(),
                                            "xml_fim": mr.end(),
                                            "txt_ini": ini,
                                            "txt_fim": fim,
                                            "xml": run_xml,
                                        }
                                    )
                                    texto_total += run_texto
                                    pos = fim

                                if not texto_total:
                                    return paragrafo

                                ranges = []
                                texto_lower = texto_total.lower()
                                for trecho in trechos_validos:
                                    t = trecho.strip()
                                    if not t:
                                        continue
                                    t_lower = t.lower()
                                    i = texto_lower.find(t_lower)
                                    while i != -1:
                                        ranges.append((i, i + len(t)))
                                        i = texto_lower.find(t_lower, i + 1)

                                if not ranges:
                                    return paragrafo

                                saida = []
                                cursor = 0
                                for run in runs:
                                    sobrepoe = False
                                    for a, b in ranges:
                                        if run["txt_fim"] > a and run["txt_ini"] < b:
                                            sobrepoe = True
                                            break
                                    run_xml_final = _adicionar_bold_run(run["xml"]) if sobrepoe else run["xml"]
                                    saida.append(paragrafo[cursor : run["xml_ini"]])
                                    saida.append(run_xml_final)
                                    cursor = run["xml_fim"]
                                saida.append(paragrafo[cursor:])
                                return "".join(saida)

                            return re.sub(
                                r"(<w:p\b[^>]*>.*?</w:p>)",
                                _processar_paragrafo,
                                texto_xml,
                                flags=re.DOTALL,
                            )

                        def _adicionar_italico_run(run_xml):
                            bloco_rpr = re.search(r"<w:rPr>.*?</w:rPr>", run_xml, flags=re.DOTALL)
                            if bloco_rpr:
                                rpr_txt = bloco_rpr.group(0)
                                if "<w:i" in rpr_txt:
                                    return run_xml
                                novo_rpr = rpr_txt.replace("</w:rPr>", "<w:i/></w:rPr>")
                                return run_xml.replace(rpr_txt, novo_rpr, 1)
                            idx = run_xml.find(">")
                            if idx < 0:
                                return run_xml
                            return run_xml[: idx + 1] + "<w:rPr><w:i/></w:rPr>" + run_xml[idx + 1 :]

                        def _aplicar_italico_trechos(texto_xml, trechos):
                            trechos_validos = [t for t in trechos if t]
                            if not trechos_validos:
                                return texto_xml

                            def _processar_paragrafo(mp):
                                paragrafo = mp.group(1)
                                runs = []
                                texto_total = ""
                                pos = 0
                                for mr in re.finditer(r"<w:r\b[^>]*>.*?</w:r>", paragrafo, flags=re.DOTALL):
                                    run_xml = mr.group(0)
                                    partes_t = re.findall(r"<w:t[^>]*>(.*?)</w:t>", run_xml, flags=re.DOTALL)
                                    run_texto = "".join(partes_t)
                                    ini = pos
                                    fim = pos + len(run_texto)
                                    runs.append(
                                        {
                                            "xml_ini": mr.start(),
                                            "xml_fim": mr.end(),
                                            "txt_ini": ini,
                                            "txt_fim": fim,
                                            "xml": run_xml,
                                        }
                                    )
                                    texto_total += run_texto
                                    pos = fim

                                if not texto_total:
                                    return paragrafo

                                ranges = []
                                texto_lower = texto_total.lower()
                                for trecho in trechos_validos:
                                    t = trecho.strip()
                                    if not t:
                                        continue
                                    t_lower = t.lower()
                                    i = texto_lower.find(t_lower)
                                    while i != -1:
                                        ranges.append((i, i + len(t)))
                                        i = texto_lower.find(t_lower, i + 1)

                                if not ranges:
                                    return paragrafo

                                saida = []
                                cursor = 0
                                for run in runs:
                                    sobrepoe = False
                                    for a, b in ranges:
                                        if run["txt_fim"] > a and run["txt_ini"] < b:
                                            sobrepoe = True
                                            break
                                    run_xml_final = _adicionar_italico_run(run["xml"]) if sobrepoe else run["xml"]
                                    saida.append(paragrafo[cursor : run["xml_ini"]])
                                    saida.append(run_xml_final)
                                    cursor = run["xml_fim"]
                                saida.append(paragrafo[cursor:])
                                return "".join(saida)

                            return re.sub(
                                r"(<w:p\b[^>]*>.*?</w:p>)",
                                _processar_paragrafo,
                                texto_xml,
                                flags=re.DOTALL,
                            )

                        def _remover_negrito_trechos(texto_xml, trechos):
                            trechos_validos = [t for t in trechos if t]
                            if not trechos_validos:
                                return texto_xml

                            def _remover_bold_run(run_xml):
                                bloco_rpr = re.search(r"<w:rPr>.*?</w:rPr>", run_xml, flags=re.DOTALL)
                                if not bloco_rpr:
                                    return run_xml
                                rpr_txt = bloco_rpr.group(0)
                                novo_rpr = re.sub(r"<w:b\b[^>]*/>", "", rpr_txt)
                                novo_rpr = re.sub(r"<w:bCs\b[^>]*/>", "", novo_rpr)
                                novo_rpr = re.sub(r"<w:rPr>\s*</w:rPr>", "", novo_rpr)
                                return run_xml.replace(rpr_txt, novo_rpr, 1)

                            def _processar_paragrafo(mp):
                                paragrafo = mp.group(1)
                                runs = []
                                texto_total = ""
                                pos = 0
                                for mr in re.finditer(r"<w:r\b[^>]*>.*?</w:r>", paragrafo, flags=re.DOTALL):
                                    run_xml = mr.group(0)
                                    partes_t = re.findall(r"<w:t[^>]*>(.*?)</w:t>", run_xml, flags=re.DOTALL)
                                    run_texto = "".join(partes_t)
                                    ini = pos
                                    fim = pos + len(run_texto)
                                    runs.append(
                                        {
                                            "xml_ini": mr.start(),
                                            "xml_fim": mr.end(),
                                            "txt_ini": ini,
                                            "txt_fim": fim,
                                            "xml": run_xml,
                                        }
                                    )
                                    texto_total += run_texto
                                    pos = fim

                                if not texto_total:
                                    return paragrafo

                                ranges = []
                                texto_lower = texto_total.lower()
                                for trecho in trechos_validos:
                                    t = trecho.strip()
                                    if not t:
                                        continue
                                    t_lower = t.lower()
                                    i = texto_lower.find(t_lower)
                                    while i != -1:
                                        ranges.append((i, i + len(t)))
                                        i = texto_lower.find(t_lower, i + 1)

                                if not ranges:
                                    return paragrafo

                                saida = []
                                cursor = 0
                                for run in runs:
                                    sobrepoe = False
                                    for a, b in ranges:
                                        if run["txt_fim"] > a and run["txt_ini"] < b:
                                            sobrepoe = True
                                            break
                                    run_xml_final = _remover_bold_run(run["xml"]) if sobrepoe else run["xml"]
                                    saida.append(paragrafo[cursor : run["xml_ini"]])
                                    saida.append(run_xml_final)
                                    cursor = run["xml_fim"]
                                saida.append(paragrafo[cursor:])
                                return "".join(saida)

                            return re.sub(
                                r"(<w:p\b[^>]*>.*?</w:p>)",
                                _processar_paragrafo,
                                texto_xml,
                                flags=re.DOTALL,
                            )

                        def _aplicar_negrito_data_realizado(texto_xml, data_extenso):
                            data_extenso = str(data_extenso or "").strip()
                            if not data_extenso:
                                return texto_xml
                            prefixo = "e realizado no dia "
                            data_lower = data_extenso.lower()

                            def _processar_paragrafo(mp):
                                paragrafo = mp.group(1)
                                runs = []
                                texto_total = ""
                                pos = 0
                                for mr in re.finditer(r"<w:r\b[^>]*>.*?</w:r>", paragrafo, flags=re.DOTALL):
                                    run_xml = mr.group(0)
                                    partes_t = re.findall(r"<w:t[^>]*>(.*?)</w:t>", run_xml, flags=re.DOTALL)
                                    run_texto = "".join(partes_t)
                                    ini = pos
                                    fim = pos + len(run_texto)
                                    runs.append(
                                        {
                                            "xml_ini": mr.start(),
                                            "xml_fim": mr.end(),
                                            "txt_ini": ini,
                                            "txt_fim": fim,
                                            "xml": run_xml,
                                        }
                                    )
                                    texto_total += run_texto
                                    pos = fim

                                if not texto_total:
                                    return paragrafo

                                texto_lower = texto_total.lower()
                                idx_prefixo = texto_lower.find(prefixo)
                                if idx_prefixo < 0:
                                    return paragrafo

                                idx_data = texto_lower.find(data_lower, idx_prefixo + len(prefixo))
                                if idx_data < 0:
                                    return paragrafo

                                a, b = idx_data, idx_data + len(data_extenso)
                                saida = []
                                cursor = 0
                                for run in runs:
                                    sobrepoe = run["txt_fim"] > a and run["txt_ini"] < b
                                    run_xml_final = _adicionar_bold_run(run["xml"]) if sobrepoe else run["xml"]
                                    saida.append(paragrafo[cursor : run["xml_ini"]])
                                    saida.append(run_xml_final)
                                    cursor = run["xml_fim"]
                                saida.append(paragrafo[cursor:])
                                return "".join(saida)

                            return re.sub(
                                r"(<w:p\b[^>]*>.*?</w:p>)",
                                _processar_paragrafo,
                                texto_xml,
                                flags=re.DOTALL,
                            )

                        texto = _aplicar_negrito_data_realizado(
                            texto,
                            self._escape_xml_text(campos.get("DATA_EXTENSO", "")),
                        )
                        texto = _aplicar_negrito_trechos(texto, [de_e_ate_token, de_a_ate_token])
                        texto = _aplicar_negrito_trechos(
                            texto,
                            [
                                funcionario1_token,
                                cpf1_token,
                                profissao1_token,
                                fim1_token,
                                comeco1_token,
                                empresa4_token,
                                cnpj4_token,
                                empresa2_token,
                                cnpj2_token,
                                empresa3_token,
                                cnpj3_token,
                                funcionario2_token,
                                cpf2_token,
                            ],
                        )
                        texto = _aplicar_italico_trechos(
                            texto,
                            [
                                admissao1_token,
                                funcionario3_token,
                                cpf3_token,
                                empresa2_token,
                                cnpj2_token,
                                empresa3_token,
                                cnpj3_token,
                                funcionario2_token,
                                cpf2_token,
                            ],
                        )
                        texto = texto.replace(de_e_ate_token, de_e_ate_xml)
                        texto = texto.replace(de_a_ate_token, de_a_ate_xml)
                        texto = texto.replace(empresa1_token, self._escape_xml_text(campos.get("EMPRESA1", "")))
                        texto = texto.replace(empresa3_token, self._escape_xml_text(campos.get("EMPRESA3", "")))
                        texto = texto.replace(empresa2_token, self._escape_xml_text(campos.get("EMPRESA2", "")))
                        texto = texto.replace(empresa4_token, self._escape_xml_text(campos.get("EMPRESA4", "")))
                        texto = texto.replace(cidade3_token, self._escape_xml_text(campos.get("CIDADE3", "")))
                        texto = texto.replace(cidade1_token, self._escape_xml_text(campos.get("CIDADE1", "")))
                        texto = texto.replace(endereco3_token, self._escape_xml_text(campos.get("ENDERECO3", "")))
                        texto = texto.replace(endereco1_token, self._escape_xml_text(campos.get("ENDERECO1", "")))
                        texto = texto.replace(cnpj1_token, self._escape_xml_text(campos.get("CNPJ1", "")))
                        texto = texto.replace(cnpj2_token, self._escape_xml_text(campos.get("CNPJ2", "")))
                        texto = texto.replace(cnpj3_token, self._escape_xml_text(campos.get("CNPJ3", "")))
                        texto = texto.replace(cnpj4_token, self._escape_xml_text(campos.get("CNPJ4", "")))
                        texto = texto.replace(profissao1_token, self._escape_xml_text(campos.get("PROFISSAO1", "")))
                        texto = texto.replace(funcionario1_token, self._escape_xml_text(campos.get("FUNCIONARIO1", "")))
                        funcionario2_sem_quebra = str(campos.get("FUNCIONARIO2", "") or "").replace(" ", "\u00A0")
                        texto = texto.replace(funcionario2_token, self._escape_xml_text(funcionario2_sem_quebra))
                        texto = texto.replace(funcionario3_token, self._escape_xml_text(campos.get("FUNCIONARIO3", "")))
                        texto = texto.replace(funcionario4_token, self._escape_xml_text(campos.get("FUNCIONARIO4", "")))
                        texto = texto.replace(cpf1_token, self._escape_xml_text(campos.get("CPF1", "")))
                        texto = texto.replace(cpf2_token, self._escape_xml_text(campos.get("CPF2", "")))
                        texto = texto.replace(cpf3_token, self._escape_xml_text(campos.get("CPF3", "")))
                        texto = texto.replace(cpf4_token, self._escape_xml_text(campos.get("CPF4", "")))
                        texto = texto.replace(admissao1_token, self._escape_xml_text(campos.get("ADMISSAO1", "")))
                        texto = texto.replace(admissao2_token, self._escape_xml_text(campos.get("ADMISSAO2", "")))
                        texto = texto.replace(admissao3_token, self._escape_xml_text(campos.get("ADMISSAO3", "")))
                        texto = texto.replace(admissao4_token, self._escape_xml_text(campos.get("ADMISSAO4", "")))
                        texto = texto.replace(comeco1_token, self._escape_xml_text(campos.get("COMECO1", "")))
                        texto = texto.replace(fim1_token, self._escape_xml_text(campos.get("FIM1", "")))
                        texto = texto.replace(fim2_token, self._escape_xml_text(campos.get("FIM2", "")))
                        texto = texto.replace(datanr_token, self._escape_xml_text(campos.get("DATANR", "")))
                        texto = texto.replace(foto1_token, "foto1")
                        texto = texto.replace(logo1_token, "logo1")
                        texto = texto.replace(logo2_token, "logo2")
                        texto = texto.replace(logo3_token, "logo3")
                        texto = texto.replace(img1_token, "img1")
                        texto = texto.replace(img2_token, "img2")
                        texto = texto.replace(img3_token, "img3")
                        # Limpa artefatos como "__CADNR_03 de fevereiro de 2026__".
                        padrao_data = (
                            r"(?:\d{1,2}\s+de\s+[A-Za-zÀ-ÿ]+\s+de\s+\d{4}|\d{2}/\d{2}/\d{4}|\d{2}/\d{4})"
                        )
                        texto = re.sub(
                            rf"__\s*CADNR\s*[_-]?\s*({padrao_data})\s*__",
                            r"\1",
                            texto,
                            flags=re.IGNORECASE,
                        )
                        texto = re.sub(
                            rf"__\s*({padrao_data})\s*__",
                            r"\1",
                            texto,
                            flags=re.IGNORECASE,
                        )
                        data = texto.encode("utf-8")
                    zout.writestr(info, data)
        except OSError:
            messagebox.showerror("NR", "Nao foi possivel gerar o arquivo Word atualizado.")
            return None

        if foto1_caminho is not None and foto1_caminho.exists():
            inseriu_foto, motivo_foto = self._inserir_imagem_por_marcador_docx(
                destino,
                foto1_caminho,
                ["foto1"],
                largura_cm=2.5,
                altura_cm=3.2,
            )
            if (not inseriu_foto) and ("nao encontrado no documento" not in str(motivo_foto or "").lower()):
                messagebox.showwarning(
                    "NR",
                    f"A foto do funcionario nao foi inserida no Word. Motivo: {motivo_foto}",
                )
        if logo_img1_caminho is not None and logo_img1_caminho.exists():
            inseriu_logo, motivo_logo = self._inserir_img1_logo_em_tabela_docx(destino, logo_img1_caminho)
            if (not inseriu_logo) and ("nao encontrado no documento" not in str(motivo_logo or "").lower()):
                messagebox.showwarning(
                    "NR",
                    f"A logo nao foi inserida no Word. Motivo: {motivo_logo}",
                )
        if logo_img2_caminho is not None and logo_img2_caminho.exists():
            inseriu_logo2, motivo_logo2 = self._inserir_img2_logo_em_tabela_docx(destino, logo_img2_caminho)
            if (not inseriu_logo2) and ("nao encontrado no documento" not in str(motivo_logo2 or "").lower()):
                messagebox.showwarning(
                    "NR",
                    f"A logo fixa da Projetta (img2) nao foi inserida no Word. Motivo: {motivo_logo2}",
                )
        if logo_img3_caminho is not None and logo_img3_caminho.exists():
            inseriu_logo3, motivo_logo3 = self._inserir_img3_logo_em_tabela_docx(destino, logo_img3_caminho)
            if (not inseriu_logo3) and ("nao encontrado no documento" not in str(motivo_logo3 or "").lower()):
                messagebox.showwarning(
                    "NR",
                    f"A logo3 da empresa nao foi inserida no Word. Motivo: {motivo_logo3}",
                )

        if exibir_mensagem:
            messagebox.showinfo("NR", f"Arquivo atualizado com sucesso:\n{destino.name}")
        return destino

    def _converter_docx_para_pdf(self, docx_path, pdf_path):
        self._ultima_falha_conversao = ""
        try:
            import win32com.client as win32com
        except Exception:
            win32com = None

        if win32com is not None:
            word = None
            documento = None
            try:
                word = win32com.DispatchEx("Word.Application")
                word.Visible = False
                documento = word.Documents.Open(str(docx_path.resolve()))
                documento.SaveAs(str(pdf_path), FileFormat=17)
                return True
            except Exception as ex:
                self._ultima_falha_conversao = f"pywin32/COM: {ex}"
            finally:
                if documento is not None:
                    documento.Close(False)
                if word is not None:
                    word.Quit()

        # Fallback Windows sem pywin32: usa automacao COM do Word via PowerShell.
        try:
            docx_ps = str(docx_path.resolve()).replace("'", "''")
            pdf_ps = str(pdf_path.resolve()).replace("'", "''")
            ps_script = (
                "$ErrorActionPreference='Stop';"
                "$word=$null;$doc=$null;"
                "try{"
                "$word=New-Object -ComObject Word.Application;"
                "$word.Visible=$false;"
                f"$doc=$word.Documents.Open('{docx_ps}');"
                f"$doc.SaveAs('{pdf_ps}',17);"
                "}finally{"
                "if($doc -ne $null){$doc.Close($false)};"
                "if($word -ne $null){$word.Quit()}"
                "}"
            )
            resultado_ps = subprocess.run(
                [
                    "powershell",
                    "-NoProfile",
                    "-ExecutionPolicy",
                    "Bypass",
                    "-Command",
                    ps_script,
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=False,
                timeout=90,
            )
            if resultado_ps.returncode == 0 and pdf_path.exists():
                return True
            stderr_ps = (resultado_ps.stderr or b"").decode("utf-8", errors="ignore").strip()
            if stderr_ps:
                self._ultima_falha_conversao = f"PowerShell COM: {stderr_ps}"
        except (OSError, subprocess.SubprocessError):
            self._ultima_falha_conversao = "PowerShell COM indisponivel no ambiente."

        candidatos = (
            "soffice",
            "soffice.bin",
            "libreoffice",
            "libreoffice7.4",
            "libreoffice7.3",
        )
        comando_soffice = next((shutil.which(c) for c in candidatos), None)
        if comando_soffice is None:
            if not self._ultima_falha_conversao:
                self._ultima_falha_conversao = "LibreOffice/soffice nao encontrado."
            return False

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_saida = Path(temp_dir)
            try:
                resultado = subprocess.run(
                    [
                        comando_soffice,
                        "--headless",
                        "--convert-to",
                        "pdf",
                        "--outdir",
                        str(temp_saida),
                        str(docx_path),
                    ],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    check=False,
                    timeout=60,
                )
            except OSError:
                self._ultima_falha_conversao = "Falha ao executar LibreOffice/soffice."
                return False

            if resultado.returncode != 0:
                stderr_soffice = (resultado.stderr or b"").decode("utf-8", errors="ignore").strip()
                if stderr_soffice:
                    self._ultima_falha_conversao = f"LibreOffice: {stderr_soffice}"
                return False

            convertido = temp_saida / f"{docx_path.stem}.pdf"
            if not convertido.exists():
                self._ultima_falha_conversao = "LibreOffice executou, mas nao gerou o PDF."
                return False

            shutil.move(str(convertido), str(pdf_path))
            return True

    @staticmethod
    def _pdf_escape_text(texto):
        return (
            str(texto or "")
            .replace("\\", "\\\\")
            .replace("(", "\\(")
            .replace(")", "\\)")
        )

    def _escrever_pdf_texto(self, destino, titulo, linhas):
        largura = 595
        altura = 842
        margem_x = 50
        topo = 790
        passo = 18

        comandos = []
        y = topo
        comandos.append("BT /F1 16 Tf 1 0 0 1 %d %d Tm (%s) Tj ET" % (
            margem_x, y, self._pdf_escape_text(titulo)
        ))
        y -= 28
        comandos.append("BT /F1 11 Tf 1 0 0 1 %d %d Tm (%s) Tj ET" % (
            margem_x, y, self._pdf_escape_text("NR selecionadas:")
        ))
        y -= 24
        if not linhas:
            comandos.append("BT /F1 11 Tf 1 0 0 1 %d %d Tm (%s) Tj ET" % (
                margem_x, y, self._pdf_escape_text("Nenhuma NR selecionada.")
            ))
        else:
            for linha in linhas:
                if y < 60:
                    break
                comandos.append("BT /F1 10 Tf 1 0 0 1 %d %d Tm (%s) Tj ET" % (
                    margem_x, y, self._pdf_escape_text(linha)
                ))
                y -= passo

        stream = "\n".join(comandos).encode("latin-1", errors="replace")
        objetos = []
        objetos.append(b"<< /Type /Catalog /Pages 2 0 R >>")
        objetos.append(b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>")
        objetos.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 {largura} {altura}] "
            f"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>".encode("ascii")
        )
        objetos.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
        objetos.append(f"<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream")

        saida = b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n"
        offsets = [0]
        for i, obj in enumerate(objetos, start=1):
            offsets.append(len(saida))
            saida += f"{i} 0 obj\n".encode("ascii") + obj + b"\nendobj\n"
        xref = len(saida)
        saida += f"xref\n0 {len(objetos) + 1}\n".encode("ascii")
        saida += b"0000000000 65535 f \n"
        for off in offsets[1:]:
            saida += f"{off:010d} 00000 n \n".encode("ascii")
        saida += (
            f"trailer << /Size {len(objetos) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref}\n%%EOF\n"
        ).encode("ascii")
        destino.write_bytes(saida)

    def _limpar_metadados_pdf(self, pdf_path):
        caminho = Path(pdf_path)
        if not caminho.exists() or not caminho.is_file():
            return False, "arquivo PDF nao encontrado"
        try:
            import fitz
        except Exception:
            return False, "PyMuPDF (fitz) indisponivel"

        temp_saida = caminho.with_name(f"{caminho.stem}.__sem_meta__.pdf")
        try:
            doc = fitz.open(str(caminho))
            meta = dict(doc.metadata or {})
            for chave in list(meta.keys()):
                meta[chave] = ""
            # Reforco para campos solicitados.
            meta["title"] = ""
            meta["author"] = ""
            meta["subject"] = ""
            doc.set_metadata(meta)
            try:
                doc.set_xml_metadata(self._xmp_minimo_pdf())
            except Exception:
                pass
            doc.save(str(temp_saida), garbage=4, deflate=True)
            doc.close()
            shutil.move(str(temp_saida), str(caminho))
            return True, ""
        except Exception as ex:
            try:
                if temp_saida.exists():
                    temp_saida.unlink()
            except OSError:
                pass
            return False, str(ex)

    @staticmethod
    def _xmp_minimo_pdf():
        return (
            '<?xpacket begin="\\ufeff" id="W5M0MpCehiHzreSzNTczkc9d"?>'
            '<x:xmpmeta xmlns:x="adobe:ns:meta/">'
            '<rdf:RDF xmlns:rdf="http://www.w3.org/1999/02/22-rdf-syntax-ns#">'
            '<rdf:Description rdf:about=""/>'
            "</rdf:RDF>"
            "</x:xmpmeta>"
            '<?xpacket end="w"?>'
        )

    def _recriar_pdf_sem_xmp(self, caminho_pdf):
        try:
            import fitz
        except Exception:
            return False, "PyMuPDF (fitz) indisponivel"

        caminho = Path(str(caminho_pdf or "")).expanduser()
        if not caminho.is_absolute():
            caminho = (Path(__file__).resolve().parent / caminho).resolve()
        if not caminho.exists() or not caminho.is_file():
            return False, "arquivo PDF nao encontrado"

        temp_saida = caminho.with_name(f"{caminho.stem}.__rebuild__.pdf")
        try:
            src = fitz.open(str(caminho))
            dst = fitz.open()
            dst.insert_pdf(src)
            src.close()
            meta = dict(dst.metadata or {})
            for chave in list(meta.keys()):
                meta[chave] = ""
            meta["title"] = ""
            meta["author"] = ""
            meta["subject"] = ""
            dst.set_metadata(meta)
            try:
                dst.set_xml_metadata(self._xmp_minimo_pdf())
            except Exception:
                pass
            dst.save(str(temp_saida), garbage=4, deflate=True)
            dst.close()
            shutil.move(str(temp_saida), str(caminho))
            return True, ""
        except Exception as ex:
            try:
                if temp_saida.exists():
                    temp_saida.unlink()
            except OSError:
                pass
            return False, str(ex)

    def _gerar_pdf_nr_imprimir(self):
        self._aviso_assinatura_exibido = False
        self._aviso_git_sync_exibido = False
        funcionario = self._funcionario_ativo_para_documento()
        if funcionario is None:
            messagebox.showwarning(
                "IMPRIMIR",
                "Selecione um funcionario em CADNR ou CADASTROS.",
            )
            return

        linhas = self._linhas_nr_selecionadas()
        empresa_funcionario = funcionario.get("empresa_id")
        outros_selecionados = [
            item
            for item in self.outros_docs_imprimir
            if item.get("empresa_id") == empresa_funcionario
        ]
        def _chave_outro_documento(item_doc):
            tipo = str(item_doc.get("tipo", "") or "").strip().casefold()
            caminho = self._normalizar_caminho_documento_db(item_doc.get("caminho", ""))
            caminho_norm = str(caminho or "").strip().replace("\\", "/").casefold()
            return (item_doc.get("empresa_id"), tipo, caminho_norm)

        # Inclui automaticamente os documentos elegiveis da aba OUTROS DOCUMENTOS
        # (incluindo Carteirinha vinculada a NR), sem depender do botao ADICIONAR.
        outros_automaticos = self._coletar_outros_documentos_disponiveis(empresa_funcionario, funcionario)
        chaves = {_chave_outro_documento(item) for item in outros_selecionados}
        adicionou_auto = False
        for item_auto in outros_automaticos:
            tipo_auto_norm = str(item_auto.get("tipo", "") or "").strip().casefold()
            # Ordem de Servico e Fit Test so devem entrar quando selecionados manualmente.
            if tipo_auto_norm in {"ordem de servico", "ordem de serviço", "fit test", "fittest"}:
                continue
            chave = _chave_outro_documento(item_auto)
            if chave in chaves:
                continue
            chaves.add(chave)
            outros_selecionados.append(dict(item_auto))
            self.outros_docs_imprimir.append(dict(item_auto))
            adicionou_auto = True
        if adicionou_auto:
            self._atualizar_lista_outros_docs_imprimir()
        if not linhas and not outros_selecionados:
            messagebox.showwarning(
                "IMPRIMIR",
                "Nenhuma NR selecionada em CERTIFICADOS e nenhum OUTRO DOCUMENTO adicionado.",
            )
            return

        nome_func = str(funcionario.get("nome", "funcionario") or "funcionario").strip()
        nome_limpo = self._obter_nome_arquivo_seguro(nome_func, "funcionario")
        nome_func_arquivo = self._obter_nome_arquivo_seguro(nome_func, "funcionario")

        try:
            pasta_destino = self._obter_pasta_saida(nome_limpo)
        except OSError:
            messagebox.showerror("IMPRIMIR", "Nao foi possivel criar a pasta de saida no projeto.")
            return

        def _limpar_lista_imprimir():
            self._limpar_areas_selecao_salvar_tudo()
            self._salvar_dados()

        def _gerar_pdf_outros_documentos():
            gerados = []
            falhas = []
            base_dir = Path(__file__).resolve().parent
            funcao_selecionada = str(funcionario.get("funcao", "") or "").strip()
            nr35_selecionada = self._nr_35_selecionada()
            vistos_doc = set()
            for item in outros_selecionados:
                chave_item = _chave_outro_documento(item)
                if chave_item in vistos_doc:
                    continue
                vistos_doc.add(chave_item)
                self._ultimo_qr_embutido_docx = False
                tipo_doc = str(item.get("tipo", "") or "").strip() or "OUTRO DOCUMENTO"
                tipo_norm = tipo_doc.casefold()
                caminho_ref = str(item.get("caminho", "") or "").strip()
                if not caminho_ref:
                    continue
                if tipo_norm in {"anuencia", "anuência"} and not nr35_selecionada:
                    falhas.append(f"{tipo_doc}: ignorado (NR 35 nao selecionada)")
                    continue
                if tipo_norm == "carteirinha":
                    if not self._arquivo_vinculado_nr_carteirinha(caminho_ref):
                        falhas.append(f"{tipo_doc}: ignorado (carteirinha nao vinculada a NR selecionada)")
                        continue
                else:
                    if tipo_norm in {"ordem de servico", "ordem de serviço"}:
                        if not self._arquivo_vinculado_nr_documento(caminho_ref):
                            falhas.append(f"{tipo_doc}: ignorado (sem vinculo com NR selecionada)")
                            continue
                        if not self._arquivo_ordem_servico_compativel_funcao(caminho_ref, funcao_selecionada):
                            falhas.append(f"{tipo_doc}: ignorado (OS sem vinculo com a funcao selecionada)")
                            continue
                    elif tipo_norm == "ficha de epi":
                        if not self._arquivo_epi_compativel_funcao(caminho_ref, funcao_selecionada):
                            falhas.append(f"{tipo_doc}: ignorado (EPI sem vinculo com a funcao selecionada)")
                            continue
                    else:
                        if not self._arquivo_vinculado_nr_documento(caminho_ref):
                            falhas.append(f"{tipo_doc}: ignorado (sem vinculo com NR selecionada)")
                            continue
                        if not self._arquivo_compativel_funcao(caminho_ref, funcao_selecionada):
                            falhas.append(f"{tipo_doc}: ignorado (sem vinculo com a funcao selecionada)")
                            continue
                origem = Path(caminho_ref)
                if not origem.is_absolute():
                    origem = (base_dir / origem).resolve()
                try:
                    if not origem.exists() or not origem.is_file():
                        caminho_norm = self._normalizar_caminho_documento_db(caminho_ref)
                        origem_alt = Path(caminho_norm)
                        if not origem_alt.is_absolute():
                            origem_alt = (base_dir / origem_alt).resolve()
                        origem = origem_alt
                except OSError:
                    falhas.append(f"{tipo_doc}: caminho invalido")
                    continue
                if not origem.exists() or not origem.is_file():
                    falhas.append(f"{tipo_doc}: arquivo nao encontrado")
                    continue
                nome_base = self._obter_nome_arquivo_seguro(
                    f"{origem.stem} {nome_func_arquivo}",
                    origem.stem or tipo_doc,
                )
                destino_pdf = self._obter_arquivo_pdf_livre(
                    pasta_destino,
                    nome_base,
                    reutilizar_existente=True,
                )
                ext = origem.suffix.lower()
                tipo_carteirinha = tipo_norm == "carteirinha"
                if ext == ".pdf":
                    try:
                        shutil.copy2(origem, destino_pdf)
                    except OSError:
                        falhas.append(f"{tipo_doc}: falha ao copiar PDF")
                        continue
                    limpou_meta, _ = self._limpar_metadados_pdf(destino_pdf)
                    if not limpou_meta:
                        falhas.append(f"{tipo_doc}: falha ao limpar metadados do PDF")
                        continue
                else:
                    tipo_os = tipo_norm in {"ordem de servico", "ordem de serviço"}
                    tipo_fit_test = tipo_norm in {"fit test", "fittest"}
                    tipo_anuencia = tipo_norm in {"anuencia", "anuência"}
                    tipo_com_parametros = tipo_os or tipo_carteirinha or tipo_fit_test or tipo_anuencia
                    if tipo_com_parametros and ext == ".doc":
                        falhas.append(f"{tipo_doc}: use arquivo .docx para aplicar parametros")
                        continue

                    origem_para_converter = origem
                    temp_param_docx = None
                    if tipo_com_parametros and ext == ".docx":
                        temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                        temp_file.close()
                        temp_param_docx = Path(temp_file.name)
                        ok_param, msg_param = self._aplicar_parametros_os_docx(
                            origem,
                            temp_param_docx,
                            funcionario,
                            forcar_nome_uma_linha=tipo_carteirinha,
                            caminho_documento_ref=caminho_ref,
                            tipo_documento=tipo_doc,
                            caminho_documento_saida=destino_pdf,
                        )
                        if not ok_param:
                            falhas.append(f"{tipo_doc}: {msg_param}")
                            try:
                                temp_param_docx.unlink(missing_ok=True)
                            except OSError:
                                pass
                            continue
                        origem_para_converter = temp_param_docx

                    if not self._converter_docx_para_pdf(origem_para_converter, destino_pdf):
                        if temp_param_docx is not None:
                            try:
                                temp_param_docx.unlink(missing_ok=True)
                            except OSError:
                                pass
                        falhas.append(f"{tipo_doc}: falha ao converter para PDF")
                        continue
                    limpou_meta, _ = self._limpar_metadados_pdf(destino_pdf)
                    if not limpou_meta:
                        if temp_param_docx is not None:
                            try:
                                temp_param_docx.unlink(missing_ok=True)
                            except OSError:
                                pass
                        falhas.append(f"{tipo_doc}: falha ao limpar metadados do PDF")
                        continue
                    if temp_param_docx is not None:
                        try:
                            temp_param_docx.unlink(missing_ok=True)
                        except OSError:
                            pass
                self._registrar_documento_salvo(
                    destino_pdf,
                    origem="imprimir_outros_documentos",
                    empresa_id=funcionario.get("empresa_id"),
                    funcionario_id=funcionario.get("id"),
                    tipo_documento=tipo_doc,
                    inserir_qr_pdf=not (tipo_carteirinha and bool(getattr(self, "_ultimo_qr_embutido_docx", False))),
                )
                gerados.append(destino_pdf)
            return gerados, falhas

        with tempfile.TemporaryDirectory() as dir_temp:
            docx_temp = Path(dir_temp)
            indices_selecionados = self._indices_nr_imprimir_adicionados()
            if not indices_selecionados:
                indices_selecionados = [
                    idx for idx, item in enumerate(self.nr_certificados) if bool(item.get("imprimir", False))
                ]
            estado_original = [
                (
                    bool(item.get("imprimir", False)),
                    bool(item.get("imprimir_adicionado", False)),
                )
                for item in self.nr_certificados
            ]
            docx_paths = []
            try:
                for idx_sel in indices_selecionados:
                    for i, item in enumerate(self.nr_certificados):
                        marcado = i == idx_sel
                        item["imprimir"] = marcado
                        item["imprimir_adicionado"] = marcado
                    docx_path = self._atualizar_word_nr(pasta_destino=docx_temp, exibir_mensagem=False)
                    if docx_path is None:
                        continue
                    unico = docx_temp / f"{Path(docx_path).stem}_{idx_sel}.docx"
                    shutil.copyfile(str(docx_path), str(unico))
                    docx_paths.append((unico, idx_sel))
            finally:
                for item, estado in zip(self.nr_certificados, estado_original):
                    item["imprimir"], item["imprimir_adicionado"] = estado

            if not docx_paths:
                outros_pdfs, falhas_outros = _gerar_pdf_outros_documentos()
                if outros_pdfs:
                    self._salvar_dados()
                    _limpar_lista_imprimir()
                    lista = "\n".join(str(p) for p in outros_pdfs)
                    if falhas_outros:
                        lista_falhas = "\n".join(falhas_outros)
                        messagebox.showwarning(
                            "IMPRIMIR",
                            f"PDF(s) gerado(s):\n{lista}\n\nFalhas:\n{lista_falhas}",
                        )
                    else:
                        messagebox.showinfo("IMPRIMIR", f"PDF(s) gerado(s) com sucesso:\n{lista}")
                    return
                if falhas_outros:
                    lista_falhas = "\n".join(falhas_outros)
                    messagebox.showwarning(
                        "IMPRIMIR",
                        f"Nenhum arquivo foi gerado.\n\nFalhas:\n{lista_falhas}",
                    )
                return

            pdfs_gerados = []
            for ordem, (docx_path, idx_sel) in enumerate(docx_paths, start=1):
                item_nr = (
                    self.nr_certificados[idx_sel]
                    if 0 <= idx_sel < len(self.nr_certificados)
                    else {}
                )
                nome_nr = self._obter_nome_arquivo_seguro(
                    item_nr.get("nome", ""),
                    f"nr_{ordem}",
                )
                pdf_destino = self._obter_arquivo_pdf_livre(
                    pasta_destino,
                    f"{nome_nr} {nome_func_arquivo}",
                    reutilizar_existente=True,
                )
                if self._converter_docx_para_pdf(docx_path, pdf_destino):
                    limpou_meta, _ = self._limpar_metadados_pdf(pdf_destino)
                    if not limpou_meta:
                        continue
                    self._registrar_documento_salvo(
                        pdf_destino,
                        origem="imprimir_nr",
                        empresa_id=funcionario.get("empresa_id"),
                        funcionario_id=funcionario.get("id"),
                        tipo_documento=item_nr.get("nome", ""),
                    )
                    pdfs_gerados.append(pdf_destino)

            if pdfs_gerados:
                outros_pdfs, falhas_outros = _gerar_pdf_outros_documentos()
                self._salvar_dados()
                _limpar_lista_imprimir()
                lista = "\n".join(str(p) for p in pdfs_gerados + outros_pdfs)
                if len(pdfs_gerados) == 1 and not outros_pdfs:
                    messagebox.showinfo("IMPRIMIR", f"PDF gerado com sucesso:\n{pdfs_gerados[0]}")
                else:
                    if falhas_outros:
                        lista_falhas = "\n".join(falhas_outros)
                        messagebox.showwarning(
                            "IMPRIMIR",
                            f"PDF(s) gerado(s):\n{lista}\n\nFalhas:\n{lista_falhas}",
                        )
                    else:
                        messagebox.showinfo("IMPRIMIR", f"PDF(s) gerado(s) com sucesso:\n{lista}")
                return

            pdf_destino = self._obter_arquivo_pdf_livre(
                pasta_destino,
                f"NR {nome_limpo}",
                reutilizar_existente=True,
            )
            try:
                self._escrever_pdf_texto(
                    pdf_destino,
                    f"Funcionario: {nome_func}",
                    linhas,
                )
                self._limpar_metadados_pdf(pdf_destino)
            except OSError:
                messagebox.showerror("IMPRIMIR", "Nao foi possivel gerar o PDF.")
                return
            self._registrar_documento_salvo(
                pdf_destino,
                origem="imprimir_nr",
                empresa_id=funcionario.get("empresa_id"),
                funcionario_id=funcionario.get("id"),
                tipo_documento="NR",
            )
        _, falhas_outros = _gerar_pdf_outros_documentos()
        self._salvar_dados()
        _limpar_lista_imprimir()
        detalhe_outros = ""
        if falhas_outros:
            detalhe_outros = "\n\nFalhas em OUTROS DOCUMENTOS:\n" + "\n".join(falhas_outros)
        messagebox.showwarning(
            "IMPRIMIR",
            "Word nao encontrado para conversao.\n"
            f"Foi gerado PDF simplificado:\n{pdf_destino}\n\n"
            f"Detalhe: {self._ultima_falha_conversao or 'Nao identificado.'}{detalhe_outros}",
        )

    def _aplicar_parametros_os_docx(
        self,
        origem_docx,
        destino_docx,
        funcionario,
        forcar_nome_uma_linha=False,
        caminho_documento_ref="",
        tipo_documento="",
        caminho_documento_saida="",
    ):
        campos = self._montar_campos_documento(funcionario or {})
        tipo_doc_norm = str(tipo_documento or "").strip().casefold()
        self._ultimo_qr_embutido_docx = False
        if tipo_doc_norm == "carteirinha":
            nr_vinculada = self._nr_vinculada_item_por_arquivo(caminho_documento_ref or origem_docx)
            if nr_vinculada is not None:
                campos.update(self._campos_data_da_nr_item(nr_vinculada))
        data35_txt = str(campos.get("DATA35", "") or "").strip()
        if tipo_doc_norm in {"anuencia", "anuência"} and not data35_txt:
            data35_txt = str(
                campos.get("DATANR", "")
                or campos.get("DATA_NR", "")
                or campos.get("DATA", "")
                or ""
            ).strip()
        funcionario2_normal = str(campos.get("FUNCIONARIO2", "") or "").strip()
        funcionario2_sem_quebra = re.sub(r"\s+", "\u00A0", funcionario2_normal) if funcionario2_normal else ""
        funcionario1_normal = str(campos.get("FUNCIONARIO1", "") or "").strip()
        funcionario1_sem_quebra = re.sub(r"\s+", "\u00A0", funcionario1_normal) if funcionario1_normal else ""
        valores = {
            "FUNCIONARIO1": str(campos.get("FUNCIONARIO1", "") or ""),
            "FUNCIONARIO2": funcionario2_sem_quebra or str(campos.get("FUNCIONARIO2", "") or ""),
            "FUNCIONARIO3": str(campos.get("FUNCIONARIO3", "") or ""),
            "FUNCIONARIO4": str(campos.get("FUNCIONARIO4", "") or ""),
            "CPF1": str(campos.get("CPF1", "") or ""),
            "CPF2": str(campos.get("CPF2", "") or ""),
            "CPF3": str(campos.get("CPF3", "") or ""),
            "CPF4": str(campos.get("CPF4", "") or ""),
            "PROFISSAO1": str(campos.get("PROFISSAO1", "") or ""),
            "ADMISSAO1": str(campos.get("ADMISSAO1", "") or ""),
            "ADMISSAO2": str(campos.get("ADMISSAO2", "") or ""),
            "ADMISSAO3": str(campos.get("ADMISSAO3", "") or ""),
            "ADMISSAO4": str(campos.get("ADMISSAO4", "") or ""),
            "COMECO1": str(campos.get("COMECO1", "") or ""),
            "FIM1": str(campos.get("FIM1", "") or ""),
            "FIM2": str(campos.get("FIM2", "") or ""),
            "DATA35": data35_txt,
            "DATA_35": data35_txt,
            "DATA 35": data35_txt,
            "EMPRESA1": str(campos.get("EMPRESA1", "") or ""),
            "EMPRESA2": str(campos.get("EMPRESA2", "") or ""),
            "EMPRESA3": str(campos.get("EMPRESA3", "") or ""),
            "EMPRESA4": str(campos.get("EMPRESA4", "") or ""),
            "EMPRESA5": str(campos.get("EMPRESA5", "") or ""),
            "CNPJ1": str(campos.get("CNPJ1", "") or ""),
            "CNPJ2": str(campos.get("CNPJ2", "") or ""),
            "CNPJ3": str(campos.get("CNPJ3", "") or ""),
            "CNPJ4": str(campos.get("CNPJ4", "") or ""),
            "CNPJ5": str(campos.get("CNPJ5", "") or ""),
            "ENDERECO1": str(campos.get("ENDERECO1", "") or ""),
            "ENDEREÇO1": str(campos.get("ENDERECO1", "") or ""),
            "CIDADE1": str(campos.get("CIDADE1", "") or ""),
            "FOTO1": "foto1",
            "LOGO1": "logo1",
            "LOGO2": "logo2",
            "LOGO3": "logo3",
            "IMG1": "logo1",
            "IMG2": "logo2",
            "IMG3": "logo3",
        }

        def _substituir_marcador_runs(texto_xml, marcador, token):
            marcador = str(marcador or "").strip()
            if not marcador:
                return texto_xml
            sep_runs = r"(?:\s|<[^>]+>)*"
            letras = [re.escape(ch) for ch in marcador]
            padrao_letras = sep_runs.join(letras)
            texto_xml = re.sub(
                rf"(?i)\[\s*{padrao_letras}\s*\]",
                token,
                texto_xml,
                flags=re.DOTALL,
            )
            texto_xml = re.sub(
                rf"(?i)(?<![A-Za-zÀ-ÿ0-9_]){padrao_letras}(?![A-Za-zÀ-ÿ0-9_])",
                token,
                texto_xml,
                flags=re.DOTALL,
            )
            return texto_xml

        try:
            with zipfile.ZipFile(origem_docx, "r") as zin, zipfile.ZipFile(destino_docx, "w") as zout:
                for info in zin.infolist():
                    data = zin.read(info.filename)
                    if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                        texto = data.decode("utf-8")
                        for marcador, valor in valores.items():
                            token = f"__CADNR_OS_{marcador}__"
                            valor_xml = self._escape_xml_text(valor)
                            texto = _substituir_marcador_runs(texto, marcador, token)
                            texto = re.sub(
                                rf"(?i)\[\s*{re.escape(marcador)}\s*\]",
                                valor_xml,
                                texto,
                            )
                            texto = re.sub(
                                rf"(?i)(?<![A-Za-zÀ-ÿ0-9_]){re.escape(marcador)}(?![A-Za-zÀ-ÿ0-9_])",
                                valor_xml,
                                texto,
                            )
                            texto = texto.replace(token, valor_xml)
                        # Reforco para evitar quebra de linha no nome completo da assinatura.
                        if funcionario2_normal and funcionario2_sem_quebra:
                            texto = texto.replace(
                                self._escape_xml_text(funcionario2_normal),
                                self._escape_xml_text(funcionario2_sem_quebra),
                            )
                        texto = texto.replace("[", "").replace("]", "")
                        padrao_data = (
                            r"(?:\d{1,2}\s+de\s+[A-Za-zÀ-ÿ]+\s+de\s+\d{4}|\d{2}/\d{2}/\d{4}|\d{2}/\d{4})"
                        )
                        texto = re.sub(
                            rf"__\s*CADNR\s*[_-]?\s*({padrao_data})\s*__",
                            r"\1",
                            texto,
                            flags=re.IGNORECASE,
                        )
                        texto = re.sub(
                            rf"__\s*({padrao_data})\s*__",
                            r"\1",
                            texto,
                            flags=re.IGNORECASE,
                        )
                        data = texto.encode("utf-8")
                    zout.writestr(info, data)
        except OSError:
            return False, "falha ao processar o arquivo OS"
        except zipfile.BadZipFile:
            return False, "arquivo OS invalido (esperado .docx)"
        except UnicodeDecodeError:
            return False, "falha ao ler conteudo do OS"

        if funcionario2_normal and funcionario2_sem_quebra:
            try:
                from docx import Document
                from docx.shared import Pt

                def _iter_paragrafos_tabelas(tabela):
                    for row in tabela.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                yield p
                            for tabela_interna in cell.tables:
                                yield from _iter_paragrafos_tabelas(tabela_interna)

                doc_nome = Document(str(destino_docx))
                paragrafos_alvo = list(doc_nome.paragraphs)
                for tabela in doc_nome.tables:
                    paragrafos_alvo.extend(list(_iter_paragrafos_tabelas(tabela)))
                for sec in doc_nome.sections:
                    paragrafos_alvo.extend(list(sec.header.paragraphs))
                    paragrafos_alvo.extend(list(sec.footer.paragraphs))
                    for tabela in sec.header.tables:
                        paragrafos_alvo.extend(list(_iter_paragrafos_tabelas(tabela)))
                    for tabela in sec.footer.tables:
                        paragrafos_alvo.extend(list(_iter_paragrafos_tabelas(tabela)))

                padrao_nome = r"\b" + r"\s+".join(
                    re.escape(parte) for parte in funcionario2_normal.split()
                ) + r"\b"
                alterou_nome = False
                paragrafos_nome = []

                def _extrair_tamanho_fonte_pt(paragrafo, padrao=11.0):
                    for run in paragrafo.runs:
                        tam = getattr(run.font, "size", None)
                        if tam is not None:
                            try:
                                return float(tam.pt)
                            except Exception:
                                pass
                    return float(padrao)

                def _largura_disponivel_pt(paragrafo, documento):
                    try:
                        parent = getattr(paragrafo, "_parent", None)
                        largura_emu = getattr(parent, "width", None)
                        if largura_emu:
                            largura_pt = float(largura_emu) / 12700.0
                        else:
                            sec = documento.sections[0] if documento.sections else None
                            if sec is None:
                                return 0.0
                            largura_emu = float(sec.page_width) - float(sec.left_margin) - float(sec.right_margin)
                            largura_pt = largura_emu / 12700.0
                        ind_esq = getattr(paragrafo.paragraph_format, "left_indent", None)
                        ind_dir = getattr(paragrafo.paragraph_format, "right_indent", None)
                        if ind_esq is not None:
                            largura_pt -= float(ind_esq) / 12700.0
                        if ind_dir is not None:
                            largura_pt -= float(ind_dir) / 12700.0
                        return max(0.0, largura_pt)
                    except Exception:
                        return 0.0

                def _largura_texto_estimada_pt(texto, fonte_pt):
                    txt = str(texto or "").replace("\u00A0", " ")
                    unidades = 0.0
                    for ch in txt:
                        if ch == " ":
                            unidades += 0.28
                        elif ch in "ilI.,:;!'|`":
                            unidades += 0.25
                        elif ch in "mwMWQO@#%&":
                            unidades += 0.90
                        elif ch.isupper():
                            unidades += 0.62
                        elif ch.isdigit():
                            unidades += 0.52
                        else:
                            unidades += 0.52
                    return unidades * float(fonte_pt)

                def _abreviar_nome_meio(nome_completo):
                    partes = [p for p in str(nome_completo or "").split() if p]
                    if len(partes) <= 2:
                        return str(nome_completo or "")
                    primeira = partes[0]
                    ultima = partes[-1]
                    meio = partes[1:-1]
                    meio_abreviado = []
                    conectores = {"de", "da", "do", "das", "dos"}
                    i = 0
                    while i < len(meio):
                        atual = str(meio[i] or "").strip()
                        atual_limpo = re.sub(r"[^A-Za-zÀ-ÿ]", "", atual)
                        if not atual_limpo:
                            i += 1
                            continue
                        atual_norm = atual_limpo.lower()
                        if (atual_norm in conectores) and (i + 1 < len(meio)):
                            prox = str(meio[i + 1] or "").strip()
                            prox_limpo = re.sub(r"[^A-Za-zÀ-ÿ]", "", prox)
                            if prox_limpo:
                                meio_abreviado.append(f"{atual_norm} {prox_limpo[0].upper()}")
                                i += 2
                                continue
                        if len(atual_limpo) >= 2:
                            meio_abreviado.append(atual_limpo[:2].title())
                        else:
                            meio_abreviado.append(atual_limpo[0].upper())
                        i += 1
                    if not meio_abreviado:
                        return f"{primeira} {ultima}".strip()
                    return f"{primeira} {' '.join(meio_abreviado)} {ultima}".strip()

                def _texto_paragrafo(paragrafo):
                    if paragrafo.runs:
                        return "".join(run.text or "" for run in paragrafo.runs)
                    return paragrafo.text or ""

                def _definir_texto_paragrafo(paragrafo, texto_novo):
                    if paragrafo.runs:
                        paragrafo.runs[0].text = texto_novo
                        for run in paragrafo.runs[1:]:
                            run.text = ""
                    else:
                        paragrafo.text = texto_novo

                def _definir_tamanho_fonte(paragrafo, tamanho_pt):
                    if not paragrafo.runs:
                        paragrafo.add_run(paragrafo.text or "")
                    for run in paragrafo.runs:
                        run.font.size = Pt(float(tamanho_pt))

                def _contem_nome_funcionario(texto):
                    texto_cmp = re.sub(r"\s+", " ", str(texto or "")).strip().casefold()
                    if not texto_cmp:
                        return False
                    candidatos = [
                        funcionario2_normal,
                        funcionario2_sem_quebra,
                        funcionario1_normal,
                        funcionario1_sem_quebra,
                    ]
                    for nome in candidatos:
                        nome_cmp = re.sub(r"\s+", " ", str(nome or "")).strip().casefold()
                        if nome_cmp and nome_cmp in texto_cmp:
                            return True
                    return False

                def _ajustar_nome_uma_linha(paragrafo):
                    texto_atual = _texto_paragrafo(paragrafo)
                    if not _contem_nome_funcionario(texto_atual):
                        return False

                    largura_disp = _largura_disponivel_pt(paragrafo, doc_nome)
                    if largura_disp <= 0:
                        return False

                    tamanho_inicial = min(_extrair_tamanho_fonte_pt(paragrafo), 12.0)
                    tamanho_min = 7.0
                    tamanho = tamanho_inicial
                    texto_teste = _texto_paragrafo(paragrafo)
                    while tamanho >= tamanho_min:
                        if _largura_texto_estimada_pt(texto_teste, tamanho) <= (largura_disp * 0.98):
                            _definir_tamanho_fonte(paragrafo, tamanho)
                            return True
                        tamanho -= 0.5

                    nome_abrev = _abreviar_nome_meio(funcionario2_normal)
                    nome_abrev_sem_quebra = re.sub(r"\s+", "\u00A0", nome_abrev).strip()
                    nome_abrev_upper_sem_quebra = nome_abrev_sem_quebra.upper()
                    texto_abrev = (
                        texto_teste.replace(funcionario2_sem_quebra, nome_abrev_sem_quebra)
                        .replace(funcionario2_normal, nome_abrev_sem_quebra)
                        .replace(funcionario1_sem_quebra, nome_abrev_upper_sem_quebra)
                        .replace(funcionario1_normal, nome_abrev_upper_sem_quebra)
                    )
                    if texto_abrev != texto_teste:
                        _definir_texto_paragrafo(paragrafo, texto_abrev)
                        texto_teste = texto_abrev

                    tamanho = tamanho_inicial
                    while tamanho >= tamanho_min:
                        if _largura_texto_estimada_pt(texto_teste, tamanho) <= (largura_disp * 0.98):
                            _definir_tamanho_fonte(paragrafo, tamanho)
                            return True
                        tamanho -= 0.5

                    _definir_tamanho_fonte(paragrafo, tamanho_min)
                    return True

                for p in paragrafos_alvo:
                    if not p.runs:
                        texto_atual = p.text or ""
                        texto_novo = re.sub(padrao_nome, funcionario2_sem_quebra, texto_atual, flags=re.IGNORECASE)
                        if texto_novo != texto_atual:
                            p.text = texto_novo
                            alterou_nome = True
                            paragrafos_nome.append(p)
                        elif forcar_nome_uma_linha and _contem_nome_funcionario(texto_atual):
                            paragrafos_nome.append(p)
                        continue
                    texto_runs = "".join(run.text or "" for run in p.runs)
                    texto_novo = re.sub(padrao_nome, funcionario2_sem_quebra, texto_runs, flags=re.IGNORECASE)
                    if texto_novo != texto_runs:
                        p.runs[0].text = texto_novo
                        for run in p.runs[1:]:
                            run.text = ""
                        alterou_nome = True
                        paragrafos_nome.append(p)
                    elif forcar_nome_uma_linha and _contem_nome_funcionario(texto_runs):
                        paragrafos_nome.append(p)
                if forcar_nome_uma_linha and paragrafos_nome:
                    for p in paragrafos_nome:
                        if _ajustar_nome_uma_linha(p):
                            alterou_nome = True
                if alterou_nome:
                    doc_nome.save(str(destino_docx))
            except Exception:
                pass

        foto1_caminho = self._resolver_foto_funcionario(campos.get("FOTO1", ""))
        logo_img1_caminho = self._resolver_logo_empresa(campos.get("LOGO1", ""))
        logo_img2_caminho = self._resolver_logo_empresa(campos.get("LOGO2", ""))
        logo_img3_caminho = self._resolver_logo_empresa(campos.get("LOGO3", ""))
        incluir_logo_em_cabecalho_rodape = tipo_doc_norm not in {"anuencia", "anuência"}
        if foto1_caminho is not None and foto1_caminho.exists():
            inseriu_foto, motivo_foto = self._inserir_imagem_por_marcador_docx(
                destino_docx,
                foto1_caminho,
                ["foto1"],
                largura_cm=2.5,
                altura_cm=3.2,
            )
            if not inseriu_foto:
                motivo_txt = str(motivo_foto or "").strip().lower()
                if "nao encontrado no documento" not in motivo_txt:
                    return False, f"falha ao inserir FOTO1: {motivo_foto}"
        if logo_img1_caminho is not None and logo_img1_caminho.exists():
            inseriu_logo, motivo_logo = self._inserir_img1_logo_em_tabela_docx(
                destino_docx,
                logo_img1_caminho,
                incluir_cabecalho_rodape=incluir_logo_em_cabecalho_rodape,
            )
            if (
                not inseriu_logo
                and not incluir_logo_em_cabecalho_rodape
                and "nao encontrado no documento" in str(motivo_logo or "").strip().lower()
            ):
                # Fallback da anuencia: se nao encontrou marcador no corpo,
                # tenta header/footer para manter compatibilidade com modelos antigos.
                inseriu_logo, motivo_logo = self._inserir_img1_logo_em_tabela_docx(
                    destino_docx,
                    logo_img1_caminho,
                    incluir_cabecalho_rodape=True,
                )
            if not inseriu_logo:
                motivo_txt = str(motivo_logo or "").strip().lower()
                if (
                    "nao encontrado no documento" not in motivo_txt
                    and "marcador img1/img nao encontrado" not in motivo_txt
                ):
                    return False, f"falha ao inserir IMG1: {motivo_logo}"
        if logo_img2_caminho is not None and logo_img2_caminho.exists():
            inseriu_logo2, motivo_logo2 = self._inserir_img2_logo_em_tabela_docx(
                destino_docx,
                logo_img2_caminho,
                incluir_cabecalho_rodape=incluir_logo_em_cabecalho_rodape,
            )
            if not inseriu_logo2:
                motivo_txt = str(motivo_logo2 or "").strip().lower()
                if "nao encontrado no documento" not in motivo_txt:
                    return False, f"falha ao inserir LOGO2: {motivo_logo2}"
        if logo_img3_caminho is not None and logo_img3_caminho.exists():
            inseriu_logo3, motivo_logo3 = self._inserir_img3_logo_em_tabela_docx(
                destino_docx,
                logo_img3_caminho,
                incluir_cabecalho_rodape=incluir_logo_em_cabecalho_rodape,
            )
            if not inseriu_logo3:
                motivo_txt = str(motivo_logo3 or "").strip().lower()
                if "nao encontrado no documento" not in motivo_txt:
                    return False, f"falha ao inserir LOGO3: {motivo_logo3}"

        if tipo_doc_norm == "carteirinha":
            caminho_saida_ref = str(caminho_documento_saida or "").strip()
            if caminho_saida_ref:
                try:
                    payload_qr = self._montar_payload_qrcode_documento(
                        caminho_saida_ref,
                        permitir_arquivo_inexistente=True,
                    )
                    if payload_qr:
                        import qrcode

                        tmp_qr = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                        tmp_qr.close()
                        caminho_tmp_qr = Path(tmp_qr.name)
                        try:
                            qr = qrcode.QRCode(
                                version=None,
                                error_correction=qrcode.constants.ERROR_CORRECT_M,
                                box_size=10,
                                border=4,
                            )
                            qr.add_data(payload_qr)
                            qr.make(fit=True)
                            img = qr.make_image()
                            img.save(str(caminho_tmp_qr))
                            self._ultimo_qr_embutido_docx = bool(
                                self._inserir_qrcode_em_carteirinha_docx(
                                    destino_docx,
                                    caminho_tmp_qr,
                                    tamanho_cm=1.6,
                                    margem_esquerda_cm=0.1,
                                )
                            )
                        finally:
                            try:
                                caminho_tmp_qr.unlink(missing_ok=True)
                            except OSError:
                                pass
                except Exception:
                    self._ultimo_qr_embutido_docx = False

        return True, ""

    def _iterar_arquivos_documentos_empresa(self, empresa):
        if not isinstance(empresa, dict):
            return
        base_dir = Path(__file__).resolve().parent
        pastas_vistas = set()
        for pasta_nome in self._pastas_candidatas_empresa(empresa):
            if not pasta_nome:
                continue
            pasta_empresa = (base_dir / pasta_nome).resolve()
            if not pasta_empresa.exists() or not pasta_empresa.is_dir():
                continue
            chave_pasta = str(pasta_empresa).casefold()
            if chave_pasta in pastas_vistas:
                continue
            pastas_vistas.add(chave_pasta)
            try:
                arquivos = sorted(pasta_empresa.iterdir(), key=lambda p: p.name.lower())
            except OSError:
                continue
            for arquivo in arquivos:
                if not arquivo.is_file():
                    continue
                if arquivo.suffix.lower() not in {".docx", ".doc", ".pdf"}:
                    continue
                yield base_dir, arquivo

    @staticmethod
    def _detectar_tipo_outro_documento(nome_arquivo):
        nome_norm = App._normalizar_texto_filtro(nome_arquivo)
        if not nome_norm:
            return ""
        if "carteirinha" in nome_norm:
            return "Carteirinha"
        if "anuencia" in nome_norm:
            return "Anuencia"
        if "fit test" in nome_norm or "fittest" in nome_norm or re.search(r"(^| )fit( |$)", nome_norm):
            return "Fit Test"
        if re.search(r"(^| )aso( |$)", nome_norm):
            return "ASO"
        if (
            "ficha epi" in nome_norm
            or (("ficha" in nome_norm) and ("epi" in nome_norm))
            or re.search(r"(^| )epi( |$)", nome_norm)
        ):
            return "Ficha de EPI"
        if "contrato" in nome_norm:
            return "Contrato"
        if (
            "ordem de servico" in nome_norm
            or re.search(r"(^| )os( |$)", nome_norm)
            or ("os " in nome_norm)
        ):
            return "Ordem de Servico"
        return ""

    def _assinatura_documentos_projeto(self):
        entradas = []
        vistos = set()
        for empresa in self.empresas:
            if not isinstance(empresa, dict):
                continue
            for base_dir, arquivo in self._iterar_arquivos_documentos_empresa(empresa):
                try:
                    rel = str(arquivo.relative_to(base_dir)).replace("\\", "/")
                    chave = rel.casefold()
                    if chave in vistos:
                        continue
                    vistos.add(chave)
                    stat = arquivo.stat()
                    entradas.append(f"{chave}|{int(stat.st_mtime)}|{int(stat.st_size)}")
                except Exception:
                    continue
        entradas.sort()
        return ";".join(entradas)

    def _reaplicar_filtros_documentos(self):
        empresa_id = self._empresa_id_selecionada_main()
        if empresa_id is not None:
            self._aplicar_filtro_nr_por_empresa(empresa_id, limpar_nr=False)
        else:
            self._render_campos_nr()
            self._atualizar_lista_nr_imprimir()
        self._carregar_outros_documentos_empresa_selecionada()
        self._atualizar_lista_outros_docs_imprimir()

    def _monitorar_documentos_projeto(self):
        try:
            assinatura = self._assinatura_documentos_projeto()
            if assinatura != self._docs_monitor_assinatura:
                self._docs_monitor_assinatura = assinatura
                self._reaplicar_filtros_documentos()
        except Exception:
            pass
        try:
            self._docs_monitor_after_id = self.after(
                int(self._docs_monitor_interval_ms),
                self._monitorar_documentos_projeto,
            )
        except Exception:
            self._docs_monitor_after_id = None

    def _construir_aba_outros_documentos(self):
        container = ttk.Frame(self.aba_outros_documentos)
        container.pack(fill="both", expand=True)
        container.columnconfigure(0, weight=1)
        container.rowconfigure(0, weight=1)

        self.outros_docs_checklist_frame = ttk.Frame(container)
        self.outros_docs_checklist_frame.grid(row=0, column=0, sticky="nsew")
        self.outros_docs_checklist_frame.columnconfigure(0, weight=1)

        f_botoes = ttk.Frame(container)
        f_botoes.grid(row=1, column=0, sticky="e", pady=(12, 0))
        ttk.Button(
            f_botoes,
            text="ADICIONAR",
            command=self._adicionar_outros_documentos_imprimir,
        ).grid(row=0, column=0)

    def _listar_arquivos_outros_documentos(self, documento):
        if not isinstance(documento, dict):
            return []
        arquivos = documento.get("arquivos", {})
        if not isinstance(arquivos, dict):
            return []
        itens = []
        for tipo in OUTROS_DOCUMENTOS_TIPOS:
            valor = arquivos.get(tipo, "")
            if isinstance(valor, list):
                for caminho_item in valor:
                    caminho_txt = str(caminho_item or "").strip()
                    if caminho_txt:
                        itens.append((tipo, caminho_txt))
            else:
                caminho_txt = str(valor or "").strip()
                if caminho_txt:
                    itens.append((tipo, caminho_txt))
        return itens

    def _coletar_outros_documentos_disponiveis(self, empresa_id_sel, funcionario_sel=None):
        if empresa_id_sel is None:
            return []
        empresa = next((e for e in self.empresas if e.get("id") == empresa_id_sel), None)
        documento = next(
            (
                item for item in self.documentos
                if isinstance(item, dict) and item.get("empresa_id") == empresa_id_sel
            ),
            None,
        )
        itens_documentos = list(self._listar_arquivos_outros_documentos(documento))
        caminhos_vistos = set()
        for _tipo_item, caminho_item in itens_documentos:
            caminho_norm = self._normalizar_caminho_documento_db(caminho_item)
            chave = str(caminho_norm or caminho_item or "").strip().replace("\\", "/").casefold()
            if chave:
                caminhos_vistos.add(chave)
        if empresa:
            for base_dir, arquivo in self._iterar_arquivos_documentos_empresa(empresa):
                tipo_auto = self._detectar_tipo_outro_documento(arquivo.stem)
                if not tipo_auto:
                    continue
                try:
                    caminho_rel = str(arquivo.relative_to(base_dir)).replace("\\", "/")
                except ValueError:
                    continue
                chave_arquivo = str(caminho_rel).strip().replace("\\", "/").casefold()
                if not chave_arquivo or chave_arquivo in caminhos_vistos:
                    continue
                itens_documentos.append((tipo_auto, caminho_rel))
                caminhos_vistos.add(chave_arquivo)
        funcao_selecionada = str((funcionario_sel or {}).get("funcao", "") or "").strip()
        nr35_selecionada = self._nr_35_selecionada()
        disponiveis = []
        for tipo, caminho in itens_documentos:
            tipo_norm = str(tipo).strip().casefold()
            if tipo_norm in {"anuencia", "anuência"} and not nr35_selecionada:
                continue
            if tipo_norm == "carteirinha":
                if not self._arquivo_vinculado_nr_carteirinha(caminho):
                    continue
            else:
                if tipo_norm in {"ordem de servico", "ordem de serviço"}:
                    if not self._arquivo_vinculado_nr_documento(caminho):
                        continue
                    if not self._arquivo_ordem_servico_compativel_funcao(caminho, funcao_selecionada):
                        continue
                elif tipo_norm == "ficha de epi":
                    if not self._arquivo_epi_compativel_funcao(caminho, funcao_selecionada):
                        continue
                else:
                    if not self._arquivo_vinculado_nr_documento(caminho):
                        continue
                    if not self._arquivo_compativel_funcao(caminho, funcao_selecionada):
                        continue
            disponiveis.append(
                {
                    "empresa_id": empresa_id_sel,
                    "tipo": tipo,
                    "caminho": caminho,
                }
            )
        return disponiveis

    @staticmethod
    def _normalizar_texto_filtro(texto):
        txt = unicodedata.normalize("NFD", str(texto or ""))
        txt = "".join(ch for ch in txt if unicodedata.category(ch) != "Mn")
        txt = txt.lower()
        txt = re.sub(r"[^a-z0-9]+", " ", txt)
        return re.sub(r"\s+", " ", txt).strip()

    def _arquivo_compativel_funcao(self, caminho_arquivo, funcao_selecionada):
        funcao_norm = self._normalizar_texto_filtro(funcao_selecionada)
        if not funcao_norm:
            return True
        nome_norm = self._normalizar_texto_filtro(Path(str(caminho_arquivo or "")).stem)
        if not nome_norm:
            return False
        if funcao_norm in nome_norm:
            return True
        funcoes_norm = [
            self._normalizar_texto_filtro(nome_funcao)
            for nome_funcao in FUNCOES_CBO.keys()
        ]
        if any(fn and fn in nome_norm for fn in funcoes_norm):
            return False
        return True

    def _arquivo_ordem_servico_compativel_funcao(self, caminho_arquivo, funcao_selecionada):
        funcao_norm = self._normalizar_texto_filtro(funcao_selecionada)
        if not funcao_norm:
            return False
        nome_norm = self._normalizar_texto_filtro(Path(str(caminho_arquivo or "")).stem)
        if not nome_norm:
            return False
        if funcao_norm in nome_norm:
            return True

        tokens_funcao = [tok for tok in funcao_norm.split() if len(tok) >= 4]
        if tokens_funcao and all(tok in nome_norm for tok in tokens_funcao):
            return True
        return False

    def _arquivo_epi_compativel_funcao(self, caminho_arquivo, funcao_selecionada):
        funcao_norm = self._normalizar_texto_filtro(funcao_selecionada)
        if not funcao_norm:
            return False
        nome_norm = self._normalizar_texto_filtro(Path(str(caminho_arquivo or "")).stem)
        if not nome_norm:
            return False
        if "epi" not in nome_norm:
            return False
        if funcao_norm in nome_norm:
            return True
        tokens_funcao = [tok for tok in funcao_norm.split() if len(tok) >= 4]
        if tokens_funcao and all(tok in nome_norm for tok in tokens_funcao):
            return True
        return False

    def _arquivo_vinculado_funcao(self, caminho_arquivo, funcao_selecionada):
        funcao_norm = self._normalizar_texto_filtro(funcao_selecionada)
        if not funcao_norm:
            return False
        nome_norm = self._normalizar_texto_filtro(Path(str(caminho_arquivo or "")).stem)
        if not nome_norm:
            return False
        return funcao_norm in nome_norm

    def _nrs_marcadas_cadnr(self):
        selecionadas = []
        for item in self.nr_certificados:
            if not isinstance(item, dict):
                continue
            marcada = bool(item.get("imprimir", False) or item.get("imprimir_adicionado", False))
            if not marcada:
                data_de = str(item.get("coluna_1", "") or "").strip()
                data_ate = str(item.get("coluna_2", "") or "").strip()
                # Fallback: se a NR recebeu periodo, tratamos como selecionada
                # para habilitar a carteirinha vinculada.
                marcada = bool(data_de or data_ate)
            if marcada:
                nome_nr = str(item.get("nome", "") or "").strip()
                if nome_nr:
                    selecionadas.append(nome_nr)
        return selecionadas

    def _nr_35_selecionada(self):
        for item in self.nr_certificados:
            if not isinstance(item, dict):
                continue
            nome_norm = self._normalizar_nome_nr(item.get("nome", ""))
            if not str(nome_norm).startswith("nr35"):
                continue
            if bool(item.get("imprimir", False) or item.get("imprimir_adicionado", False)):
                return True
        return False

    def _dados_nr_para_carteirinha(self, nome_nr):
        nome_norm = self._normalizar_texto_filtro(nome_nr)
        if not nome_norm:
            return None, set(), set()
        m_num = re.search(r"\bnr\s*(\d+)\b", nome_norm)
        nr_num = m_num.group(1) if m_num else None
        codigos = set()
        if re.search(r"\bemp\b", nome_norm) or re.search(r"\bempilhadeira\b", nome_norm):
            codigos.update({"emp", "empilhadeira"})
        if re.search(r"\bpr\b", nome_norm):
            codigos.update({"pr", "ponterolante", "ponte rolante"})
        if re.search(r"\bmunck\b", nome_norm):
            codigos.update({"munck"})
        if re.search(r"\bgarra\b", nome_norm):
            codigos.update({"garra"})
        if re.search(r"\bpta\b", nome_norm) or re.search(r"\bplataforma\b", nome_norm):
            codigos.update({"pta", "plataforma"})
        if re.search(r"\bpemt\b", nome_norm):
            codigos.update({"pemt"})
        if re.search(r"\bguin\b", nome_norm) or re.search(r"\bguindaste\b", nome_norm):
            codigos.update({"guin", "guindaste"})
        if re.search(r"\band\b", nome_norm) or "andaime" in nome_norm:
            codigos.update({"and", "andaime"})
        cargas = set()
        if m_num:
            trecho_pos_nr = nome_norm[m_num.end() :]
            for n in re.findall(r"\b\d{1,3}\b", trecho_pos_nr):
                n_limpo = str(int(n))
                if n_limpo:
                    cargas.add(n_limpo)
        return nr_num, codigos, cargas

    def _arquivo_vinculado_nr_carteirinha(self, caminho_arquivo):
        nome_arquivo = self._normalizar_texto_filtro(Path(str(caminho_arquivo or "")).stem)
        if not nome_arquivo:
            return False
        nrs_marcadas = self._nrs_marcadas_cadnr()
        if not nrs_marcadas:
            # Sem NR marcada, ainda exibimos/reconhecemos carteirinhas disponiveis.
            return True

        nr_arq, codigos_arq, cargas_arq = self._dados_nr_para_carteirinha(nome_arquivo)
        if not nr_arq:
            return False

        for nr in nrs_marcadas:
            nr_num, codigos, cargas = self._dados_nr_para_carteirinha(nr)
            if not nr_num:
                continue
            if str(nr_arq) != str(nr_num):
                continue

            # Se a NR marcada exige carga horaria, o arquivo deve conter a mesma carga.
            if cargas and not (cargas_arq & cargas):
                # Fallback pragmatico: para subtipos de equipamento (ex.: munck/pr/emp),
                # se o subtipo bate, aceitamos mesmo com carga diferente no nome.
                if not (codigos and (codigos_arq & codigos)):
                    continue
            # Se a NR marcada exige subtipo/codigo, o arquivo deve conter o mesmo subtipo.
            if codigos and not (codigos_arq & codigos):
                continue
            # Para NR generica, exige arquivo generico (sem subtipo e sem carga extra),
            # evitando vincular carteirinha especializada por engano.
            if (not codigos and not cargas) and (codigos_arq or cargas_arq):
                continue
            return True
        return False

    def _arquivo_vinculado_nr_documento(self, caminho_arquivo):
        nome_arquivo = self._normalizar_texto_filtro(Path(str(caminho_arquivo or "")).stem)
        if not nome_arquivo:
            return False
        if "nr" not in nome_arquivo:
            # Documento sem NR no nome nao depende da NR selecionada.
            return True

        nrs_marcadas = self._nrs_marcadas_cadnr()
        if not nrs_marcadas:
            return False

        nr_arq, codigos_arq, cargas_arq = self._dados_nr_para_carteirinha(nome_arquivo)
        if not nr_arq:
            return False
        for nr in nrs_marcadas:
            nr_num, codigos, cargas = self._dados_nr_para_carteirinha(nr)
            if not nr_num:
                continue
            if str(nr_arq) != str(nr_num):
                continue
            if cargas and not (cargas_arq & cargas):
                continue
            if codigos and not (codigos_arq & codigos):
                continue
            if (not codigos and not cargas) and (codigos_arq or cargas_arq):
                continue
            return True
        return False

    def _nr_vinculada_item_por_arquivo(self, caminho_arquivo):
        nome_arquivo = self._normalizar_texto_filtro(Path(str(caminho_arquivo or "")).stem)
        if not nome_arquivo:
            return None
        nr_arq, codigos_arq, cargas_arq = self._dados_nr_para_carteirinha(nome_arquivo)
        if not nr_arq:
            return None

        melhor_item = None
        melhor_score = -1
        for item in self.nr_certificados:
            if not isinstance(item, dict):
                continue
            marcada = bool(item.get("imprimir", False) or item.get("imprimir_adicionado", False))
            if not marcada:
                data_de = str(item.get("coluna_1", "") or "").strip()
                data_ate = str(item.get("coluna_2", "") or "").strip()
                marcada = bool(data_de or data_ate)
            if not marcada:
                continue

            nome_nr = str(item.get("nome", "") or "").strip()
            nr_num, codigos, cargas = self._dados_nr_para_carteirinha(nome_nr)
            if not nr_num or str(nr_arq) != str(nr_num):
                continue
            if cargas and not (cargas_arq & cargas):
                continue
            if codigos and not (codigos_arq & codigos):
                continue
            if (not codigos and not cargas) and (codigos_arq or cargas_arq):
                continue

            score = 0
            if codigos and (codigos_arq & codigos):
                score += 3
            if cargas and (cargas_arq & cargas):
                score += 3
            if str(item.get("coluna_2", "") or "").strip():
                score += 1
            if str(item.get("coluna_1", "") or "").strip():
                score += 1

            if score > melhor_score:
                melhor_score = score
                melhor_item = item
        return melhor_item

    def _campos_data_da_nr_item(self, item):
        if not isinstance(item, dict):
            return {}

        texto_de = str(item.get("coluna_1", "") or "").strip()
        texto_ate = str(item.get("coluna_2", "") or "").strip()
        data_de = self._parse_data_br(texto_de)
        data_ate = self._parse_data_br(texto_ate)
        data_ref = data_ate or data_de
        if data_ref is None:
            return {}

        data_de_ref = data_de or data_ref
        data_ate_ref = data_ate or data_ref
        data_ref_extenso = self._formatar_data_extenso_br(data_ref)
        data_de_extenso = self._formatar_data_extenso_br(data_de_ref)
        data_ate_extenso = self._formatar_data_extenso_br(data_ate_ref)
        data_ate_extenso_dd = (
            f"{data_ate_ref.day:02d} de {self._formatar_data_extenso_br(data_ate_ref).split(' de ', 1)[1]}"
        )
        data_de_mes_ano = self._formatar_data_dd_mes_extenso_br(data_de_ref)
        if data_de and data_ate:
            de_e_ate_extenso = (
                f"{self._formatar_data_extenso_br_sem_ano(data_de)} "
                f"e {self._formatar_data_extenso_br(data_ate)}"
            )
            de_a_ate_extenso = (
                f"{self._formatar_data_extenso_br_sem_ano(data_de)} "
                f"a {self._formatar_data_extenso_br(data_ate)}"
            )
        else:
            de_e_ate_extenso = self._formatar_data_extenso_br(data_ref)
            de_a_ate_extenso = self._formatar_data_extenso_br(data_ref)

        data_nr_txt = texto_ate or texto_de or self._formatar_data_br(data_ref)
        return {
            "COMECO1": data_de_mes_ano,
            "FIM1": data_ate_extenso_dd,
            "FIM2": data_ate_extenso_dd,
            "DATA": data_ref_extenso,
            "DATA_EXTENSO": data_ref_extenso,
            "DATA_DE_EXTENSO": data_de_extenso,
            "DATA_ATE_EXTENSO": data_ate_extenso,
            "DATANR": data_nr_txt,
            "DATA_NR": data_nr_txt,
            "DE_E_ATE_EXTENSO": de_e_ate_extenso,
            "DE_A_ATE_EXTENSO": de_a_ate_extenso,
        }

    def _atualizar_outros_documentos_empresas(self, preservar_id=None):
        self._carregar_outros_documentos_empresa_selecionada()

    def _carregar_outros_documentos_empresa_selecionada(self):
        if not hasattr(self, "outros_docs_checklist_frame"):
            return
        for widget in self.outros_docs_checklist_frame.winfo_children():
            widget.destroy()
        self.outros_docs_check_vars = {}
        self.outros_docs_check_paths = {}

        empresa_id_sel = self._empresa_id_selecionada_main()
        if empresa_id_sel is None:
            ttk.Label(
                self.outros_docs_checklist_frame,
                text="Selecione a empresa na aba CADNR para visualizar os documentos.",
            ).grid(row=0, column=0, sticky="w")
            return

        funcionario_sel = next(
            (
                f for f in self.funcionarios
                if isinstance(f, dict) and f.get("id") == self._funcionario_id_selecionado_main()
            ),
            None,
        )
        funcao_selecionada = str((funcionario_sel or {}).get("funcao", "") or "").strip()
        row = 0
        disponiveis = self._coletar_outros_documentos_disponiveis(empresa_id_sel, funcionario_sel)
        for idx_item, item_doc in enumerate(disponiveis):
            tipo = str(item_doc.get("tipo", "") or "").strip()
            caminho = str(item_doc.get("caminho", "") or "").strip()
            var = tk.BooleanVar(value=False)
            chave_item = f"{tipo}|{idx_item}"
            self.outros_docs_check_vars[chave_item] = var
            self.outros_docs_check_paths[chave_item] = {
                "tipo": tipo,
                "caminho": caminho,
            }
            nome_exibicao = Path(caminho).name or caminho
            ttk.Checkbutton(
                self.outros_docs_checklist_frame,
                text=f"{tipo} - {nome_exibicao}",
                variable=var,
            ).grid(row=row, column=0, sticky="w", pady=2)
            row += 1

        if row == 0:
            msg = "Nao ha documentos cadastrados para esta empresa."
            if funcao_selecionada:
                msg = f"Nao ha documentos cadastrados para a funcao '{funcao_selecionada}'."
            ttk.Label(
                self.outros_docs_checklist_frame,
                text=msg,
            ).grid(row=0, column=0, sticky="w")

    def _adicionar_outros_documentos_imprimir(self):
        empresa_id_sel = self._empresa_id_selecionada_main()
        if empresa_id_sel is None:
            messagebox.showwarning("OUTROS DOCUMENTOS", "Selecione uma empresa na aba CADNR.")
            return
        funcionario_sel = next(
            (
                f for f in self.funcionarios
                if isinstance(f, dict) and f.get("id") == self._funcionario_id_selecionado_main()
            ),
            None,
        )
        if not funcionario_sel:
            messagebox.showwarning("OUTROS DOCUMENTOS", "Selecione um funcionario na aba CADNR.")
            return
        funcao_selecionada = str(funcionario_sel.get("funcao", "") or "").strip()
        if not funcao_selecionada:
            messagebox.showwarning(
                "OUTROS DOCUMENTOS",
                "O funcionario selecionado nao possui funcao cadastrada.",
            )
            return
        selecionados = []
        for chave_item, var in self.outros_docs_check_vars.items():
            if bool(var.get()):
                info_doc = self.outros_docs_check_paths.get(chave_item, {})
                if not isinstance(info_doc, dict):
                    continue
                tipo = str(info_doc.get("tipo", "") or "").strip()
                caminho = str(info_doc.get("caminho", "") or "").strip()
                if caminho:
                    selecionados.append((tipo, caminho))
        if not selecionados:
            messagebox.showwarning("OUTROS DOCUMENTOS", "Selecione ao menos um documento.")
            return

        selecionados_vinculados = []
        selecionados_sem_vinculo = []
        for tipo, caminho in selecionados:
            tipo_norm = str(tipo or "").strip().casefold()
            if tipo_norm == "carteirinha":
                vinculado = self._arquivo_vinculado_nr_carteirinha(caminho)
            else:
                if tipo_norm in {"ordem de servico", "ordem de serviço"}:
                    if not self._arquivo_vinculado_nr_documento(caminho):
                        selecionados_sem_vinculo.append((tipo, caminho))
                        continue
                    vinculado = self._arquivo_ordem_servico_compativel_funcao(caminho, funcao_selecionada)
                elif tipo_norm == "ficha de epi":
                    vinculado = self._arquivo_epi_compativel_funcao(caminho, funcao_selecionada)
                else:
                    if not self._arquivo_vinculado_nr_documento(caminho):
                        selecionados_sem_vinculo.append((tipo, caminho))
                        continue
                    vinculado = self._arquivo_compativel_funcao(caminho, funcao_selecionada)
            if vinculado:
                selecionados_vinculados.append((tipo, caminho))
            else:
                selecionados_sem_vinculo.append((tipo, caminho))

        if not selecionados_vinculados:
            messagebox.showwarning(
                "OUTROS DOCUMENTOS",
                f"Nenhum documento selecionado possui vinculo com a funcao '{funcao_selecionada}'.",
            )
            return

        adicionados = 0
        for tipo, caminho in selecionados_vinculados:
            duplicado = any(
                item.get("empresa_id") == empresa_id_sel
                and item.get("tipo") == tipo
                and item.get("caminho") == caminho
                for item in self.outros_docs_imprimir
            )
            if duplicado:
                continue
            self.outros_docs_imprimir.append(
                {
                    "empresa_id": empresa_id_sel,
                    "tipo": tipo,
                    "caminho": caminho,
                }
            )
            adicionados += 1
        self._atualizar_lista_outros_docs_imprimir()
        if adicionados > 0 and selecionados_sem_vinculo:
            messagebox.showwarning(
                "OUTROS DOCUMENTOS",
                f"{adicionados} documento(s) adicionado(s). "
                f"{len(selecionados_sem_vinculo)} sem vinculo com a funcao '{funcao_selecionada}' foram ignorados.",
            )
        elif adicionados > 0:
            messagebox.showinfo("OUTROS DOCUMENTOS", f"{adicionados} documento(s) adicionados na aba IMPRIMIR.")
        else:
            messagebox.showinfo("OUTROS DOCUMENTOS", "Os documentos selecionados ja estavam na aba IMPRIMIR.")

    def _atualizar_lista_outros_docs_imprimir(self):
        if not hasattr(self, "lista_outros_docs_imprimir"):
            return
        self.lista_outros_docs_imprimir.delete(0, tk.END)
        for item in self.outros_docs_imprimir:
            tipo = str(item.get("tipo", "") or "").strip()
            caminho = str(item.get("caminho", "") or "").strip()
            nome_arquivo = Path(caminho).name or caminho
            self.lista_outros_docs_imprimir.insert(tk.END, f"{tipo} - {nome_arquivo}")

    def _excluir_outro_documento_imprimir_selecionado(self):
        if not hasattr(self, "lista_outros_docs_imprimir"):
            return
        selecao = self.lista_outros_docs_imprimir.curselection()
        if not selecao:
            messagebox.showwarning("IMPRIMIR", "Selecione um documento para excluir.")
            return
        idx = int(selecao[0])
        if idx < 0 or idx >= len(self.outros_docs_imprimir):
            return
        del self.outros_docs_imprimir[idx]
        self._atualizar_lista_outros_docs_imprimir()

    def _on_empresa_imprimir_selected(self, _event=None):
        empresa_id = self._empresa_id_selecionada_imprimir()
        self._atualizar_imprimir_funcionarios(empresa_id)
        self._sincronizar_empresa_entre_abas(empresa_id, origem="imprimir")

    def _on_funcionario_imprimir_selected(self, _event=None):
        self._sincronizar_funcionario_entre_abas(self._funcionario_id_selecionado_imprimir(), origem="imprimir")
        self._atualizar_preview_nr()
        self._atualizar_lista_nr_imprimir()

    def _atualizar_imprimir_empresas(self, preservar_id=None):
        self.imprimir_empresa_ids = [None] + [empresa["id"] for empresa in self.empresas]
        self.imprimir_select_empresa["values"] = [""] + [
            self._empresa_label(empresa) for empresa in self.empresas
        ]

        if preservar_id in self.imprimir_empresa_ids:
            idx = self.imprimir_empresa_ids.index(preservar_id)
        else:
            idx = 0

        self.imprimir_select_empresa.current(idx)
        self._atualizar_imprimir_funcionarios(self.imprimir_empresa_ids[idx])

    def _atualizar_imprimir_funcionarios(self, empresa_id):
        if empresa_id is None:
            self.imprimir_funcionario_ids = [None]
            self.imprimir_select_funcionario["values"] = [""]
            self.imprimir_select_funcionario.current(0)
            self._atualizar_preview_nr()
            return

        funcs = [f for f in self.funcionarios if f["empresa_id"] == empresa_id]
        self.imprimir_funcionario_ids = [None] + [f["id"] for f in funcs]
        self.imprimir_select_funcionario["values"] = [""] + [f["nome"] for f in funcs]
        if funcs:
            self.imprimir_select_funcionario.current(1)
        else:
            self.imprimir_select_funcionario.current(0)
        self._atualizar_preview_nr()

    def _excluir_empresa_imprimir(self):
        empresa_id = self._empresa_id_selecionada_imprimir()
        if empresa_id is None:
            messagebox.showwarning("Exclusao", "Selecione uma empresa para excluir.")
            return

        empresa = next((e for e in self.empresas if e["id"] == empresa_id), None)
        if empresa is None:
            messagebox.showerror("Exclusao", "Empresa nao encontrada.")
            return

        qtd_func = sum(1 for f in self.funcionarios if f["empresa_id"] == empresa_id)
        msg = f"Deseja excluir a empresa '{empresa['nome']}'?"
        if qtd_func:
            msg += f"\n{qtd_func} funcionario(s) vinculado(s) tambem sera(ao) excluido(s)."
        if not messagebox.askyesno("Confirmar exclusao", msg):
            return

        self.empresas = [e for e in self.empresas if e["id"] != empresa_id]
        self.funcionarios = [f for f in self.funcionarios if f["empresa_id"] != empresa_id]
        self._salvar_dados()
        self._atualizar_select_empresas()
        messagebox.showinfo("Exclusao", "Cadastro de empresa excluido com sucesso.")

    def _excluir_funcionario_imprimir(self):
        funcionario_id = self._funcionario_id_selecionado_imprimir()
        if funcionario_id is None:
            messagebox.showwarning("Exclusao", "Selecione um funcionario para excluir.")
            return

        funcionario = next((f for f in self.funcionarios if f["id"] == funcionario_id), None)
        if funcionario is None:
            messagebox.showerror("Exclusao", "Funcionario nao encontrado.")
            return

        if not messagebox.askyesno(
            "Confirmar exclusao",
            f"Deseja excluir o funcionario '{funcionario['nome']}'?",
        ):
            return

        self.funcionarios = [f for f in self.funcionarios if f["id"] != funcionario_id]
        self._salvar_dados()
        self._atualizar_select_empresas(preservar_id=funcionario["empresa_id"])
        messagebox.showinfo("Exclusao", "Cadastro de funcionario excluido com sucesso.")

    def _abrir_edicao_empresa_imprimir(self):
        empresa_id = self._empresa_id_selecionada_imprimir()
        if empresa_id is None:
            messagebox.showwarning("Edicao", "Selecione uma empresa para editar.")
            return

        empresa = next((e for e in self.empresas if e["id"] == empresa_id), None)
        if empresa is None:
            messagebox.showerror("Edicao", "Empresa nao encontrada.")
            return

        popup = CadastroPopup(self, "Editar Empresa")
        popup.columnconfigure(0, weight=1)

        r = 0
        popup.secao("Empresa", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Escolher Empresa:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nome_pasta = ttk.Entry(f)
        nome_pasta.grid(row=0, column=1, sticky="ew")
        nome_pasta.insert(0, empresa.get("nome_pasta", ""))

        r += 1
        popup.secao("Insira uma Logo", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Logo:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        logo_salva = self._caminho_logo_empresa(empresa)
        logo_var = tk.StringVar(value=logo_salva or str(empresa.get("logo", "") or ""))
        ttk.Entry(f, textvariable=logo_var, state="readonly").grid(row=0, column=1, sticky="ew")
        preview_logo = ttk.Label(f, text="Pre-visualizacao indisponivel")
        preview_logo.grid(row=1, column=1, sticky="w", pady=(6, 0))
        logo_origem = {"path": ""}
        self._atualizar_preview_logo(preview_logo, logo_salva or str(empresa.get("logo", "") or ""))

        def escolher_logo():
            caminho = filedialog.askopenfilename(
                title="Selecionar logo da empresa",
                filetypes=[
                    ("Imagens", "*.png *.jpg *.jpeg *.bmp *.gif *.webp"),
                    ("Todos os arquivos", "*.*"),
                ],
            )
            if not caminho:
                return
            logo_origem["path"] = caminho
            logo_var.set(caminho)
            self._atualizar_preview_logo(preview_logo, caminho)

        ttk.Button(f, text="Procurar...", command=escolher_logo).grid(row=0, column=2, padx=(8, 0))

        r += 1
        popup.secao("Nome da Empresa", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Razao Social:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nome_empresa = ttk.Entry(f)
        nome_empresa.grid(row=0, column=1, sticky="ew")
        nome_empresa.insert(0, empresa.get("nome", ""))

        r += 1
        popup.secao("CNPJ", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="CNPJ:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        cnpj = ttk.Entry(f)
        cnpj.grid(row=0, column=1, sticky="ew")
        cnpj.insert(0, empresa.get("cnpj", ""))
        self._bind_mask(cnpj, mascara_cnpj)

        r += 1
        popup.secao("Email", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="email:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        email = ttk.Entry(f)
        email.grid(row=0, column=1, sticky="ew")
        email.insert(0, empresa.get("email", ""))

        r += 1
        popup.secao("nº do celular", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Contato:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        contato = ttk.Entry(f)
        contato.grid(row=0, column=1, sticky="ew")
        contato.insert(0, empresa.get("contato", ""))
        self._bind_mask(contato, mascara_celular)

        r += 1
        popup.secao("Endereco da Empresa", r)
        r += 1
        tipo_log, nome_log = self._linha_logradouro(popup, r)
        tipo_log.set(empresa.get("logradouro_tipo", ""))
        nome_log.insert(0, empresa.get("logradouro_nome", ""))

        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12, pady=2)
        f.columnconfigure(1, weight=1)
        f.columnconfigure(3, weight=1)
        ttk.Label(f, text="nº:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        numero = ttk.Entry(f)
        numero.grid(row=0, column=1, sticky="ew")
        numero.insert(0, empresa.get("numero", ""))
        ttk.Label(f, text="Complemento:").grid(row=0, column=2, sticky="w", padx=(14, 6))
        complemento = ttk.Entry(f)
        complemento.grid(row=0, column=3, sticky="ew")
        complemento.insert(0, empresa.get("complemento", ""))

        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12, pady=2)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Bairro:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        bairro = ttk.Entry(f)
        bairro.grid(row=0, column=1, sticky="ew")
        bairro.insert(0, empresa.get("bairro", ""))

        r += 1
        uf, cidade = self._combo_uf_cidade(popup, r)
        uf_inicial = empresa.get("uf", "")
        cidade_inicial = empresa.get("cidade", "")
        if uf_inicial:
            uf.set(uf_inicial)
            uf.event_generate("<<ComboboxSelected>>")
            if cidade_inicial:
                cidade.set(cidade_inicial)

        r += 1
        rodape = ttk.Frame(popup, height=56)
        rodape.grid(row=r, column=0, sticky="ew", padx=12)
        rodape.grid_propagate(False)
        botoes = ttk.Frame(popup)
        botoes.place(relx=1.0, rely=1.0, anchor="se", x=-12, y=-12)

        def salvar():
            novo_nome = nome_empresa.get().strip()
            if not novo_nome:
                messagebox.showerror("Validacao", "Preencha a Razao Social.")
                return
            nome_pasta_txt = nome_pasta.get().strip()
            logo_atual = str(empresa.get("logo", "") or "").strip()
            logo_novo = logo_atual
            if logo_origem["path"]:
                try:
                    logo_novo = self._salvar_logo_empresa(
                        logo_origem["path"],
                        empresa_id,
                        nome_pasta_txt or novo_nome,
                    )
                except OSError as exc:
                    messagebox.showerror("Validacao", f"Nao foi possivel salvar a logo: {exc}")
                    return
            empresa["nome_pasta"] = nome_pasta_txt
            empresa["logo"] = logo_novo
            empresa["nome"] = novo_nome
            empresa["cnpj"] = cnpj.get().strip()
            empresa["email"] = email.get().strip()
            empresa["contato"] = contato.get().strip()
            empresa["logradouro_tipo"] = tipo_log.get().strip()
            empresa["logradouro_nome"] = nome_log.get().strip()
            empresa["numero"] = numero.get().strip()
            empresa["complemento"] = complemento.get().strip()
            empresa["bairro"] = bairro.get().strip()
            empresa["uf"] = uf.get().strip()
            empresa["cidade"] = cidade.get().strip()
            if logo_origem["path"] and logo_atual and logo_atual != logo_novo:
                self._remover_logo_empresa(logo_atual)
            self._salvar_dados()
            self._atualizar_select_empresas(preservar_id=empresa_id)
            messagebox.showinfo("Edicao", "Empresa atualizada com sucesso.")
            popup.destroy()

        ttk.Button(botoes, text="Cancelar", command=popup.destroy).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(botoes, text="Salvar", command=salvar).grid(row=0, column=1)

        popup.ajustar_tamanho()
        nome_empresa.focus_set()

    def _abrir_edicao_funcionario_imprimir(self):
        funcionario_id = self._funcionario_id_selecionado_imprimir()
        if funcionario_id is None:
            messagebox.showwarning("Edicao", "Selecione um funcionario para editar.")
            return

        funcionario = next((f for f in self.funcionarios if f["id"] == funcionario_id), None)
        if funcionario is None:
            messagebox.showerror("Edicao", "Funcionario nao encontrado.")
            return

        popup = CadastroPopup(self, "Editar Funcionario")
        popup.columnconfigure(0, weight=1)
        form_parent = popup.habilitar_rolagem()
        form_parent.columnconfigure(0, weight=1)

        r = 0
        popup.secao("Nome da Pasta", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Nome da Pasta:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nome_pasta_valores = []
        for empresa_ref in self.empresas:
            valor_pasta = str(empresa_ref.get("nome_pasta") or empresa_ref.get("nome") or "").strip()
            if valor_pasta and valor_pasta not in nome_pasta_valores:
                nome_pasta_valores.append(valor_pasta)
        nome_pasta = ttk.Combobox(f, state="readonly", values=nome_pasta_valores)
        nome_pasta.grid(row=0, column=1, sticky="ew")
        nome_pasta_atual = str(funcionario.get("nome_pasta") or "").strip()
        if nome_pasta_atual:
            nome_pasta.set(nome_pasta_atual)
        else:
            empresa_ref = next((e for e in self.empresas if e["id"] == funcionario.get("empresa_id")), None)
            if empresa_ref:
                nome_pasta_fallback = str(empresa_ref.get("nome_pasta") or empresa_ref.get("nome") or "").strip()
                if nome_pasta_fallback:
                    nome_pasta.set(nome_pasta_fallback)

        r += 1
        popup.secao("Nome do Funcionario", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Nome:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nome = ttk.Entry(f)
        nome.grid(row=0, column=1, sticky="ew")
        foto_salva = self._resolver_foto_funcionario(funcionario.get("foto", ""))
        ttk.Label(f, text="Foto:").grid(row=1, column=0, sticky="w", padx=(0, 8), pady=(6, 0))
        foto_var = tk.StringVar(value=str(foto_salva) if foto_salva else str(funcionario.get("foto", "") or ""))
        ttk.Entry(f, textvariable=foto_var, state="readonly").grid(row=1, column=1, sticky="ew", pady=(6, 0))
        preview_foto = ttk.Label(f, text="Pre-visualizacao indisponivel")
        preview_foto.grid(row=2, column=1, sticky="w", pady=(6, 0))
        foto_origem = {"path": ""}
        self._atualizar_preview_logo(
            preview_foto,
            str(foto_salva) if foto_salva else str(funcionario.get("foto", "") or ""),
        )

        def escolher_foto_funcionario():
            caminho = filedialog.askopenfilename(
                title="Selecionar foto do funcionario",
                filetypes=[
                    ("Imagens", "*.png *.jpg *.jpeg *.bmp *.gif *.webp"),
                    ("Todos os arquivos", "*.*"),
                ],
            )
            if caminho:
                foto_origem["path"] = caminho
                foto_var.set(caminho)
                self._atualizar_preview_logo(preview_foto, caminho)

        ttk.Button(f, text="Procurar...", command=escolher_foto_funcionario).grid(
            row=1, column=2, padx=(8, 0), pady=(6, 0)
        )
        nome.insert(0, funcionario.get("nome", ""))

        r += 1
        popup.secao("Insira o nº do CPF", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="CPF:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        cpf = ttk.Entry(f)
        cpf.grid(row=0, column=1, sticky="ew")
        cpf.insert(0, funcionario.get("cpf", ""))
        self._bind_mask(cpf, mascara_cpf)

        r += 1
        popup.secao("Insira o nº do RG", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="RG:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        rg = ttk.Entry(f)
        rg.grid(row=0, column=1, sticky="ew")
        rg.insert(0, funcionario.get("rg", ""))

        r += 1
        popup.secao("Insira a data de nascimento", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Nascimento:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nascimento = ttk.Entry(f)
        nascimento.grid(row=0, column=1, sticky="ew")
        self._set_date_placeholder(nascimento)
        if funcionario.get("nascimento"):
            nascimento.delete(0, tk.END)
            nascimento.insert(0, funcionario.get("nascimento", ""))
        self._bind_mask(nascimento, mascara_data)
        self._attach_calendar(f, nascimento)

        r += 1
        popup.secao("nº do celular", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Contato:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        celular = ttk.Entry(f)
        celular.grid(row=0, column=1, sticky="ew")
        celular.insert(0, funcionario.get("contato", ""))
        self._bind_mask(celular, mascara_celular)

        r += 1
        popup.secao("Endereco do Funcionario", r, parent=form_parent)
        r += 1
        tipo_log, nome_log = self._linha_logradouro(form_parent, r)
        tipo_log.set(funcionario.get("logradouro_tipo", ""))
        nome_log.insert(0, funcionario.get("logradouro_nome", ""))

        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12, pady=2)
        f.columnconfigure(1, weight=1)
        f.columnconfigure(3, weight=1)
        ttk.Label(f, text="nº:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        numero = ttk.Entry(f)
        numero.grid(row=0, column=1, sticky="ew")
        numero.insert(0, funcionario.get("numero", ""))
        ttk.Label(f, text="Complemento:").grid(row=0, column=2, sticky="w", padx=(14, 6))
        complemento = ttk.Entry(f)
        complemento.grid(row=0, column=3, sticky="ew")
        complemento.insert(0, funcionario.get("complemento", ""))

        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12, pady=2)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Bairro:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        bairro = ttk.Entry(f)
        bairro.grid(row=0, column=1, sticky="ew")
        bairro.insert(0, funcionario.get("bairro", ""))

        r += 1
        uf, cidade = self._combo_uf_cidade(form_parent, r)
        uf_inicial = funcionario.get("uf", "")
        cidade_inicial = funcionario.get("cidade", "")
        if uf_inicial:
            uf.set(uf_inicial)
            uf.event_generate("<<ComboboxSelected>>")
            if cidade_inicial:
                cidade.set(cidade_inicial)

        r += 1
        popup.secao("Selecione a data de admissão", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Admissão:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        admissao = ttk.Entry(f)
        admissao.grid(row=0, column=1, sticky="ew")
        self._set_date_placeholder(admissao)
        if funcionario.get("admissao"):
            admissao.delete(0, tk.END)
            admissao.insert(0, funcionario.get("admissao", ""))
        self._bind_mask(admissao, mascara_data)
        self._attach_calendar(f, admissao)

        r += 1
        popup.secao("Selecione o cargo do funcionario", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Função:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        funcao = ttk.Combobox(f, state="readonly", values=sorted(FUNCOES_CBO.keys()))
        funcao.grid(row=0, column=1, sticky="ew")
        funcao_inicial = funcionario.get("funcao", "")
        if funcao_inicial in FUNCOES_CBO:
            funcao.set(funcao_inicial)

        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="CBO:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        cbo_var = tk.StringVar(value=funcionario.get("cbo", ""))
        cbo = ttk.Entry(f, textvariable=cbo_var, state="readonly")
        cbo.grid(row=0, column=1, sticky="ew")

        def atualizar_cbo(_event=None):
            cbo_var.set(FUNCOES_CBO.get(funcao.get(), ""))

        funcao.bind("<<ComboboxSelected>>", atualizar_cbo)
        if funcao.get() and not cbo_var.get():
            cbo_var.set(FUNCOES_CBO.get(funcao.get(), ""))

        r += 1
        popup.secao("Insira o valor do salário", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Salário:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        salario = ttk.Entry(f)
        salario.grid(row=0, column=1, sticky="ew")
        salario_val = funcionario.get("salario", "")
        salario.insert(0, salario_val if salario_val else "R$ 0,00")
        self._bind_mask(salario, mascara_moeda_br)

        r += 1
        botoes = ttk.Frame(form_parent)
        botoes.grid(row=r, column=0, sticky="e", padx=12, pady=18)

        def salvar():
            novo_nome = nome.get().strip()
            nome_pasta_txt = nome_pasta.get().strip()
            if not novo_nome:
                messagebox.showerror("Validacao", "Preencha o nome do funcionario.")
                return
            if not nome_pasta_txt:
                messagebox.showerror("Validacao", "Selecione a empresa.")
                return
            if len(re.sub(r"\D", "", cpf.get())) != 11:
                messagebox.showerror("Validacao", "CPF invalido.")
                return

            empresa_ref = next(
                (
                    e
                    for e in self.empresas
                    if str(e.get("nome_pasta") or e.get("nome") or "").strip() == nome_pasta_txt
                ),
                None,
            )
            if not empresa_ref:
                messagebox.showerror("Validacao", "Empresa nao encontrada para o item selecionado.")
                return
            novo_empresa_id = empresa_ref["id"]
            foto_atual = str(funcionario.get("foto", "") or "").strip()
            foto_nova = foto_atual
            if foto_origem["path"]:
                try:
                    foto_nova = self._salvar_foto_funcionario(
                        foto_origem["path"],
                        funcionario["id"],
                        novo_nome,
                    )
                except OSError as exc:
                    messagebox.showerror("Validacao", f"Nao foi possivel salvar a foto: {exc}")
                    return
            funcionario["nome"] = novo_nome
            funcionario["nome_pasta"] = nome_pasta_txt
            funcionario["empresa_id"] = novo_empresa_id
            funcionario["foto"] = foto_nova
            funcionario["cpf"] = cpf.get().strip()
            funcionario["rg"] = rg.get().strip()
            funcionario["nascimento"] = nascimento.get().strip()
            funcionario["contato"] = celular.get().strip()
            funcionario["logradouro_tipo"] = tipo_log.get().strip()
            funcionario["logradouro_nome"] = nome_log.get().strip()
            funcionario["numero"] = numero.get().strip()
            funcionario["complemento"] = complemento.get().strip()
            funcionario["bairro"] = bairro.get().strip()
            funcionario["uf"] = uf.get().strip()
            funcionario["cidade"] = cidade.get().strip()
            funcionario["admissao"] = admissao.get().strip()
            funcionario["funcao"] = funcao.get().strip()
            funcionario["cbo"] = cbo_var.get().strip()
            funcionario["salario"] = salario.get().strip()
            if foto_origem["path"] and foto_atual and foto_atual != foto_nova:
                self._remover_foto_funcionario(foto_atual)
            self._salvar_dados()
            self._atualizar_select_empresas(preservar_id=novo_empresa_id)
            messagebox.showinfo("Edicao", "Funcionario atualizado com sucesso.")
            popup.destroy()

        ttk.Button(botoes, text="Cancelar", command=popup.destroy).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(botoes, text="Salvar", command=salvar).grid(row=0, column=1)

        popup.ajustar_tamanho()
        nome.focus_set()

    def _carregar_dados(self):
        if not self.dados_path.exists():
            return

        try:
            conteudo = self.dados_path.read_text(encoding="utf-8")
            dados = json.loads(conteudo)
        except (OSError, json.JSONDecodeError, UnicodeDecodeError):
            return

        empresas = dados.get("empresas", [])
        funcionarios = dados.get("funcionarios", [])
        documentos = dados.get("documentos", [])
        documentos_salvos = dados.get("documentos_salvos", [])
        assinatura_digital = dados.get("assinatura_digital", {})
        github_config = dados.get("github_config", {})
        if not isinstance(empresas, list) or not isinstance(funcionarios, list):
            return

        empresas_validas = []
        for e in empresas:
            if not isinstance(e, dict):
                continue
            eid = e.get("id")
            nome = e.get("nome")
            if isinstance(eid, int) and isinstance(nome, str) and nome.strip():
                empresas_validas.append(
                    {
                        "id": eid,
                        "nome_pasta": str(e.get("nome_pasta", "") or ""),
                        "logo": str(e.get("logo", "") or ""),
                        "nome": nome.strip(),
                        "cnpj": str(e.get("cnpj", "") or ""),
                        "email": str(e.get("email", "") or ""),
                        "contato": str(e.get("contato", "") or ""),
                        "logradouro_tipo": str(e.get("logradouro_tipo", "") or ""),
                        "logradouro_nome": str(e.get("logradouro_nome", "") or ""),
                        "numero": str(e.get("numero", "") or ""),
                        "complemento": str(e.get("complemento", "") or ""),
                        "bairro": str(e.get("bairro", "") or ""),
                        "uf": str(e.get("uf", "") or ""),
                        "cidade": str(e.get("cidade", "") or ""),
                    }
                )

        ids_empresas = {e["id"] for e in empresas_validas}
        funcionarios_validos = []
        for f in funcionarios:
            if not isinstance(f, dict):
                continue
            fid = f.get("id")
            nome = f.get("nome")
            empresa_id = f.get("empresa_id")
            if (
                isinstance(fid, int)
                and isinstance(nome, str)
                and nome.strip()
                and isinstance(empresa_id, int)
                and empresa_id in ids_empresas
            ):
                funcionarios_validos.append(
                    {
                        "id": fid,
                        "nome": nome.strip(),
                        "nome_pasta": str(f.get("nome_pasta", "") or ""),
                        "empresa_id": empresa_id,
                        "foto": str(f.get("foto", "") or ""),
                        "cpf": str(f.get("cpf", "") or ""),
                        "rg": str(f.get("rg", "") or ""),
                        "nascimento": str(f.get("nascimento", "") or ""),
                        "contato": str(f.get("contato", "") or ""),
                        "logradouro_tipo": str(f.get("logradouro_tipo", "") or ""),
                        "logradouro_nome": str(f.get("logradouro_nome", "") or ""),
                        "numero": str(f.get("numero", "") or ""),
                        "complemento": str(f.get("complemento", "") or ""),
                        "bairro": str(f.get("bairro", "") or ""),
                        "uf": str(f.get("uf", "") or ""),
                        "cidade": str(f.get("cidade", "") or ""),
                        "admissao": str(f.get("admissao", "") or ""),
                        "funcao": str(f.get("funcao", "") or ""),
                        "cbo": str(f.get("cbo", "") or ""),
                        "salario": str(f.get("salario", "") or ""),
                    }
                )

        self.empresas = sorted(empresas_validas, key=lambda x: x["id"])
        self.funcionarios = sorted(funcionarios_validos, key=lambda x: x["id"])
        documentos_validos = []
        if isinstance(documentos, list):
            for d in documentos:
                if not isinstance(d, dict):
                    continue
                did = d.get("id")
                empresa_id = d.get("empresa_id")
                tipos = d.get("tipos", [])
                arquivos = d.get("arquivos", {})
                certificado_arquivo = str(d.get("certificado_arquivo", "") or "").strip()
                certificado_dias_raw = d.get("certificado_dias", 1)
                try:
                    certificado_dias = max(1, int(certificado_dias_raw))
                except (TypeError, ValueError):
                    certificado_dias = 1
                if (
                    isinstance(did, int)
                    and isinstance(empresa_id, int)
                    and empresa_id in ids_empresas
                    and isinstance(tipos, list)
                ):
                    tipos_validos = [
                        str(t).strip()
                        for t in tipos
                        if (
                            isinstance(t, str)
                            and str(t).strip()
                            and str(t).strip() in OUTROS_DOCUMENTOS_TIPOS
                        )
                    ]
                    arquivos_validos = {}
                    if isinstance(arquivos, dict):
                        for k, v in arquivos.items():
                            nome_tipo = str(k or "").strip()
                            if not nome_tipo or nome_tipo not in OUTROS_DOCUMENTOS_TIPOS:
                                continue
                            if isinstance(v, list):
                                caminhos_tipo = [
                                    str(item or "").strip()
                                    for item in v
                                    if str(item or "").strip()
                                ]
                                if caminhos_tipo:
                                    arquivos_validos[nome_tipo] = caminhos_tipo
                            else:
                                caminho_tipo = str(v or "").strip()
                                if caminho_tipo:
                                    arquivos_validos[nome_tipo] = caminho_tipo
                    # Compatibilidade com registros antigos que tinham "Certificado" em tipos.
                    if not certificado_arquivo and any(t.casefold() == "certificado" for t in tipos_validos):
                        certificado_arquivo = ""
                    documentos_validos.append(
                        {
                            "id": did,
                            "empresa_id": empresa_id,
                            "certificado_arquivo": certificado_arquivo,
                            "certificado_dias": certificado_dias,
                            "tipos": [t for t in tipos_validos if t.casefold() != "certificado"],
                            "arquivos": arquivos_validos,
                        }
                    )
        self.documentos = sorted(documentos_validos, key=lambda x: x["id"])
        documentos_salvos_validos = []
        if isinstance(documentos_salvos, list):
            for item in documentos_salvos:
                if not isinstance(item, dict):
                    continue
                caminho = self._normalizar_caminho_documento_db(item.get("caminho", ""))
                if not caminho:
                    continue
                qrcode_caminho = self._normalizar_caminho_documento_db(item.get("qrcode", ""))
                documentos_salvos_validos.append(
                    {
                        "caminho": caminho,
                        "qrcode": qrcode_caminho,
                        "origem": str(item.get("origem", "") or "").strip() or "sistema",
                        "empresa_id": item.get("empresa_id")
                        if isinstance(item.get("empresa_id"), int) else None,
                        "funcionario_id": item.get("funcionario_id")
                        if isinstance(item.get("funcionario_id"), int) else None,
                        "tipo_documento": str(item.get("tipo_documento", "") or "").strip(),
                        "data_ultima_gravacao": str(
                            item.get("data_ultima_gravacao", "") or ""
                        ).strip(),
                    }
                )
        self.documentos_salvos = documentos_salvos_validos
        self._sincronizar_vinculo_funcionarios_empresas()
        nr_certificados = dados.get("nr_certificados", None)
        if isinstance(nr_certificados, list):
            nr_validos = []
            for item in nr_certificados:
                if not isinstance(item, dict):
                    continue
                nome = str(item.get("nome", "") or "").strip()
                if not nome:
                    continue
                nome = self._nome_nr_canonico(nome)
                if nome.strip().upper() == "NR 20":
                    nome = "NR 20 (08)"
                reciclagem_raw = item.get("reciclagem", False)
                if isinstance(reciclagem_raw, bool):
                    reciclagem = reciclagem_raw
                else:
                    reciclagem = str(reciclagem_raw).strip().lower() in {
                        "1",
                        "true",
                        "sim",
                        "yes",
                        "on",
                    }
                imprimir_raw = item.get("imprimir", False)
                if isinstance(imprimir_raw, bool):
                    imprimir = imprimir_raw
                else:
                    imprimir = str(imprimir_raw).strip().lower() in {
                        "1",
                        "true",
                        "sim",
                        "yes",
                        "on",
                    }
                nr_validos.append(
                    {
                        "nome": nome,
                        "coluna_1": str(item.get("coluna_1", "") or ""),
                        "coluna_2": str(item.get("coluna_2", "") or ""),
                        "dias": self._duracao_nr_item(item),
                        "reciclagem": reciclagem,
                        "imprimir": imprimir,
                        "imprimir_adicionado": bool(item.get("imprimir_adicionado", False)),
                    }
                )
            nr_unicos = {}
            for item in nr_validos:
                chave = self._nome_nr_canonico(item.get("nome", "")).strip().lower()
                if not chave:
                    continue
                if chave not in nr_unicos:
                    nr_unicos[chave] = item
                    nr_unicos[chave]["nome"] = self._nome_nr_canonico(item.get("nome", ""))
                else:
                    nr_unicos[chave]["reciclagem"] = bool(
                        nr_unicos[chave].get("reciclagem", False) or item.get("reciclagem", False)
                    )
                    nr_unicos[chave]["imprimir"] = bool(
                        nr_unicos[chave].get("imprimir", False) or item.get("imprimir", False)
                    )
                    nr_unicos[chave]["imprimir_adicionado"] = bool(
                        nr_unicos[chave].get("imprimir_adicionado", False)
                        or item.get("imprimir_adicionado", False)
                    )
            nr_validos = list(nr_unicos.values())
            nomes_existentes = {item["nome"].strip().lower() for item in nr_validos}
            for item_padrao in self._nr_certificados_padrao():
                nome_padrao = item_padrao["nome"].strip().lower()
                if nome_padrao not in nomes_existentes:
                    nr_validos.append(item_padrao.copy())
            self.nr_certificados = sorted(nr_validos, key=self._chave_ordenacao_nr)
            for item in self.nr_certificados:
                item["coluna_1"] = ""
                item["coluna_2"] = ""
                item["dias"] = self._duracao_nr_item(item)
        self.prox_empresa_id = max((e["id"] for e in self.empresas), default=0) + 1
        self.prox_funcionario_id = (
            max((f["id"] for f in self.funcionarios), default=0) + 1
        )
        self.prox_documento_id = max((d["id"] for d in self.documentos), default=0) + 1
        if isinstance(assinatura_digital, dict):
            self.assinatura_digital_habilitada = bool(assinatura_digital.get("habilitada", False))
            self.assinatura_digital_pfx = str(assinatura_digital.get("pfx", "") or "").strip()
            self.assinatura_digital_senha = str(assinatura_digital.get("senha", "") or "")
            self.assinatura_digital_pfx1 = str(assinatura_digital.get("pfx1", "") or "").strip()
            self.assinatura_digital_senha1 = str(assinatura_digital.get("senha1", "") or "")
            self.assinatura_digital_img1 = str(assinatura_digital.get("img1", "") or "").strip()
            self.assinatura_digital_pfx2 = str(
                assinatura_digital.get("pfx2", "") or self.assinatura_digital_pfx or ""
            ).strip()
            self.assinatura_digital_senha2 = str(
                assinatura_digital.get("senha2", "") or self.assinatura_digital_senha or ""
            )
            certs = assinatura_digital.get("certificados", [])
            certs_validos = []
            if isinstance(certs, list):
                for cert in certs:
                    if not isinstance(cert, dict):
                        continue
                    pfx_txt = str(cert.get("pfx", "") or "").strip()
                    if not pfx_txt:
                        continue
                    certs_validos.append(
                        {
                            "nome": str(cert.get("nome", "") or Path(pfx_txt).stem).strip() or Path(pfx_txt).stem,
                            "pfx": pfx_txt,
                            "senha": str(cert.get("senha", "") or ""),
                        }
                    )
            self.assinatura_digital_certificados = certs_validos
        if isinstance(github_config, dict):
            repo_txt = str(github_config.get("repo", "") or "").strip()
            repo_norm = self._normalizar_repo_github(repo_txt)
            self.github_repo = repo_norm or "Elizangela2805/documentos"
            self.github_branch = str(github_config.get("branch", "") or "").strip() or "main"
            self.github_dir = str(github_config.get("dir", "") or "").strip().strip("/") or "_pdf_gerados"
            self.github_pages_base = self._normalizar_pages_base(github_config.get("pages_base", "")) or "https://elizangela2805.github.io/documentos"
            self.github_token = str(github_config.get("token", "") or "").strip()
        self._aplicar_configuracao_github_ambiente()

    def _salvar_dados(self):
        self._sincronizar_vinculo_funcionarios_empresas()
        self._garantir_pastas_empresas()
        nr_certificados = []
        if self.nr_certificados_widgets:
            for idx, nr in enumerate(self.nr_certificados):
                item = (
                    self.nr_certificados_widgets[idx]
                    if idx < len(self.nr_certificados_widgets)
                    else None
                )
                nome = str(nr.get("nome", "") or "").strip()
                if not nome:
                    nome = f"NR {idx + 1}"
                if item:
                    coluna_1 = item["coluna_1"].get().strip() if item.get("coluna_1") is not None else ""
                    coluna_2 = item["coluna_2"].get().strip() if item.get("coluna_2") is not None else ""
                    if item.get("dias") is not None:
                        try:
                            dias = max(1, int(item["dias"].get().strip() or "1"))
                        except (TypeError, ValueError):
                            dias = self._duracao_nr_item(nr)
                    else:
                        dias = self._duracao_nr_item(nr)
                    reciclagem = bool(item["reciclagem_var"].get())
                else:
                    coluna_1 = str(nr.get("coluna_1", "") or "")
                    coluna_2 = str(nr.get("coluna_2", "") or "")
                    dias = self._duracao_nr_item(nr)
                    reciclagem = bool(nr.get("reciclagem", False))
                nr_certificados.append(
                    {
                        "nome": nome,
                        "coluna_1": coluna_1,
                        "coluna_2": coluna_2,
                        "dias": dias,
                        "reciclagem": reciclagem,
                        "imprimir": bool(
                            nr.get("imprimir", False)
                        )
                        if idx < len(self.nr_certificados) else False,
                        "imprimir_adicionado": bool(
                            nr.get("imprimir_adicionado", False)
                        )
                        if idx < len(self.nr_certificados) else False,
                    }
                )
            self.nr_certificados = nr_certificados
        else:
            for idx, item in enumerate(self.nr_certificados):
                nome = str(item.get("nome", "") or "").strip()
                if not nome:
                    nome = f"NR {idx + 1}"
                nr_certificados.append(
                    {
                        "nome": nome,
                        "coluna_1": str(item.get("coluna_1", "") or ""),
                        "coluna_2": str(item.get("coluna_2", "") or ""),
                        "dias": self._duracao_nr_item(item),
                        "reciclagem": bool(item.get("reciclagem", False)),
                        "imprimir": bool(item.get("imprimir", False)),
                        "imprimir_adicionado": bool(item.get("imprimir_adicionado", False)),
                    }
                )

        dados = {
            "empresas": self.empresas,
            "funcionarios": self.funcionarios,
            "documentos": self.documentos,
            "documentos_salvos": self.documentos_salvos,
            "nr_certificados": nr_certificados,
            "assinatura_digital": {
                "habilitada": bool(self.assinatura_digital_habilitada),
                "pfx": str(self.assinatura_digital_pfx or "").strip(),
                "senha": str(self.assinatura_digital_senha or ""),
                "pfx1": str(self.assinatura_digital_pfx1 or "").strip(),
                "senha1": str(self.assinatura_digital_senha1 or ""),
                "img1": str(self.assinatura_digital_img1 or "").strip(),
                "pfx2": str(self.assinatura_digital_pfx2 or "").strip(),
                "senha2": str(self.assinatura_digital_senha2 or ""),
                "certificados": self.assinatura_digital_certificados,
            },
            "github_config": {
                "repo": str(self.github_repo or "").strip() or "Elizangela2805/documentos",
                "branch": str(self.github_branch or "").strip() or "main",
                "dir": str(self.github_dir or "").strip().strip("/") or "_pdf_gerados",
                "pages_base": str(self.github_pages_base or "").strip().rstrip("/")
                or "https://elizangela2805.github.io/documentos",
                "token": str(self.github_token or "").strip(),
            },
        }
        try:
            self.dados_path.write_text(
                json.dumps(dados, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        except OSError:
            messagebox.showerror(
                "Erro",
                "Nao foi possivel salvar os dados em cadnr_dados.json.",
            )

    def _buscar_cidades_por_uf(self, uf_sigla):
        if not uf_sigla:
            return []

        if uf_sigla in self.cidades_cache:
            return self.cidades_cache[uf_sigla]

        url = (
            f"https://servicodados.ibge.gov.br/api/v1/localidades/estados/"
            f"{uf_sigla}/municipios?orderBy=nome"
        )
        try:
            req = request.Request(
                url,
                headers={
                    "Accept-Encoding": "identity",
                    "User-Agent": "CADNR/1.0",
                },
            )
            with request.urlopen(req, timeout=8) as response:
                payload = response.read()
                content_encoding = response.headers.get("Content-Encoding", "").lower()
                if "gzip" in content_encoding:
                    payload = gzip.decompress(payload)

                try:
                    dados = json.loads(payload.decode("utf-8"))
                except UnicodeDecodeError:
                    # Alguns proxies ignoram Accept-Encoding e retornam gzip sem header.
                    dados = json.loads(gzip.decompress(payload).decode("utf-8"))

            cidades = [item.get("nome", "").strip() for item in dados if item.get("nome")]
            if not cidades:
                cidades = UF_CIDADES.get(uf_sigla, [])
        except (error.URLError, TimeoutError, json.JSONDecodeError, ValueError):
            cidades = UF_CIDADES.get(uf_sigla, [])

        self.cidades_cache[uf_sigla] = cidades
        return cidades

    def _combo_uf_cidade(self, parent, row):
        frame = ttk.Frame(parent)
        frame.grid(row=row, column=0, sticky="ew", padx=12, pady=2)
        frame.columnconfigure(1, weight=1)
        frame.columnconfigure(3, weight=1)

        ttk.Label(frame, text="UF:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        uf = ttk.Combobox(frame, state="readonly", values=sorted(UF_CIDADES.keys()), width=8)
        uf.grid(row=0, column=1, sticky="ew")

        ttk.Label(frame, text="Cidade:").grid(row=0, column=2, sticky="w", padx=(14, 6))
        cidade = ttk.Combobox(frame, state="readonly")
        cidade.grid(row=0, column=3, sticky="ew")

        def atualizar_cidades(_evt=None):
            cidades = self._buscar_cidades_por_uf(uf.get())
            cidade["values"] = cidades
            cidade.set("")

        uf.bind("<<ComboboxSelected>>", atualizar_cidades)
        return uf, cidade

    def _linha_logradouro(self, parent, row):
        frame = ttk.Frame(parent)
        frame.grid(row=row, column=0, sticky="ew", padx=12, pady=2)
        frame.columnconfigure(1, weight=1)
        frame.columnconfigure(3, weight=2)

        ttk.Label(frame, text="Rua:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        tipo = ttk.Combobox(frame, state="readonly", values=TIPOS_LOGRADOURO, width=14)
        tipo.grid(row=0, column=1, sticky="ew")

        ttk.Label(frame, text="Nome:").grid(row=0, column=2, sticky="w", padx=(14, 6))
        nome = ttk.Entry(frame)
        nome.grid(row=0, column=3, sticky="ew")
        return tipo, nome

    def _attach_calendar(
        self,
        container,
        entry_widget,
        row=0,
        column=2,
        padx=(8, 0),
        button_text="Calendario",
        width=None,
    ):
        kwargs = {
            "text": button_text,
            "command": lambda: CalendarioPopup(container.winfo_toplevel(), entry_widget),
        }
        if width is not None:
            kwargs["width"] = width
        ttk.Button(container, **kwargs).grid(row=row, column=column, padx=padx)
        entry_widget.bind(
            "<Double-Button-1>",
            lambda _e: CalendarioPopup(container.winfo_toplevel(), entry_widget),
        )

    @staticmethod
    def _set_date_placeholder(entry_widget):
        placeholder = "dd/mm/aaaa"
        entry_widget.insert(0, placeholder)

        def on_focus_in(_event=None):
            if entry_widget.get() == placeholder:
                entry_widget.delete(0, tk.END)

        def on_focus_out(_event=None):
            if not entry_widget.get().strip():
                entry_widget.insert(0, placeholder)

        entry_widget.bind("<FocusIn>", on_focus_in)
        entry_widget.bind("<FocusOut>", on_focus_out)

    def abrir_cadastro_empresa(self):
        popup = CadastroPopup(self, "Cadastro de Empresa")
        popup.columnconfigure(0, weight=1)

        r = 0
        popup.secao("Qual Empresa?", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Empresa:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nome_pasta = ttk.Entry(f)
        nome_pasta.grid(row=0, column=1, sticky="ew")

        r += 1
        popup.secao("Insira uma Logo", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Logo:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        logo_var = tk.StringVar()
        ttk.Entry(f, textvariable=logo_var, state="readonly").grid(row=0, column=1, sticky="ew")
        preview_logo = ttk.Label(f, text="Pre-visualizacao indisponivel")
        preview_logo.grid(row=1, column=1, sticky="w", pady=(6, 0))
        logo_origem = {"path": ""}

        def escolher_logo():
            caminho = filedialog.askopenfilename(
                title="Selecionar logo da empresa",
                filetypes=[
                    ("Imagens", "*.png *.jpg *.jpeg *.bmp *.gif *.webp"),
                    ("Todos os arquivos", "*.*"),
                ],
            )
            if not caminho:
                return
            logo_origem["path"] = caminho
            logo_var.set(caminho)
            self._atualizar_preview_logo(preview_logo, caminho)

        ttk.Button(f, text="Procurar...", command=escolher_logo).grid(row=0, column=2, padx=(8, 0))

        r += 1
        popup.secao("Nome da Empresa", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Razao Social:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nome_empresa = ttk.Entry(f)
        nome_empresa.grid(row=0, column=1, sticky="ew")

        r += 1
        popup.secao("CNPJ", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="CNPJ:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        cnpj = ttk.Entry(f)
        cnpj.grid(row=0, column=1, sticky="ew")
        self._bind_mask(cnpj, mascara_cnpj)

        r += 1
        popup.secao("Email", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="email:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        email = ttk.Entry(f)
        email.grid(row=0, column=1, sticky="ew")

        r += 1
        popup.secao("nº do celular", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Contato:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        contato = ttk.Entry(f)
        contato.grid(row=0, column=1, sticky="ew")
        self._bind_mask(contato, mascara_celular)

        r += 1
        popup.secao("Endereco da Empresa", r)
        r += 1
        tipo_log, nome_log = self._linha_logradouro(popup, r)

        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12, pady=2)
        f.columnconfigure(1, weight=1)
        f.columnconfigure(3, weight=1)
        ttk.Label(f, text="nº:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        numero = ttk.Entry(f)
        numero.grid(row=0, column=1, sticky="ew")
        ttk.Label(f, text="Complemento:").grid(row=0, column=2, sticky="w", padx=(14, 6))
        complemento = ttk.Entry(f)
        complemento.grid(row=0, column=3, sticky="ew")

        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12, pady=2)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Bairro:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        bairro = ttk.Entry(f)
        bairro.grid(row=0, column=1, sticky="ew")

        r += 1
        uf, cidade = self._combo_uf_cidade(popup, r)

        r += 1
        botoes = ttk.Frame(popup)
        botoes.grid(row=r, column=0, sticky="e", padx=12, pady=18)

        def salvar():
            nome_empresa_txt = nome_empresa.get().strip()
            if not nome_empresa_txt:
                messagebox.showerror("Validacao", "Preencha a Razao Social.")
                return
            nome_pasta_txt = nome_pasta.get().strip()
            logo_rel = ""
            if logo_origem["path"]:
                try:
                    logo_rel = self._salvar_logo_empresa(
                        logo_origem["path"],
                        self.prox_empresa_id,
                        nome_pasta_txt or nome_empresa_txt,
                    )
                except OSError as exc:
                    messagebox.showerror("Validacao", f"Nao foi possivel salvar a logo: {exc}")
                    return
            empresa = {
                "id": self.prox_empresa_id,
                "nome_pasta": nome_pasta_txt,
                "logo": logo_rel,
                "nome": nome_empresa_txt,
                "cnpj": cnpj.get().strip(),
                "email": email.get().strip(),
                "contato": contato.get().strip(),
                "logradouro_tipo": tipo_log.get().strip(),
                "logradouro_nome": nome_log.get().strip(),
                "numero": numero.get().strip(),
                "complemento": complemento.get().strip(),
                "bairro": bairro.get().strip(),
                "uf": uf.get().strip(),
                "cidade": cidade.get().strip(),
            }
            self.prox_empresa_id += 1
            self.empresas.append(empresa)
            self._salvar_dados()
            self._atualizar_select_empresas(preservar_id=empresa["id"])
            messagebox.showinfo("Cadastro", "Empresa cadastrada com sucesso.")
            popup.destroy()

        ttk.Button(botoes, text="Cancelar", command=popup.destroy).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(botoes, text="Salvar", command=salvar).grid(row=0, column=1)

        tipo_log.set("")
        uf.set("")
        cidade.set("")
        popup.ajustar_tamanho()
        nome_log.focus_set()

    def abrir_cadastro_funcao(self):
        popup = CadastroPopup(self, "Cadastro de Funcao")
        popup.columnconfigure(0, weight=1)

        r = 0
        popup.secao("Cadastrar Nova Funcao", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=2)
        f.columnconfigure(3, weight=1)
        ttk.Label(f, text="Funcao:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nova_funcao = ttk.Entry(f)
        nova_funcao.grid(row=0, column=1, sticky="ew")
        ttk.Label(f, text="CBO:").grid(row=0, column=2, sticky="w", padx=(12, 8))
        novo_cbo = ttk.Entry(f)
        novo_cbo.grid(row=0, column=3, sticky="ew")

        r += 1
        popup.secao("Editar Funcao", r)
        r += 1
        f_edit = ttk.Frame(popup)
        f_edit.grid(row=r, column=0, sticky="ew", padx=12)
        f_edit.columnconfigure(1, weight=2)
        f_edit.columnconfigure(3, weight=2)
        f_edit.columnconfigure(5, weight=1)
        ttk.Label(f_edit, text="Atual:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        funcao_atual = ttk.Combobox(f_edit, state="readonly")
        funcao_atual.grid(row=0, column=1, sticky="ew")
        ttk.Label(f_edit, text="Nova Funcao:").grid(row=0, column=2, sticky="w", padx=(12, 8))
        funcao_nova = ttk.Entry(f_edit)
        funcao_nova.grid(row=0, column=3, sticky="ew")
        ttk.Label(f_edit, text="Novo CBO:").grid(row=0, column=4, sticky="w", padx=(12, 8))
        cbo_novo_edit = ttk.Entry(f_edit)
        cbo_novo_edit.grid(row=0, column=5, sticky="ew")

        r += 1
        popup.secao("Excluir Funcao", r)
        r += 1
        f_del = ttk.Frame(popup)
        f_del.grid(row=r, column=0, sticky="ew", padx=12)
        f_del.columnconfigure(1, weight=1)
        ttk.Label(f_del, text="Funcao:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        funcao_excluir = ttk.Combobox(f_del, state="readonly")
        funcao_excluir.grid(row=0, column=1, sticky="ew")

        def atualizar_listas():
            valores = sorted(FUNCOES_CBO.keys())
            funcao_atual["values"] = valores
            funcao_excluir["values"] = valores
            if valores and not funcao_atual.get():
                funcao_atual.current(0)
            if valores and not funcao_excluir.get():
                funcao_excluir.current(0)

        def cadastrar():
            nome_novo = nova_funcao.get().strip()
            cbo_novo = re.sub(r"\D", "", novo_cbo.get())
            if not nome_novo:
                messagebox.showerror("Validacao", "Preencha a nova funcao.")
                return
            if not cbo_novo:
                messagebox.showerror("Validacao", "Preencha o CBO da nova funcao.")
                return
            FUNCOES_CBO[nome_novo] = cbo_novo
            atualizar_listas()
            messagebox.showinfo("Cadastro", "Funcao cadastrada com sucesso.")
            nova_funcao.delete(0, tk.END)
            novo_cbo.delete(0, tk.END)

        def editar():
            atual = funcao_atual.get().strip()
            if not atual or atual not in FUNCOES_CBO:
                messagebox.showerror("Validacao", "Selecione a funcao atual.")
                return
            nome_novo = funcao_nova.get().strip() or atual
            cbo_novo = re.sub(r"\D", "", cbo_novo_edit.get()) or FUNCOES_CBO.get(atual, "")
            if not cbo_novo:
                messagebox.showerror("Validacao", "Preencha o novo CBO.")
                return
            if nome_novo != atual and nome_novo in FUNCOES_CBO:
                messagebox.showerror("Validacao", "Ja existe uma funcao com esse nome.")
                return
            if nome_novo != atual:
                del FUNCOES_CBO[atual]
            FUNCOES_CBO[nome_novo] = cbo_novo
            for func in self.funcionarios:
                if str(func.get("funcao", "")).strip() == atual:
                    func["funcao"] = nome_novo
                    func["cbo"] = cbo_novo
            self._salvar_dados()
            atualizar_listas()
            funcao_atual.set(nome_novo)
            funcao_excluir.set(nome_novo)
            messagebox.showinfo("Cadastro", "Funcao atualizada com sucesso.")

        def excluir():
            nome_excluir = funcao_excluir.get().strip()
            if not nome_excluir:
                messagebox.showerror("Validacao", "Selecione a funcao para excluir.")
                return
            if nome_excluir not in FUNCOES_CBO:
                messagebox.showerror("Validacao", "Funcao nao encontrada.")
                return
            if not messagebox.askyesno("Confirmar", f"Excluir a funcao '{nome_excluir}'?"):
                return
            del FUNCOES_CBO[nome_excluir]
            for func in self.funcionarios:
                if str(func.get("funcao", "")).strip() == nome_excluir:
                    func["funcao"] = ""
                    func["cbo"] = ""
            self._salvar_dados()
            atualizar_listas()
            funcao_atual.set("")
            funcao_excluir.set("")
            messagebox.showinfo("Cadastro", "Funcao excluida com sucesso.")

        r += 1
        botoes = ttk.Frame(popup)
        botoes.grid(row=r, column=0, sticky="e", padx=12, pady=18)
        ttk.Button(botoes, text="Cadastrar", command=cadastrar).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(botoes, text="Editar", command=editar).grid(row=0, column=1, padx=(0, 8))
        ttk.Button(botoes, text="Excluir", command=excluir).grid(row=0, column=2, padx=(0, 8))
        ttk.Button(botoes, text="Fechar", command=popup.destroy).grid(row=0, column=3)

        atualizar_listas()
        popup.ajustar_tamanho()

    def abrir_cadastro_documento(self):
        popup = CadastroPopup(self, "Cadastro de Documento")
        popup.columnconfigure(0, weight=1)

        r = 0
        popup.secao("Qual Empresa?", r)
        r += 1
        f = ttk.Frame(popup)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Empresa:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        empresa_ids = [empresa["id"] for empresa in self.empresas]
        empresa_labels = [self._empresa_label(empresa) for empresa in self.empresas]
        select_empresa_doc = ttk.Combobox(f, state="readonly", values=empresa_labels)
        select_empresa_doc.grid(row=0, column=1, sticky="ew")

        r += 1
        popup.secao("Certificado", r)
        r += 1
        f_cert = ttk.Frame(popup)
        f_cert.grid(row=r, column=0, sticky="ew", padx=12)
        f_cert.columnconfigure(1, weight=1)
        ttk.Label(f_cert, text="Arquivo:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        certificado_arquivo_var = tk.StringVar()
        ttk.Entry(f_cert, textvariable=certificado_arquivo_var, state="readonly").grid(
            row=0, column=1, sticky="ew"
        )

        def escolher_certificado():
            caminho = filedialog.askopenfilename(
                title="Selecionar documento de Certificado",
                filetypes=[
                    ("Documentos", "*.docx *.doc *.pdf"),
                    ("Todos os arquivos", "*.*"),
                ],
            )
            if caminho:
                certificado_arquivo_var.set(caminho)

        ttk.Button(f_cert, text="Procurar...", command=escolher_certificado).grid(
            row=0, column=2, padx=(8, 0)
        )

        r += 1
        f_dias = ttk.Frame(popup)
        f_dias.grid(row=r, column=0, sticky="ew", padx=12, pady=(4, 0))
        ttk.Label(f_dias, text="Dias:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        certificado_dias = ttk.Entry(f_dias, width=8)
        certificado_dias.grid(row=0, column=1, sticky="w")
        certificado_dias.insert(0, "1")

        r += 1
        f_salvar_doc = ttk.Frame(popup)
        f_salvar_doc.grid(row=r, column=0, sticky="ew", padx=12, pady=(6, 0))
        f_salvar_doc.columnconfigure(0, weight=1)
        btn_salvar_doc = ttk.Button(f_salvar_doc, text="Salvar")
        btn_salvar_doc.grid(row=0, column=1, sticky="e")

        r += 1
        popup.secao("Atualizar NR", r)
        r += 1
        f_nr = ttk.Frame(popup)
        f_nr.grid(row=r, column=0, sticky="ew", padx=12, pady=(4, 0))
        f_nr.columnconfigure(1, weight=1)
        ttk.Label(f_nr, text="Atualizar NR:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nr_indices = []
        nr_labels = []
        select_nr_dias = ttk.Combobox(f_nr, state="readonly", values=nr_labels)
        select_nr_dias.grid(row=0, column=1, sticky="ew")
        r += 1
        f_nr_info = ttk.Frame(popup)
        f_nr_info.grid(row=r, column=0, sticky="ew", padx=12, pady=(2, 0))
        ttk.Label(f_nr_info, text="Dias da NR selecionada:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        dias_nr_selecionada_var = tk.StringVar(value="1")
        dias_nr_selecionada_entry = ttk.Entry(f_nr_info, textvariable=dias_nr_selecionada_var, width=8)
        dias_nr_selecionada_entry.grid(row=0, column=1, sticky="w")

        def atualizar_info_nr(*_args):
            idx_nr_sel = select_nr_dias.current()
            if 0 <= idx_nr_sel < len(nr_indices):
                idx_nr = nr_indices[idx_nr_sel]
                if 0 <= idx_nr < len(self.nr_certificados):
                    dias_nr_selecionada_var.set(str(self._duracao_nr_item(self.nr_certificados[idx_nr])))
                    return
            dias_nr_selecionada_var.set("1")

        def atualizar_nr_por_empresa(*_args):
            nr_indices.clear()
            nr_labels.clear()
            idx_empresa = select_empresa_doc.current()
            if 0 <= idx_empresa < len(empresa_ids):
                empresa_id_sel = empresa_ids[idx_empresa]
                nomes_nr_empresa = self._nomes_nr_na_pasta_empresa(empresa_id_sel) or {}
                chaves_empresa = set(nomes_nr_empresa.keys())
                for idx, item in enumerate(self.nr_certificados):
                    nome_nr = str(item.get("nome", "") or "").strip()
                    if not nome_nr:
                        continue
                    chave = self._normalizar_nome_nr(nome_nr)
                    if chave not in chaves_empresa:
                        continue
                    dias_nr = self._duracao_nr_item(item)
                    nr_indices.append(idx)
                    nr_labels.append(f"{nome_nr} ({dias_nr} dia(s))")
            select_nr_dias.configure(values=nr_labels)
            if nr_labels:
                select_nr_dias.current(0)
            atualizar_info_nr()

        def atualizar_certificado_por_empresa(*_args):
            idx_empresa = select_empresa_doc.current()
            dias_valor = 1
            if 0 <= idx_empresa < len(empresa_ids):
                empresa_id_sel = empresa_ids[idx_empresa]
                documento_empresa = next(
                    (
                        item for item in self.documentos
                        if isinstance(item, dict) and item.get("empresa_id") == empresa_id_sel
                    ),
                    None,
                )
                if documento_empresa:
                    try:
                        dias_valor = max(1, int(documento_empresa.get("certificado_dias", 1)))
                    except (TypeError, ValueError):
                        dias_valor = 1
            certificado_dias.delete(0, tk.END)
            certificado_dias.insert(0, str(dias_valor))
            certificado_arquivo_var.set("")

        select_nr_dias.bind("<<ComboboxSelected>>", atualizar_info_nr, add="+")
        select_empresa_doc.bind("<<ComboboxSelected>>", atualizar_nr_por_empresa, add="+")
        select_empresa_doc.bind("<<ComboboxSelected>>", atualizar_certificado_por_empresa, add="+")

        r += 1
        f_nr_botao = ttk.Frame(popup)
        f_nr_botao.grid(row=r, column=0, sticky="ew", padx=12, pady=(2, 0))
        f_nr_botao.columnconfigure(0, weight=1)

        def salvar_atualizacao_dias_nr(exibir_mensagem=True):
            dias_txt = re.sub(r"\D", "", dias_nr_selecionada_entry.get().strip())
            dias_valor = int(dias_txt) if dias_txt else 1
            dias_valor = max(1, dias_valor)
            dias_nr_selecionada_entry.delete(0, tk.END)
            dias_nr_selecionada_entry.insert(0, str(dias_valor))

            idx_nr_sel = select_nr_dias.current()
            if idx_nr_sel < 0 or idx_nr_sel >= len(nr_indices):
                if exibir_mensagem:
                    messagebox.showwarning("Cadastro", "Selecione uma NR para atualizar os dias.")
                return None
            idx_nr = nr_indices[idx_nr_sel]
            if not (0 <= idx_nr < len(self.nr_certificados)):
                if exibir_mensagem:
                    messagebox.showwarning("Cadastro", "NR selecionada invalida.")
                return None

            self.nr_certificados[idx_nr]["dias"] = dias_valor
            if idx_nr_sel < len(nr_labels):
                nome_nr = str(self.nr_certificados[idx_nr].get("nome", "") or "").strip()
                nr_labels[idx_nr_sel] = f"{nome_nr} ({dias_valor} dia(s))"
                select_nr_dias.configure(values=nr_labels)
                select_nr_dias.current(idx_nr_sel)
            dias_nr_selecionada_var.set(str(dias_valor))
            self._render_campos_nr()
            self._salvar_dados()
            if exibir_mensagem:
                messagebox.showinfo("Cadastro", "Dias da NR atualizados com sucesso.")
            return dias_valor

        def excluir_nr_selecionada():
            idx_nr_sel = select_nr_dias.current()
            if idx_nr_sel < 0 or idx_nr_sel >= len(nr_indices):
                messagebox.showwarning("Cadastro", "Selecione uma NR para excluir.")
                return
            idx_nr = nr_indices[idx_nr_sel]
            if not (0 <= idx_nr < len(self.nr_certificados)):
                messagebox.showwarning("Cadastro", "NR selecionada invalida.")
                return
            nome_nr = str(self.nr_certificados[idx_nr].get("nome", "") or "").strip() or "NR"
            if not messagebox.askyesno("Confirmar exclusao", f"Deseja excluir a NR '{nome_nr}'?"):
                return
            self._remover_linha_nr(idx_nr)
            atualizar_nr_por_empresa()
            messagebox.showinfo("Cadastro", "NR excluida com sucesso.")

        ttk.Button(
            f_nr_botao,
            text="Salvar",
            command=lambda: salvar_atualizacao_dias_nr(True),
        ).grid(row=0, column=1, sticky="e", padx=(0, 8))
        ttk.Button(
            f_nr_botao,
            text="Excluir NR",
            command=excluir_nr_selecionada,
        ).grid(row=0, column=2, sticky="e")

        r += 1
        popup.secao("Excluir OUTROS DOCUMENTOS", r)
        r += 1
        f_outros_exc = ttk.Frame(popup)
        f_outros_exc.grid(row=r, column=0, sticky="ew", padx=12, pady=(4, 0))
        f_outros_exc.columnconfigure(1, weight=1)
        ttk.Label(f_outros_exc, text="Documento:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        outros_docs_indices = []
        outros_docs_labels = []
        select_outro_excluir = ttk.Combobox(f_outros_exc, state="readonly", values=outros_docs_labels)
        select_outro_excluir.grid(row=0, column=1, sticky="ew")

        def atualizar_outros_por_empresa(*_args):
            outros_docs_indices.clear()
            outros_docs_labels.clear()
            idx_empresa = select_empresa_doc.current()
            if 0 <= idx_empresa < len(empresa_ids):
                empresa_id_sel = empresa_ids[idx_empresa]
                documento_empresa = next(
                    (
                        item for item in self.documentos
                        if isinstance(item, dict) and item.get("empresa_id") == empresa_id_sel
                    ),
                    None,
                )
                for tipo, caminho in self._listar_arquivos_outros_documentos(documento_empresa):
                    tipo_txt = str(tipo or "").strip()
                    caminho_txt = str(caminho or "").strip()
                    if not tipo_txt or not caminho_txt:
                        continue
                    nome_arq = Path(caminho_txt).name or caminho_txt
                    outros_docs_indices.append({"tipo": tipo_txt, "caminho": caminho_txt})
                    outros_docs_labels.append(f"{tipo_txt} - {nome_arq}")
            select_outro_excluir.configure(values=outros_docs_labels)
            if outros_docs_labels:
                select_outro_excluir.current(0)
            else:
                select_outro_excluir.set("")

        def excluir_outro_documento_cadastrado():
            idx_empresa = select_empresa_doc.current()
            if idx_empresa < 0 or idx_empresa >= len(empresa_ids):
                messagebox.showwarning("Cadastro", "Selecione uma empresa.")
                return
            idx_sel = select_outro_excluir.current()
            if idx_sel < 0 or idx_sel >= len(outros_docs_indices):
                messagebox.showwarning("Cadastro", "Selecione um documento para excluir.")
                return
            empresa_id_sel = empresa_ids[idx_empresa]
            doc_idx = next(
                (
                    i for i, item in enumerate(self.documentos)
                    if isinstance(item, dict) and item.get("empresa_id") == empresa_id_sel
                ),
                None,
            )
            if doc_idx is None:
                messagebox.showwarning("Cadastro", "Nenhum cadastro de documento encontrado para a empresa.")
                return

            item_sel = outros_docs_indices[idx_sel]
            tipo_sel = str(item_sel.get("tipo", "") or "").strip()
            caminho_sel = str(item_sel.get("caminho", "") or "").strip()
            nome_exibicao = Path(caminho_sel).name or caminho_sel
            if not messagebox.askyesno(
                "Confirmar exclusao",
                f"Deseja excluir '{tipo_sel} - {nome_exibicao}' do cadastro?",
            ):
                return

            documento = self.documentos[doc_idx]
            arquivos = documento.get("arquivos", {})
            if not isinstance(arquivos, dict):
                arquivos = {}
            valor_tipo = arquivos.get(tipo_sel, "")
            removido = False
            if isinstance(valor_tipo, list):
                nova_lista = []
                ja_removeu = False
                for caminho_item in valor_tipo:
                    caminho_txt = str(caminho_item or "").strip()
                    if (not ja_removeu) and caminho_txt == caminho_sel:
                        ja_removeu = True
                        removido = True
                        continue
                    if caminho_txt:
                        nova_lista.append(caminho_txt)
                if nova_lista:
                    arquivos[tipo_sel] = nova_lista
                else:
                    arquivos.pop(tipo_sel, None)
            else:
                caminho_txt = str(valor_tipo or "").strip()
                if caminho_txt == caminho_sel:
                    arquivos.pop(tipo_sel, None)
                    removido = True

            if not removido:
                messagebox.showwarning("Cadastro", "Documento selecionado nao foi encontrado no cadastro.")
                atualizar_outros_por_empresa()
                return

            documento["arquivos"] = arquivos
            documento["tipos"] = [
                tipo for tipo in OUTROS_DOCUMENTOS_TIPOS
                if tipo in arquivos and (
                    (isinstance(arquivos.get(tipo), list) and any(str(x or "").strip() for x in arquivos.get(tipo)))
                    or (not isinstance(arquivos.get(tipo), list) and str(arquivos.get(tipo) or "").strip())
                )
            ]
            if not str(documento.get("certificado_arquivo", "") or "").strip() and not self._listar_arquivos_outros_documentos(documento):
                del self.documentos[doc_idx]
            else:
                self.documentos[doc_idx] = documento

            self._salvar_dados()
            self._carregar_outros_documentos_empresa_selecionada()
            atualizar_outros_por_empresa()
            messagebox.showinfo("Cadastro", "Documento excluido com sucesso.")

        ttk.Button(
            f_outros_exc,
            text="Excluir",
            command=excluir_outro_documento_cadastrado,
        ).grid(row=0, column=2, sticky="e", padx=(8, 0))
        select_empresa_doc.bind("<<ComboboxSelected>>", atualizar_outros_por_empresa, add="+")

        r += 1
        botoes = ttk.Frame(popup)
        botoes.grid(row=r, column=0, sticky="ew", padx=12, pady=(28, 18))
        botoes.columnconfigure(0, weight=1)

        def salvar():
            idx_empresa = select_empresa_doc.current()
            if idx_empresa < 0 or idx_empresa >= len(empresa_ids):
                messagebox.showerror("Validacao", "Selecione uma empresa.")
                return
            empresa_id_sel = empresa_ids[idx_empresa]
            empresa_sel = next((e for e in self.empresas if e.get("id") == empresa_id_sel), None)
            if not empresa_sel:
                messagebox.showerror("Validacao", "Empresa selecionada nao encontrada.")
                return
            nome_empresa_arquivo = str(
                empresa_sel.get("nome_pasta", "") or empresa_sel.get("nome", "")
            ).strip()
            if not nome_empresa_arquivo:
                nome_empresa_arquivo = empresa_labels[idx_empresa]
            indices_empresa = [
                i for i, item in enumerate(self.documentos)
                if isinstance(item, dict) and item.get("empresa_id") == empresa_id_sel
            ]
            documento_existente = self.documentos[indices_empresa[0]] if indices_empresa else None
            arquivo_cert = certificado_arquivo_var.get().strip()
            dias_txt = re.sub(r"\D", "", certificado_dias.get().strip())
            dias_valor = int(dias_txt) if dias_txt else 1
            dias_valor = max(1, dias_valor)
            certificado_dias.delete(0, tk.END)
            certificado_dias.insert(0, str(dias_valor))

            arquivos_existentes = (documento_existente or {}).get("arquivos", {})
            if not isinstance(arquivos_existentes, dict):
                arquivos_existentes = {}
            arquivos_outros = {}
            for k, v in arquivos_existentes.items():
                nome_tipo = str(k or "").strip()
                if not nome_tipo:
                    continue
                if isinstance(v, list):
                    caminhos_tipo = [
                        str(item or "").strip()
                        for item in v
                        if str(item or "").strip()
                    ]
                    if caminhos_tipo:
                        arquivos_outros[nome_tipo] = caminhos_tipo
                else:
                    caminho_tipo = str(v or "").strip()
                    if caminho_tipo:
                        arquivos_outros[nome_tipo] = caminho_tipo

            certificado_atual = str(
                (documento_existente or {}).get("certificado_arquivo", "") or ""
            ).strip()
            tipo_outro_novo = ""
            arquivo_salvo_rel = ""

            if arquivo_cert:
                tipo_outro_detectado = self._detectar_tipo_outro_documento(Path(arquivo_cert).stem)
                tipo_para_registro = tipo_outro_detectado or "Certificado"
                try:
                    arquivo_salvo_rel = self._salvar_documento_no_projeto(
                        arquivo_cert,
                        nome_empresa_arquivo,
                        tipo_para_registro,
                        empresa_id=empresa_id_sel,
                    )
                except OSError as exc:
                    messagebox.showerror("Validacao", f"Nao foi possivel salvar o documento: {exc}")
                    return

                tipo_outro_novo = tipo_outro_detectado
                if tipo_outro_novo:
                    atual_tipo = arquivos_outros.get(tipo_outro_novo, [])
                    if isinstance(atual_tipo, list):
                        lista_tipo = list(atual_tipo)
                    else:
                        valor_atual = str(atual_tipo or "").strip()
                        lista_tipo = [valor_atual] if valor_atual else []
                    if arquivo_salvo_rel not in lista_tipo:
                        lista_tipo.append(arquivo_salvo_rel)
                    arquivos_outros[tipo_outro_novo] = lista_tipo
                else:
                    certificado_atual = arquivo_salvo_rel
            else:
                if not certificado_atual and not arquivos_outros:
                    messagebox.showwarning(
                        "Cadastro",
                        "Selecione um arquivo para cadastrar.",
                    )
                    return

            documento = {
                "empresa_id": empresa_id_sel,
                "certificado_arquivo": certificado_atual,
                "certificado_dias": dias_valor,
                "tipos": list(arquivos_outros.keys()),
                "arquivos": arquivos_outros,
            }
            if indices_empresa:
                idx_base = indices_empresa[0]
                documento["id"] = self.documentos[idx_base].get("id", self.prox_documento_id)
                self.documentos[idx_base] = documento
                # Remove duplicados antigos da mesma empresa, mantendo apenas o registro atualizado.
                for idx_dup in reversed(indices_empresa[1:]):
                    del self.documentos[idx_dup]
            else:
                documento["id"] = self.prox_documento_id
                self.prox_documento_id += 1
                self.documentos.append(documento)

            # Ao cadastrar uma NR por arquivo, persiste os dias na lista de NRs
            # somente quando o arquivo foi tratado como Certificado (nao OUTROS DOCUMENTOS).
            caminho_nr_ref = arquivo_salvo_rel or certificado_atual
            if (not tipo_outro_novo) and caminho_nr_ref:
                nome_nr_ref = self._nome_nr_do_arquivo(Path(caminho_nr_ref).stem, nome_empresa_arquivo)
                if re.match(r"(?i)^nr\b", str(nome_nr_ref or "").strip()):
                    chave_nr_ref = self._normalizar_nome_nr(self._nome_nr_canonico(nome_nr_ref))
                    idx_nr_existente = next(
                        (
                            idx_nr
                            for idx_nr, item_nr in enumerate(self.nr_certificados)
                            if self._normalizar_nome_nr(
                                self._nome_nr_canonico(item_nr.get("nome", ""))
                            ) == chave_nr_ref
                        ),
                        None,
                    )
                    if idx_nr_existente is None:
                        self.nr_certificados.append(
                            {
                                "nome": nome_nr_ref,
                                "coluna_1": "",
                                "coluna_2": "",
                                "dias": dias_valor,
                                "reciclagem": False,
                                "imprimir": False,
                                "imprimir_adicionado": False,
                            }
                        )
                    else:
                        self.nr_certificados[idx_nr_existente]["nome"] = nome_nr_ref
                        self.nr_certificados[idx_nr_existente]["dias"] = dias_valor
                    self.nr_certificados.sort(key=self._chave_ordenacao_nr)

            self._salvar_dados()
            empresa_main_id = self._empresa_id_selecionada_main()
            if empresa_main_id is not None:
                self._aplicar_filtro_nr_por_empresa(empresa_main_id)
            else:
                self._render_campos_nr()
                self._atualizar_lista_nr_imprimir()
            self._carregar_outros_documentos_empresa_selecionada()
            atualizar_outros_por_empresa()
            certificado_arquivo_var.set("")
            if tipo_outro_novo:
                messagebox.showinfo(
                    "Cadastro",
                    f"Documento cadastrado com sucesso em OUTROS DOCUMENTOS ({tipo_outro_novo}).",
                )
            else:
                messagebox.showinfo(
                    "Cadastro",
                    "Documento(s) cadastrado(s) com sucesso.",
                )

        ttk.Button(botoes, text="Fechar", command=popup.destroy).grid(row=0, column=0, sticky="e")
        btn_salvar_doc.configure(command=salvar)

        if empresa_labels:
            select_empresa_doc.current(0)
            atualizar_nr_por_empresa()
            atualizar_certificado_por_empresa()
            atualizar_outros_por_empresa()
        popup.ajustar_tamanho()
        select_empresa_doc.focus_set()

    def abrir_cadastro_funcionario(self):
        if not self.empresas:
            messagebox.showwarning("Cadastro", "Cadastre uma empresa antes de cadastrar funcionario.")
            return

        popup = CadastroPopup(self, "Cadastro de Funcionario")
        popup.columnconfigure(0, weight=1)
        form_parent = popup.habilitar_rolagem()
        form_parent.columnconfigure(0, weight=1)

        r = 0
        popup.secao("Nome da Pasta", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Nome da Pasta:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nome_pasta_valores = []
        for empresa_ref in self.empresas:
            valor_pasta = str(empresa_ref.get("nome_pasta") or empresa_ref.get("nome") or "").strip()
            if valor_pasta and valor_pasta not in nome_pasta_valores:
                nome_pasta_valores.append(valor_pasta)
        nome_pasta = ttk.Combobox(f, state="readonly", values=nome_pasta_valores)
        nome_pasta.grid(row=0, column=1, sticky="ew")
        empresa_id_main = self._empresa_id_selecionada_main()
        empresa_main = next((e for e in self.empresas if e["id"] == empresa_id_main), None)
        if empresa_main:
            nome_pasta_main = str(empresa_main.get("nome_pasta") or empresa_main.get("nome") or "").strip()
            if nome_pasta_main:
                nome_pasta.set(nome_pasta_main)
        elif nome_pasta_valores:
            nome_pasta.current(0)

        r += 1
        popup.secao("Nome do Funcionario", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Nome:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nome = ttk.Entry(f)
        nome.grid(row=0, column=1, sticky="ew")
        ttk.Label(f, text="Foto:").grid(row=1, column=0, sticky="w", padx=(0, 8), pady=(6, 0))
        foto_var = tk.StringVar()
        ttk.Entry(f, textvariable=foto_var, state="readonly").grid(row=1, column=1, sticky="ew", pady=(6, 0))
        preview_foto = ttk.Label(f, text="Pre-visualizacao indisponivel")
        preview_foto.grid(row=2, column=1, sticky="w", pady=(6, 0))
        foto_origem = {"path": ""}

        def escolher_foto_funcionario():
            caminho = filedialog.askopenfilename(
                title="Selecionar foto do funcionario",
                filetypes=[
                    ("Imagens", "*.png *.jpg *.jpeg *.bmp *.gif *.webp"),
                    ("Todos os arquivos", "*.*"),
                ],
            )
            if caminho:
                foto_origem["path"] = caminho
                foto_var.set(caminho)
                self._atualizar_preview_logo(preview_foto, caminho)

        ttk.Button(f, text="Procurar...", command=escolher_foto_funcionario).grid(
            row=1, column=2, padx=(8, 0), pady=(6, 0)
        )

        r += 1
        popup.secao("Insira o nº do CPF", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="CPF:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        cpf = ttk.Entry(f)
        cpf.grid(row=0, column=1, sticky="ew")
        self._bind_mask(cpf, mascara_cpf)

        r += 1
        popup.secao("Insira o nº do RG", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="RG:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        rg = ttk.Entry(f)
        rg.grid(row=0, column=1, sticky="ew")

        r += 1
        popup.secao("Insira a data de nascimento", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Nascimento:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        nascimento = ttk.Entry(f)
        nascimento.grid(row=0, column=1, sticky="ew")
        self._set_date_placeholder(nascimento)
        self._bind_mask(nascimento, mascara_data)
        self._attach_calendar(f, nascimento)

        r += 1
        popup.secao("nº do celular", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Contato:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        celular = ttk.Entry(f)
        celular.grid(row=0, column=1, sticky="ew")
        self._bind_mask(celular, mascara_celular)

        r += 1
        popup.secao("Endereco do Funcionario", r, parent=form_parent)
        r += 1
        tipo_log, nome_log = self._linha_logradouro(form_parent, r)

        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12, pady=2)
        f.columnconfigure(1, weight=1)
        f.columnconfigure(3, weight=1)
        ttk.Label(f, text="nº:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        numero = ttk.Entry(f)
        numero.grid(row=0, column=1, sticky="ew")
        ttk.Label(f, text="Complemento:").grid(row=0, column=2, sticky="w", padx=(14, 6))
        complemento = ttk.Entry(f)
        complemento.grid(row=0, column=3, sticky="ew")

        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12, pady=2)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Bairro:").grid(row=0, column=0, sticky="w", padx=(0, 6))
        bairro = ttk.Entry(f)
        bairro.grid(row=0, column=1, sticky="ew")

        r += 1
        uf, cidade = self._combo_uf_cidade(form_parent, r)

        r += 1
        popup.secao("Selecione a data de admissão", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Admissão:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        admissao = ttk.Entry(f)
        admissao.grid(row=0, column=1, sticky="ew")
        self._set_date_placeholder(admissao)
        self._bind_mask(admissao, mascara_data)
        self._attach_calendar(f, admissao)

        r += 1
        popup.secao("Selecione o cargo do funcionario", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Função:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        funcao = ttk.Combobox(f, state="readonly", values=sorted(FUNCOES_CBO.keys()))
        funcao.grid(row=0, column=1, sticky="ew")

        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="CBO:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        cbo_var = tk.StringVar(value="")
        cbo = ttk.Entry(f, textvariable=cbo_var, state="readonly")
        cbo.grid(row=0, column=1, sticky="ew")

        def atualizar_cbo(_event=None):
            cbo_var.set(FUNCOES_CBO.get(funcao.get(), ""))

        funcao.bind("<<ComboboxSelected>>", atualizar_cbo)

        r += 1
        popup.secao("Insira o valor do salário", r, parent=form_parent)
        r += 1
        f = ttk.Frame(form_parent)
        f.grid(row=r, column=0, sticky="ew", padx=12)
        f.columnconfigure(1, weight=1)
        ttk.Label(f, text="Salário:").grid(row=0, column=0, sticky="w", padx=(0, 8))
        salario = ttk.Entry(f)
        salario.grid(row=0, column=1, sticky="ew")
        salario.insert(0, "R$ 0,00")
        self._bind_mask(salario, mascara_moeda_br)

        r += 1
        botoes = ttk.Frame(form_parent)
        botoes.grid(row=r, column=0, sticky="e", padx=12, pady=18)

        def salvar():
            nome_pasta_txt = nome_pasta.get().strip()
            if not nome_pasta_txt:
                messagebox.showerror("Validacao", "Selecione a empresa.")
                return
            if not nome.get().strip():
                messagebox.showerror("Validacao", "Preencha o nome do funcionario.")
                return
            if len(re.sub(r"\D", "", cpf.get())) != 11:
                messagebox.showerror("Validacao", "CPF invalido.")
                return
            empresa_ref = next(
                (
                    e
                    for e in self.empresas
                    if str(e.get("nome_pasta") or e.get("nome") or "").strip() == nome_pasta_txt
                ),
                None,
            )
            if not empresa_ref:
                messagebox.showerror("Validacao", "Empresa nao encontrada para o item selecionado.")
                return
            empresa_id = empresa_ref["id"]
            foto_rel = ""
            if foto_origem["path"]:
                try:
                    foto_rel = self._salvar_foto_funcionario(
                        foto_origem["path"],
                        self.prox_funcionario_id,
                        nome.get().strip(),
                    )
                except OSError as exc:
                    messagebox.showerror("Validacao", f"Nao foi possivel salvar a foto: {exc}")
                    return
            funcionario = {
                "id": self.prox_funcionario_id,
                "nome": nome.get().strip(),
                "nome_pasta": nome_pasta_txt,
                "empresa_id": empresa_id,
                "foto": foto_rel,
                "cpf": cpf.get().strip(),
                "rg": rg.get().strip(),
                "nascimento": nascimento.get().strip(),
                "contato": celular.get().strip(),
                "logradouro_tipo": tipo_log.get().strip(),
                "logradouro_nome": nome_log.get().strip(),
                "numero": numero.get().strip(),
                "complemento": complemento.get().strip(),
                "bairro": bairro.get().strip(),
                "uf": uf.get().strip(),
                "cidade": cidade.get().strip(),
                "admissao": admissao.get().strip(),
                "funcao": funcao.get().strip(),
                "cbo": cbo_var.get().strip(),
                "salario": salario.get().strip(),
            }
            self.prox_funcionario_id += 1
            self.funcionarios.append(funcionario)
            self._salvar_dados()
            self._atualizar_select_empresas(preservar_id=empresa_id)
            messagebox.showinfo("Cadastro", "Funcionario cadastrado com sucesso.")
            popup.destroy()

        ttk.Button(botoes, text="Cancelar", command=popup.destroy).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(botoes, text="Salvar", command=salvar).grid(row=0, column=1)

        tipo_log.set("")
        uf.set("")
        cidade.set("")
        popup.ajustar_tamanho()
        nome.focus_set()

    def _bind_mask(self, entry_widget, mask_fn):
        def aplicar(_event=None):
            self._apply_mask(entry_widget, mask_fn)

        entry_widget.bind("<KeyRelease>", aplicar)
        entry_widget.bind("<<Paste>>", lambda _e: entry_widget.after_idle(aplicar))

    @staticmethod
    def _apply_mask(entry_widget, mask_fn):
        texto_atual = entry_widget.get()
        cursor_atual = entry_widget.index(tk.INSERT)
        digitos_antes_cursor = len(re.sub(r"\D", "", texto_atual[:cursor_atual]))

        texto_mascarado = mask_fn(texto_atual)
        if texto_mascarado == texto_atual:
            return

        entry_widget.delete(0, tk.END)
        entry_widget.insert(0, texto_mascarado)

        if digitos_antes_cursor <= 0:
            nova_pos = 0
        else:
            cont = 0
            nova_pos = len(texto_mascarado)
            for i, ch in enumerate(texto_mascarado):
                if ch.isdigit():
                    cont += 1
                    if cont >= digitos_antes_cursor:
                        nova_pos = i + 1
                        break

        entry_widget.icursor(nova_pos)


if __name__ == "__main__":
    app = App()
    app.mainloop()
